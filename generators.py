"""
Document generators for NY Surrogate's Court Probate HQ
Generates filled Word docs and PDFs from case data
"""

import io
import os
import re
import traceback
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
import fitz

TEMPLATES_DIR       = os.path.join(os.path.dirname(__file__), "templates")
ADMIN_TEMPLATES_DIR = os.path.join(TEMPLATES_DIR, "Admin")
PROBATE_TEMPLATES_DIR = os.path.join(TEMPLATES_DIR, "Probate")
WORD_TEMPLATES_DIR  = os.path.join(TEMPLATES_DIR, "Not Using Word Docs")
PDFS_DIR            = os.path.dirname(__file__)

COUNTY_INFO = {
    "Bronx": {
        "address": "851 Grand Concourse, 3rd Floor",
        "city_state_zip": "Bronx, NY 10451",
        "dept_probate": "Probate Department",
        "dept_admin": "Administration Department",
    },
    "Kings": {
        "address": "2 Johnson Street",
        "city_state_zip": "Brooklyn, NY 11201",
        "dept_probate": "Probate Department",
        "dept_admin": "Administration Department",
    },
    "Nassau": {
        "address": "262 Old Country Road",
        "city_state_zip": "Mineola, NY 11501",
        "dept_probate": "Probate Department",
        "dept_admin": "Administration Department",
    },
    "New York": {
        "address": "31 Chambers Street",
        "city_state_zip": "New York, NY 10007",
        "dept_probate": "Probate Department",
        "dept_admin": "Administration Department",
    },
    "Queens": {
        "address": "88-11 Sutphin Blvd",
        "city_state_zip": "Jamaica, NY 11435",
        "dept_probate": "Probate Department",
        "dept_admin": "Administration Department",
    },
    "Richmond": {
        "address": "18 Richmond Terrace",
        "city_state_zip": "Staten Island, NY 10301",
        "dept_probate": "Probate Department",
        "dept_admin": "Administration Department",
    },
    "Suffolk": {
        "address": "320 Center Drive",
        "city_state_zip": "Riverhead, NY 11901",
        "dept_probate": "Probate Department",
        "dept_admin": "Administration Department",
    },
}

SIGNERS = {
    "Jessica Wilson": "Jessica Wilson, Esq.",
    "Robyn Foresta": "Robyn Foresta, Legal Assistant",
}

# Relationship keywords that trigger Rule 207.16(c) on their own
# (grandparents, aunts/uncles, first cousins, first cousins once removed)
_DISTANT_REL_KEYWORDS = [
    "grandparent", "grandfather", "grandmother",
    "aunt", "uncle",
    "cousin",
]


def needs_family_tree_affidavit(data):
    """Return True if Rule 207.16(c) requires an Affidavit of Family Tree.

    Triggers when:
      - 0 or 1 distributee survives, OR
      - any distributee's relationship is grandparents, aunts/uncles,
        first cousins, or first cousins once removed.
    """
    dists = [d for d in data.get("distributees", []) if d.get("name")]
    if len(dists) <= 1:
        return True
    for d in dists:
        rel = (d.get("relationship") or "").lower()
        if any(k in rel for k in _DISTANT_REL_KEYWORDS):
            return True
    return False


def needs_family_tree_diagram(data):
    """Return True if the FT-1 diagram is also required (Rule 207.16(c)).

    The diagram is NOT required when the sole distributee is the spouse
    or only child of the decedent.
    """
    if not needs_family_tree_affidavit(data):
        return False
    dists = [d for d in data.get("distributees", []) if d.get("name")]
    if len(dists) == 1:
        rel = (dists[0].get("relationship") or "").lower()
        if any(k in rel for k in ["spouse", "child", "son", "daughter"]):
            return False
    return True


def family_tree_trigger_reason(data):
    """Return a short human-readable string explaining why 207.16(c) fired
    (used in the case summary).  Returns empty string if not triggered."""
    dists = [d for d in data.get("distributees", []) if d.get("name")]
    if len(dists) == 0:
        return "no distributees"
    if len(dists) == 1:
        return f"only one distributee ({dists[0].get('name', '')})"
    for d in dists:
        rel = (d.get("relationship") or "").lower()
        if any(k in rel for k in _DISTANT_REL_KEYWORDS):
            return f"distributee relationship: {d.get('relationship', '')}"
    return ""


def today():
    return datetime.now().strftime("%B %d, %Y")


def format_date_long(date_str):
    """Convert MM/DD/YYYY or YYYY-MM-DD to 'Month DD, YYYY'
    (e.g. '03/15/1945' → 'March 15, 1945')."""
    for fmt in ("%m/%d/%Y", "%Y-%m-%d"):
        try:
            return datetime.strptime(date_str, fmt).strftime("%B %d, %Y")
        except Exception:
            continue
    return date_str


def nonzero(v):
    """Return v only if it's a non-empty, non-zero value."""
    s = str(v or "").strip()
    return s if s and s not in ("0", "0.0", "0.00") else ""


def replace_in_doc(doc, replacements):
    """Replace placeholder text throughout a Word document.
    Handles placeholders split across multiple runs and multiple occurrences
    of the same placeholder within a paragraph."""
    def replace_in_para(para):
        for key, value in replacements.items():
            if key not in para.text:
                continue
            # First pass: replace within each run that contains the key.
            # (Don't break — the same key may appear in multiple runs.)
            replaced_in_run = False
            for run in para.runs:
                if key in run.text:
                    run.text = run.text.replace(key, value or "")
                    replaced_in_run = True
            # If the key still appears (e.g., split across runs), consolidate
            # the full paragraph text into the first run.
            if key in para.text:
                full_text = para.text.replace(key, value or "")
                if para.runs:
                    para.runs[0].text = full_text
                    for run in para.runs[1:]:
                        run.text = ""
                else:
                    para.add_run(full_text)

    for para in doc.paragraphs:
        replace_in_para(para)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    replace_in_para(para)


def replace_para(para, old_text, new_text):
    """Replace text within a paragraph's runs, preserving formatting."""
    full = para.text
    if old_text not in full and old_text != full:
        return
    for run in para.runs:
        if old_text in run.text:
            run.text = run.text.replace(old_text, new_text)
            return
    new_full = full.replace(old_text, new_text)
    for run in para.runs:
        run.text = ""
    if para.runs:
        para.runs[0].text = new_full
    else:
        para.add_run(new_full)


def _validate_docx(doc, generator_name):
    """Scan a generated Word doc for unreplaced [...] placeholder tokens
    and print warnings to the server log.

    Soft check — some bracketed text in templates is intentional (e.g.
    instructional [Name] / [Note: ...] markers meant for the end-user
    to fill in by hand), so warnings don't fail generation. They just
    give us a fast signal when a substitution path was missed.
    """
    token_rx = re.compile(r'\[[^\[\]\n]{1,80}\]')
    # Checkbox furniture on official court forms — "[    ]" / "[ X ]" — is
    # intentional, not an unreplaced placeholder.
    checkbox_rx = re.compile(r'\[\s*X?\s*\]')
    found = []

    def _scan(p, location):
        for m in token_rx.finditer(p.text or ""):
            if checkbox_rx.fullmatch(m.group(0)):
                continue
            found.append((location, m.group(0)))

    for pi, p in enumerate(doc.paragraphs):
        _scan(p, f"para[{pi}]")
    for ti, t in enumerate(doc.tables):
        for ri, row in enumerate(t.rows):
            for ci, cell in enumerate(row.cells):
                for pi, p in enumerate(cell.paragraphs):
                    _scan(p, f"table[{ti}].r{ri}c{ci}.p{pi}")

    if found:
        unique = sorted(set(found))
        print(f"[VALIDATE] {generator_name}: {len(unique)} unreplaced placeholder token(s):")
        for loc, tok in unique:
            print(f"  {loc}: {tok}")


def make_docx_bytes(doc):
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ─── INTEREST-TEXT PHRASE MAP ─────────────────────────────────────────────────
# Render-time cleanup applied to every "Description of Legacy / Devise / Other
# Interest" string before it goes into the petition or Notice of Probate.
# Order matters — earlier entries run first. To add a new mapping, append a
# tuple of (description, compiled regex, replacement). To disable one,
# comment it out.
INTEREST_PHRASE_MAP = [
    ("Drop conditional parens like '(if she survives decedent)'",
     re.compile(r'\s*\(if [^)]*\)', re.IGNORECASE), ''),

    ("Drop trust-name aliases like '(Family Trust)'",
     re.compile(r'\s*\((?:Family|Marital|Bypass|QTIP|Disclaimer|'
                r'Generation[- ]?Skipping|Charitable)\s+Trust\)',
                re.IGNORECASE), ''),

    ("Drop HEMS standard clause",
     re.compile(r',?\s*for\s+health,?\s+education,?\s+maintenance,?\s*'
                r'(?:or|and)\s+support', re.IGNORECASE), ''),

    ("'Income and principal beneficiary of' → 'Beneficiary of'",
     re.compile(r'\bIncome and principal beneficiary of\b', re.IGNORECASE),
     'Beneficiary of'),

    ("'Income beneficiary of' → 'Beneficiary of'",
     re.compile(r'\bIncome beneficiary of\b', re.IGNORECASE),
     'Beneficiary of'),

    ("'Legatee of (the) net residuary estate (outright)' → 'Residuary Beneficiary'",
     re.compile(r'\bLegatee of (?:the )?net residuary estate(?:\s+outright)?\b',
                re.IGNORECASE), 'Residuary Beneficiary'),

    ("Drop 'named in Will' from fiduciary roles "
     "('Executor named in Will' → 'Executor')",
     re.compile(r'\s+named in Will\b'), ''),

    ("'Article' → 'Art' (everywhere, not just after 'under')",
     re.compile(r'\bArticle\b'), 'Art'),

    # Run AFTER the Article→Art transform so we can match 'Art' here.
    # 'Trustee of trusts under Art FOUR under Art SEVEN(a)'
    #   → 'Trustee under Art SEVEN(a)'
    ("Drop 'of trusts under Art X' middle clause when another 'under Art' "
     "follows (Trustee chain)",
     re.compile(r'\s+of trusts under Art \S+(?=\s+under Art\b)', re.IGNORECASE),
     ''),

    ("Collapse extra whitespace",
     re.compile(r'\s{2,}'), ' '),
]


def _scrub_interest(text):
    """Apply INTEREST_PHRASE_MAP to clean up verbose interest text.
    Both petition and Notice of Probate render through this so they
    can't drift."""
    if not text:
        return text
    for _, rx, repl in INTEREST_PHRASE_MAP:
        text = rx.sub(repl, text)
    return text.strip(' ;,')


# ─── EPTL 4-1.1 SURVIVAL CHAIN ────────────────────────────────────────────────
# Single source of truth for "who's a distributee." The seven classes in
# EPTL 4-1.1(a) priority order; the FIRST class with a surviving member
# takes the entire intestate share, cutting off all subsequent classes.

EPTL_CLASSES = (
    ("spouse", 0),
    ("children", 1),
    ("parents", 2),
    ("siblings", 3),
    ("grandparents", 4),
    ("auntsUncles", 5),
    ("firstCousins", 6),
)

# Keyword map used as a fallback (and for counting per-class survivors)
# when the form's relationship strings are the only available signal.
DIST_REL_MAP = {
    "spouse": 0, "husband": 0, "wife": 0,
    "son": 1, "daughter": 1, "child": 1, "children": 1, "issue": 1,
    "grandchild": 1, "grandson": 1, "granddaughter": 1,
    "mother": 2, "father": 2, "parent": 2,
    "sister": 3, "brother": 3, "sibling": 3,
    "half-sister": 3, "half-brother": 3,
    "niece": 3, "nephew": 3,
    "grandmother": 4, "grandfather": 4, "grandparent": 4,
    "aunt": 5, "uncle": 5, "cousin": 5,
}


def distributee_classes(data):
    """Return the set of EPTL 4-1.1 class indices whose members are actual
    distributees (per the family-tree questionnaire).

    Special-case for EPTL 4-1.1(a)(1): when both spouse AND children
    survive, both classes share the intestate estate, so both are
    distributees. Otherwise the first surviving class takes alone.
    """
    ft = data.get("ft") or {}
    spouse_alive = ft.get("spouse") is True
    children_alive = ft.get("children") is True
    if spouse_alive and children_alive:
        return {0, 1}
    if spouse_alive:
        return {0}
    if children_alive:
        return {1}
    fs = first_surviving_class(data)
    return {fs} if fs is not None else set()


def first_surviving_class(data):
    """Return the EPTL 4-1.1 class index of the first surviving class,
    or None if it can't be determined.

    Priority of signals:
      1. ``data.ft`` family-tree questionnaire (legal source of truth) —
         each key is True/False/None; first key with True wins.
      2. Legacy ``data.surviving*`` flat fields, if any are populated.
      3. Auto-derive by counting ``data.distributees`` relationship
         keywords against DIST_REL_MAP — only used when ft and the
         legacy fields are both empty.
    """
    # 1. ft (the FTW result)
    ft = data.get("ft") or {}
    for key, idx in EPTL_CLASSES:
        if ft.get(key) is True:
            return idx

    # 2. Legacy surviving* fields
    legacy = (
        ("survivingSpouse", 0), ("survivingChildren", 1),
        ("survivingParents", 2), ("survivingSiblings", 3),
        ("survivingGrandparents", 4), ("survivingAuntsUncles", 5),
        ("survivingFirstCousinsOnceRemoved", 6),
    )
    for key, idx in legacy:
        raw = data.get(key)
        if raw and str(raw).strip().lower() not in ("false", "0", "no", ""):
            return idx

    # 3. Fallback: derive from distributees array
    for idx in range(7):
        for d in (data.get("distributees") or []):
            rel = (d.get("relationship") or "").strip().lower()
            for kw, ci in DIST_REL_MAP.items():
                if kw in rel and ci == idx:
                    return idx
    return None


# ─── COVER LETTER ─────────────────────────────────────────────────────────────

def generate_cover_letter(data):
    county = data.get("county", "")
    proceeding = data.get("proceedingType", "Probate")
    signer_key = data.get("signer", "Jessica Wilson")
    signer = SIGNERS.get(signer_key, signer_key)
    decedent = decedent_full(data)
    efile_date = data.get("efileDate", today())
    enclosures = data.get("enclosures", [])

    county_info = COUNTY_INFO.get(county, {})
    address = county_info.get("address", "")
    city_state_zip = county_info.get("city_state_zip", "")
    dept = county_info.get("dept_probate" if proceeding == "Probate" else "dept_admin", "")

    doc = Document()

    # Set default style to single-spaced
    style = doc.styles['Normal']
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.space_before = Pt(0)

    def _para(text="", space_after=0):
        p = doc.add_paragraph(text)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        p.paragraph_format.space_after = Pt(space_after)
        p.paragraph_format.space_before = Pt(0)
        return p

    # Date
    _para(today(), space_after=12)

    # Addressee
    _para(f"Surrogate's Court, {county} County")
    _para(f"Attn: {dept}")
    _para(address)
    _para(city_state_zip, space_after=12)

    # RE line
    _para(f"RE: Estate of {decedent}", space_after=12)

    _para("Greetings,", space_after=6)

    proc_word = {
        "NonDomiciliary": "non-domiciliary administration",
        "AdminCTA": "administration c.t.a.",
        "SmallEstate": "voluntary administration (small estate)",
    }.get(proceeding, proceeding.lower())
    _para(
        f"Our office efiled the above referenced petition for {proc_word} on {efile_date}. "
        f"Please find enclosed the following original documents required by the Court:",
        space_after=6
    )

    # Enclosures as bullet list
    for enc in enclosures:
        p = doc.add_paragraph(style="List Bullet")
        p.text = enc
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        p.paragraph_format.space_after = Pt(0)

    _para("", space_after=6)
    _para("Please do not hesitate to call our office if you have concerns and questions.")
    _para("", space_after=6)
    _para("Sincerely,")
    _para("")
    _para("")
    _para(signer)
    _para("Enc.")

    _validate_docx(doc, "generate_cover_letter")
    return make_docx_bytes(doc)


# ─── 805 AFFIDAVIT ────────────────────────────────────────────────────────────

def generate_805(data):
    """Build the 805 Affidavit of Assets & Liabilities from scratch for
    consistent formatting (Times New Roman 12pt, 1-inch margins)."""
    doc = Document()

    # ── Page margins: 1" all sides ────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1)
        section.right_margin  = Inches(1)

    county    = data.get("county", "")
    decedent  = decedent_full(data)
    petitioner = petitioner_full(data)
    file_no   = data.get("fileNo", "")
    aka       = (data.get("decedentAKA") or "").strip()
    year      = datetime.now().strftime("%Y")

    # ── Helpers ───────────────────────────────────────────────────────────────
    FONT = "Times New Roman"
    SIZE = Pt(12)

    def _run(para, text, bold=False, italic=False):
        r = para.add_run(text)
        r.font.name  = FONT
        r.font.size  = SIZE
        r.bold       = bold
        r.italic     = italic
        return r

    def line(text="", bold=False, italic=False, center=False,
             space_before=0, space_after=0, left_indent=None):
        p = doc.add_paragraph()
        p.paragraph_format.alignment    = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after  = Pt(space_after)
        if left_indent is not None:
            p.paragraph_format.left_indent = Inches(left_indent)
        if text:
            _run(p, text, bold=bold, italic=italic)
        return p

    def blank(n=1):
        for _ in range(n):
            line()

    # ── Caption ──────────────────────────────────────────────────────────────
    proceeding = data.get("proceedingType", "Administration")
    letters_type = data.get("lettersType", "Letters of Administration")

    line("SURROGATE\u2019S COURT OF THE STATE OF NEW YORK", bold=True)
    line(f"COUNTY OF {county.upper()}", bold=True, space_after=2)

    divider = "\u2500" * 43 + "x"
    line(divider, space_after=0)

    # Two-column caption using a borderless table
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    def _no_border(cell):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = OxmlElement("w:tcBorders")
        for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
            el = OxmlElement(f"w:{side}")
            el.set(qn("w:val"), "none")
            el.set(qn("w:sz"), "0")
            el.set(qn("w:space"), "0")
            el.set(qn("w:color"), "auto")
            tcBorders.append(el)
        tcPr.append(tcBorders)

    # Build caption rows — left column has matter text, right has doc title
    aka_line = f"    a/k/a {aka}," if aka else ""
    if proceeding == "Probate":
        left_lines = [
            "PROBATE PROCEEDING, WILL OF",
            "",
            f"    {decedent.upper()},",
        ]
    else:
        left_lines = [
            "In the Matter of the Application for",
            "",
            f"{letters_type} of the Estate of",
            "",
            f"    {decedent.upper()},",
        ]
    if aka_line:
        left_lines.append(aka_line)
    left_lines += [
        "",
        "                Deceased.",
    ]

    right_lines = [
        ("AFFIDAVIT OF ASSETS", True),
        ("& LIABILITIES", True),
        ("(SCPA 805)", True),
        ("", False),
        (f"File No. {file_no}" if file_no else "", False),
    ]
    # Pad right column to match left
    while len(right_lines) < len(left_lines):
        right_lines.append(("", False))

    cap_tbl = doc.add_table(rows=len(left_lines), cols=2)
    cap_tbl.style = "Table Grid"

    # Set table width and column widths
    tbl_el = cap_tbl._tbl
    tblPr = tbl_el.tblPr if tbl_el.tblPr is not None else OxmlElement("w:tblPr")
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"), "0")
    tblW.set(qn("w:type"), "auto")
    tblPr.append(tblW)
    # Remove table borders at the table level too
    tblBorders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "none")
        el.set(qn("w:sz"), "0")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "auto")
        tblBorders.append(el)
    tblPr.append(tblBorders)

    for row_i, row in enumerate(cap_tbl.rows):
        for col_i, cell in enumerate(row.cells):
            _no_border(cell)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            if col_i == 0:
                _run(p, left_lines[row_i])
            else:
                # Add vertical bar separator before right-column text
                txt, bld = right_lines[row_i]
                cell_text = f"\u2502  {txt}" if txt else "\u2502"
                _run(p, cell_text, bold=bld)

    line(divider, space_after=4)

    # ── Venue block ───────────────────────────────────────────────────────────
    line("STATE OF NEW YORK\t\t\t\t)")
    line("\t\t\t\t\t\t) ss:")
    line(f"COUNTY OF {county.upper()}\t\t\t\t)")
    blank()

    # ── Oath paragraph ────────────────────────────────────────────────────────
    line(
        "I, the undersigned being duly sworn, depose and say:  I have personal knowledge "
        "as to the assets, debts and/or liabilities of the estate of the decedent. "
        "The assets of the estate, including real and/or personal property held solely "
        "by the decedent consist of:",
        space_after=4
    )

    # Assets — use individual asset tracker entries if available, else summary fields
    tracked_assets = [a for a in data.get("assets", []) if a.get("institution")]
    asset_lines = []

    if tracked_assets:
        for a in tracked_assets:
            val = nonzero(a.get("value"))
            inst = a.get("institution", "")
            cat = a.get("category", "")
            acct = a.get("accountNumber", "")
            desc = f"{cat} – {inst}" if cat and inst else (inst or cat)
            if acct:
                desc += f" (acct ...{acct[-4:]})" if len(acct) >= 4 else f" (acct {acct})"
            if val:
                asset_lines.append(f"{desc}:  ${val}")
            else:
                asset_lines.append(desc)

        # Also include real property from summary fields (not tracked in asset cards)
        ir = nonzero(data.get("improvedRealProperty"))
        ur = nonzero(data.get("unimprovedRealProperty"))
        rd = (data.get("realPropertyDescription") or "").strip()
        gr = nonzero(data.get("grossRents18mo"))
        if ir: asset_lines.append(f"Improved Real Property (NY):  ${ir}")
        if ur: asset_lines.append(f"Unimproved Real Property (NY):  ${ur}")
        if rd: asset_lines.append(f"Description:  {rd}")
        if gr: asset_lines.append(f"Gross Rents (18 months):  ${gr}")
    else:
        pp = nonzero(data.get("personalPropertyValue"))
        ir = nonzero(data.get("improvedRealProperty"))
        ur = nonzero(data.get("unimprovedRealProperty"))
        rd = (data.get("realPropertyDescription") or "").strip()
        gr = nonzero(data.get("grossRents18mo"))
        if pp: asset_lines.append(f"Personal Property:  ${pp}")
        if ir: asset_lines.append(f"Improved Real Property (NY):  ${ir}")
        if ur: asset_lines.append(f"Unimproved Real Property (NY):  ${ur}")
        if rd: asset_lines.append(f"Description:  {rd}")
        if gr: asset_lines.append(f"Gross Rents (18 months):  ${gr}")

    if not asset_lines:
        asset_lines = ["NONE"]

    for asset in asset_lines:
        line(asset, left_indent=0.5)

    blank()

    # ── Liabilities ───────────────────────────────────────────────────────────
    line(
        "All the liabilities of the decedent known to me are as follows "
        "(Indicate AMOUNT DUE or answer \u201cNONE\u201d):",
        space_before=4, space_after=4
    )

    mort = (data.get("mortgageAmount") or "").strip()
    fp   = (data.get("funeralPaid") or "").strip()
    fo   = (data.get("funeralOutstanding") or "").strip()
    misc = (data.get("miscDebts") or "").strip()

    line(f"Amount of outstanding mortgages:  {mort or 'NONE'}", left_indent=0.5)
    line(
        f"Amount of funeral expenses paid (attach copy of paid funeral bill):  {fp or 'NONE'}",
        left_indent=0.5
    )
    line(f"Amount of funeral expenses still outstanding:  {fo or 'NONE'}", left_indent=0.5)
    blank()
    line(
        "Itemize and specify amount of any miscellaneous expenses payable "
        "(i.e. credit card, utility bills, insurance premiums, etc.  "
        "Use attachments if more space is required.)",
        italic=True, space_after=2
    )
    line("NOTE: ANY UNSECURED DEBT MAY BE BONDED", bold=True, space_after=4)

    if misc:
        for ln in misc.splitlines():
            if ln.strip():
                line(ln.strip(), left_indent=0.5)
    else:
        line("NONE", left_indent=0.5)

    blank()

    # ── WHEREFORE clause ──────────────────────────────────────────────────────
    bond_status, bond_reason = compute_bond_status(data)
    pet_role = "Executor" if proceeding == "Probate" else "Administrator"
    role_desc = pet_role

    if bond_status == "dispense":
        wherefore_text = (
            f"WHEREFORE, your deponent prays, that the filing of a bond by {petitioner} "
            f"as {role_desc} be dispensed with."
        )
    else:
        wherefore_text = (
            f"WHEREFORE, your deponent prays for the issuance of Letters to {petitioner} "
            f"as {role_desc}, upon the filing of a bond."
        )
    line(wherefore_text, space_before=6, space_after=18)

    # ── Signature block ───────────────────────────────────────────────────────
    line("__________________________________", space_after=2)
    line(petitioner, space_after=14)

    line(f"Sworn to before me this _________")
    line(f"day of __________________, {year}")
    blank()
    line("__________________________________", space_after=2)
    line("Notary Public")

    _validate_docx(doc, "generate_805")
    return make_docx_bytes(doc)


# ─── AFFIDAVIT OF HEIRSHIP ────────────────────────────────────────────────────

def generate_heirship(data):
    proceeding = data.get("proceedingType", "Administration")
    if proceeding == "Probate":
        template = "Affidavit_of_Heirship_Full_Probate.docx"
        letters_phrase = "Letters Testamentary"
    else:
        template = "Affidavit_of_Heirship_Full_Admin.docx"
        letters_phrase = "Letters of Administration"
    doc = Document(os.path.join(WORD_TEMPLATES_DIR, template))
    decedent = decedent_full(data)
    county = data.get("county", "")
    petitioner = petitioner_full(data)
    # `.get(key, default)` only uses default when the key is MISSING; it does
    # NOT trigger when value is "". Saved cases populate every field as ""
    # so we need `or` fallbacks instead.
    petitioner_full_addr = ", ".join(filter(None, [
        data.get("petitionerStreet", ""), data.get("petitionerCity", ""),
        data.get("petitionerState", ""), data.get("petitionerZip", ""),
    ]))
    deponent = (data.get("deponentName") or "").strip() or petitioner
    deponent_address = (data.get("deponentAddress") or "").strip() or petitioner_full_addr or "_________________________"
    deponent_rel = (data.get("deponentRelationship") or "").strip() or (data.get("petitionerRelationship") or "").strip() or "______________"
    years_known = (data.get("yearsKnown") or "").strip() or "_____"
    dob = data.get("decedentDOB", "")
    dod = data.get("decedentDOD", "")
    marital_status    = (data.get("maritalStatus") or "").strip()      # never_married / married / divorced / widowed
    spouse_name       = (data.get("spouseName") or "").strip()
    divorce_year      = (data.get("divorceYear") or "").strip()
    prior_spouse_death = (data.get("priorSpouseDeathDate") or "").strip()
    children_note = data.get("childrenNote", "").strip()
    mother_name = data.get("motherName", "")
    mother_dod = data.get("motherDOD", "")
    father_name = data.get("fatherName", "")
    father_dod = data.get("fatherDOD", "")
    sole_distributee = (data.get("soleDistributee") or "").strip() or petitioner

    was_married = marital_status in ("married", "divorced", "widowed")
    has_children = bool(children_note and "never had" not in children_note.lower())

    # Build the marriage sentence for para 21
    if marital_status == "married":
        marriage_sentence = (
            f"Decedent was married to {spouse_name} at the time of death "
            f"and was never divorced."
        )
    elif marital_status == "divorced":
        yr = f"in {divorce_year}" if divorce_year else "prior to death"
        marriage_sentence = (
            f"Decedent was married to {spouse_name}, which said marriage ended in "
            f"divorce {yr}. The decedent never remarried after said divorce."
        )
    elif marital_status == "widowed":
        when = f"on {prior_spouse_death}" if prior_spouse_death else "prior to the decedent's death"
        marriage_sentence = (
            f"Decedent was married to {spouse_name}, who predeceased the decedent "
            f"{when}. The decedent never remarried after the death of said spouse."
        )
    else:
        marriage_sentence = None  # never married — use para 23 instead

    paras_to_delete = []

    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()

        if "which said marriage ended in divorce" in text:
            if marriage_sentence:
                replace_para(para, para.text, marriage_sentence)
            else:
                paras_to_delete.append(i)

        elif "The decedent was/never married" in text:
            if was_married:
                paras_to_delete.append(i)
            else:
                replace_para(para, para.text,
                             "The decedent was never married.")

        elif "never had any children" in text or "did have children" in text:
            if has_children:
                replace_para(para, "The decedent never had any children, adopted, out of wedlock nor marital. Or did have children.", children_note)
            else:
                replace_para(para, "The decedent never had any children, adopted, out of wedlock nor marital. Or did have children.",
                    "The decedent never had any children, adopted, out of wedlock nor marital.")

        elif "The marriage of" in text and "bore no children" in text:
            if was_married:
                replace_para(para, "___________ and ____________",
                             f"{decedent} and {spouse_name}")
            else:
                paras_to_delete.append(i)

        elif "There were no children of the decedent" in text:
            if was_married or has_children:
                paras_to_delete.append(i)

    for i in sorted(paras_to_delete, reverse=True):
        p = doc.paragraphs[i]._element
        p.getparent().remove(p)

    # Format dates as written-out (e.g. "March 15, 1945")
    dob_long = format_date_long(dob)
    dod_long = format_date_long(dod)

    replace_in_doc(doc, {
        "COUNTY OF _____________": f"COUNTY OF {county.upper()}",
        "___________________\t\t\t\t\tAFFIDAVIT OF HEIRSHIP": f"{decedent}\t\t\t\t\tAFFIDAVIT OF HEIRSHIP",
        "A/K/A ___________________\t\t\t\tFile No.:": f"A/K/A {data.get('decedentAKA', '')}\t\t\t\tFile No.: {data.get('fileNo', '')}",
        "COUNTY OF \t\t\t)": "COUNTY OF \t\t\t)",  # Leave blank — affiant may be out of state
        "STATE OF NEW YORK\t)": "STATE OF \t\t)",  # Leave blank — affiant may be out of state
        "______day of __________, 2022": "______day of __________, 2026",
        "\tI, ______________, being duly sworn, deposes and says:": f"\tI, {deponent}, being duly sworn, deposes and says:",
        "I reside at _________________________.  I am over the age of eighteen (18) years and I am fully familiar with the facts and circumstances herein, the decedent\u2019s family tree, as I am the ______________of the Decedent and have known the Decedent for over _____ years.":
            f"I reside at {deponent_address}.  I am over the age of eighteen (18) years and I am fully familiar with the facts and circumstances herein, the decedent\u2019s family tree, as I am the {deponent_rel} of the Decedent and have known the Decedent for over {years_known} years.",
        "The Decedent was born on ___________ and died on __________________.": f"The Decedent was born on {dob_long} and died on {dod_long}.",
        "Mother: ": f"Mother: {mother_name}",
        "Father: ": f"Father: {father_name}",
        f"Therefore, ______________ is the sole distributee of the Estate of ______________":
            f"Therefore, {sole_distributee} is the sole distributee of the Estate of {decedent}",
        f"This affidavit is made with my personal knowledge knowing the ______________ County Surrogate\u2019s Court will rely thereon in issuing Letters Testamentary to _________________, the petitioner." if proceeding == "Probate" else
        f"This affidavit is made with my personal knowledge knowing the ______________ County Surrogate\u2019s Court will rely thereon in issuing Letters of Administration to _________________, the petitioner.":
            f"This affidavit is made with my personal knowledge knowing the {county} County Surrogate\u2019s Court will rely thereon in issuing {letters_phrase} to {petitioner}, the petitioner.",
    })

    mother_dod_filled = False
    for para in doc.paragraphs:
        t = para.text.strip()
        if t.startswith("Date of Death:"):
            if not mother_dod_filled:
                replace_para(para, "Date of Death:", f"Date of Death: {mother_dod}")
                mother_dod_filled = True
            else:
                replace_para(para, "Date of Death:", f"Date of Death: {father_dod}")

    _validate_docx(doc, "generate_heirship")
    return make_docx_bytes(doc)


# ─── WAIVER COVER LETTER ──────────────────────────────────────────────────────

def generate_waiver_cover(data, distributee):
    doc = Document(os.path.join(WORD_TEMPLATES_DIR, "Waiver_cover_letter.docx"))
    decedent = decedent_full(data)
    petitioner = petitioner_full(data)
    dist_name = distributee.get("name", "")
    dist_rel = distributee.get("relationship", "")
    dist_addr = distributee.get("address", "")

    replace_in_doc(doc, {
        "September 27, 2022": today(),
        "(Distributee)": dist_name,
        "(Distributee Address)": dist_addr,
        "(Deceased)": decedent,
        "(Petitioner)": petitioner,
    })

    _validate_docx(doc, "generate_waiver_cover")
    return make_docx_bytes(doc)


# ─── AFFIDAVIT AS TO DESIGNEE ─────────────────────────────────────────────────

def generate_designee_affidavit(data, distributee):
    """Affidavit as to Designee — signed by a distributee who is stepping aside
    and designating the petitioner (the "designee") to serve as fiduciary.

    Generated once per distributee when the petitioner's interest is "Designee".
    Caption, file no., decedent and proposed administrator (the petitioner) are
    filled from the case; the affiant is the distributee. The *reason* the affiant
    does not wish to serve is left as a blank line to complete by hand.
    """
    county     = (data.get("county", "") or "").strip()
    file_no    = (data.get("fileNo", "") or "").strip()
    aka        = (data.get("decedentAKA", "") or "").strip()
    decedent   = decedent_full(data)
    petitioner = petitioner_full(data)
    aff_name   = (distributee.get("name", "") or "").strip()
    aff_rel    = (distributee.get("relationship", "") or "").strip()
    aff_addr   = (distributee.get("address", "") or "").strip()

    proceeding = data.get("proceedingType") or data.get("proceeding") or "Administration"
    title = "Administrator c.t.a." if proceeding == "AdminCTA" else "Administrator"

    doc = Document()
    for s in doc.sections:
        s.left_margin = s.right_margin = Inches(1)
        s.top_margin = s.bottom_margin = Inches(1)

    def p(text="", *, bold=False, align=None, space_after=6):
        para = doc.add_paragraph()
        para.paragraph_format.space_after = Pt(space_after)
        para.paragraph_format.line_spacing = 1.0
        if align is not None:
            para.alignment = align
        run = para.add_run(text)
        run.bold = bold
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)
        return para

    RIGHT = WD_ALIGN_PARAGRAPH.RIGHT
    rule = "—" * 52 + "X"

    p("SURROGATE'S COURT OF THE STATE OF NEW YORK", bold=True, space_after=0)
    p(f"COUNTY OF {county.upper()}", bold=True, space_after=0)
    p(rule, space_after=0)
    p("ADMINISTRATION PROCEEDING,", space_after=0)
    p(f"ESTATE OF {decedent},", space_after=0)
    if aka:
        p(f"        a/k/a {aka}", space_after=0)
    p("                                Deceased.", space_after=0)
    p(rule, space_after=0)
    p("AFFIDAVIT AS TO DESIGNEE", bold=True, align=RIGHT, space_after=0)
    p(f"File No.: {file_no}", align=RIGHT)

    p("STATE OF NEW YORK\t)", space_after=0)
    p("\t\t\t)  ss.:", space_after=0)
    p(f"COUNTY OF {county.upper()}\t)")

    p(f"I, {aff_name}, being duly sworn, deposes and says:")
    p(f"1.\tI reside at {aff_addr or '________________________________________'}.")
    p(f"2.\tI am the Decedent's {aff_rel or '________________'}.")
    p(f"3.\tI have asked {petitioner} to be designated as {title} of the Estate of "
      f"{decedent}.  ________________________________________________ (reason).")
    p(f"4.\t{petitioner} will carry out this duty efficiently, fairly and in the best "
      f"interest of the family.  I do not want to act as {title} because "
      f"________________________ and it is too difficult.")
    p(f"5.\tThis affidavit is made with my personal knowledge knowing the {county} "
      f"County Surrogate's Court will rely thereon in issuing Letters of "
      f"Administration to {petitioner}, the petitioner herein.")

    p("_______________________________", align=RIGHT, space_after=0)
    p(aff_name, align=RIGHT)
    p("Sworn to before me on", space_after=0)
    p("___________________, 20___")
    p("_______________________", space_after=0)
    p("Notary Public", space_after=0)

    _validate_docx(doc, "generate_designee_affidavit")
    return make_docx_bytes(doc)


# ─── ATTORNEY CERTIFICATION ───────────────────────────────────────────────────

def generate_attorney_cert(data):
    doc = Document(os.path.join(WORD_TEMPLATES_DIR, "newcertform_6_59_19_PM.docx"))
    replace_in_doc(doc, {
        "Dated:": f"Dated: {today()}",
    })
    _validate_docx(doc, "generate_attorney_cert")
    return make_docx_bytes(doc)


# ─── PDF FILLING (pymupdf/fitz) ──────────────────────────────────────────────

def _ensure_acroform_root(pdf_bytes):
    """Reattach the Root /AcroForm dictionary that fitz drops on save.

    Without /AcroForm in Root:
    - pypdf and other strict readers see 0 form fields (silent failure)
    - Court e-file systems that validate AcroForm structure may reject
    - Some viewers won't show fields as interactive

    /NeedAppearances True — tells viewers (esp. Acrobat) to regenerate
    field appearances from /V + /DA at display time. fitz bakes
    appearance streams that strict-clip at the widget rect; flipping
    this flag lets Acrobat re-render with its own multiline text engine,
    which wraps overflow into the form's pre-printed continuation lines
    the same way it does when the form is filled manually.

    No global /DR on the AcroForm dict: each widget in this template
    carries its own /DR with its font reference (e.g. /Helv), so a global
    one would be ignored anyway and might shadow per-widget fonts.

    Only top-level widgets (no /Parent) go in /Fields — defensive against
    field hierarchies even though our templates are flat.
    """
    try:
        import pypdf
        from pypdf.generic import (
            DictionaryObject, ArrayObject, NameObject, BooleanObject,
        )
    except Exception as e:
        print(f"[ACROFORM] pypdf unavailable, skipping reattach: {e}")
        return pdf_bytes

    try:
        reader = pypdf.PdfReader(io.BytesIO(pdf_bytes))
        writer = pypdf.PdfWriter(clone_from=reader)
        field_refs = []
        for page in writer.pages:
            annots = page.get("/Annots") or []
            for a in annots:
                obj = a.get_object() if hasattr(a, "get_object") else a
                if obj.get("/Subtype") == "/Widget" and "/Parent" not in obj:
                    field_refs.append(a)
        af = DictionaryObject()
        af[NameObject("/Fields")] = ArrayObject(field_refs)
        af[NameObject("/NeedAppearances")] = BooleanObject(True)
        # pypdf preserves /AcroForm only if it's an indirect object reference
        # — direct dict assignment to _root_object is dropped during write().
        af_ref = writer._add_object(af)
        writer._root_object[NameObject("/AcroForm")] = af_ref
        out = io.BytesIO()
        writer.write(out)
        return out.getvalue()
    except Exception as e:
        print(f"[ACROFORM] reattach failed, returning original bytes: {e}")
        traceback.print_exc()
        return pdf_bytes


def fill_pdf(template_path, fields, font_overrides=None):
    """Universal PDF form filler using pymupdf/fitz.

    Behavior:
    - Every text field defaults to 8pt (unless template explicitly sets a non-zero size).
    - Checkbox-style 'X' values are sized to the box height so they fill the box.
    - font_overrides: optional dict mapping field name → font size in points.
      When the field's name is a key, that size wins over the default 8pt.

    Note: Multiline flag is NOT set. Adding Multiline to single-line template
    fields causes Acrobat to draw a "+" scroll indicator at every field's
    right edge (visible both on screen and on print). Acrobat with
    /NeedAppearances=True handles long-content wrapping naturally without
    needing the Multiline flag, so we leave each field's flags as the
    template author set them.

    Handles text, checkboxes (True/False), radio buttons, and combo/dropdown fields.
    Calls widget.update() to bake appearance streams so fields render in all viewers.
    """
    doc = fitz.open(template_path)
    for page in doc:
        for widget in page.widgets():
            name = widget.field_name
            if name not in fields:
                # XFA-flattened templates (e.g. probcta.pdf) carry full dotted
                # paths like "topmostSubform[0].Page1[0].TextField13[0]" while
                # mappings use the bare leaf name. Fall back to the last dotted
                # segment — only for dotted names, so plain-named templates are
                # unaffected.
                if "." in name and name.rsplit(".", 1)[1] in fields:
                    name = name.rsplit(".", 1)[1]
                else:
                    continue
            value = fields[name]
            if widget.field_type == fitz.PDF_WIDGET_TYPE_CHECKBOX:
                widget.field_value = bool(value)
            elif widget.field_type == fitz.PDF_WIDGET_TYPE_RADIOBUTTON:
                widget.field_value = str(value).lstrip("/")
            else:
                s = str(value) if value is not None else ""
                widget.field_value = s
                if s == "X":
                    if font_overrides and name in font_overrides:
                        # Explicit per-field override wins.
                        widget.text_fontsize = font_overrides[name]
                    else:
                        # Default for ALL checkbox-style "X" cells:
                        # fontsize=0 triggers PDF spec auto-sizing —
                        # Acrobat (which regenerates appearances under
                        # /NeedAppearances=True) picks a size that
                        # makes the X fit the cell. Works for both
                        # tiny 6x6 cells and larger 10x10 ones.
                        widget.text_fontsize = 0
                else:
                    # Per-field override wins. Otherwise universal 8pt
                    # default — fits the cell heights on these templates
                    # and matches the §1–§7 visual style. Template
                    # fontsize (often 12pt) is always overridden because
                    # 12pt overflows narrow form cells.
                    if font_overrides and name in font_overrides:
                        widget.text_fontsize = font_overrides[name]
                    else:
                        widget.text_fontsize = 8
            widget.update()
    buf = io.BytesIO()
    doc.save(buf)
    doc.close()
    buf.seek(0)
    # Reattach Root /AcroForm — fitz drops it on save, which makes pypdf
    # and other strict readers see zero form fields.
    return _ensure_acroform_root(buf.read())


def _extract_pages(pdf_bytes, page_indices):
    """Extract specific pages from PDF bytes, preserving form widgets."""
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    doc.select(page_indices)
    buf = io.BytesIO()
    doc.save(buf, garbage=4, deflate=True)
    doc.close()
    buf.seek(0)
    # fitz.save drops Root /AcroForm — reattach for pypdf / strict reader compatibility
    return _ensure_acroform_root(buf.read())


def extract_pdf_pages(template_path, fields, page_indices):
    """Fill a PDF template and extract specific pages."""
    filled = fill_pdf(template_path, fields)
    return _extract_pages(filled, page_indices)


# ─── PROBATE PDF (P-1 + OATH + WITNESS) ──────────────────────────────────────


def _auto_compute_property(data):
    """Auto-compute property values from asset tracker when available."""
    tracked = [a for a in data.get("assets", []) if a.get("institution")]
    if not tracked:
        return
    personal = 0
    real = 0
    for a in tracked:
        try:
            val = float(str(a.get("value", "0")).replace(",", "").replace("$", ""))
        except (ValueError, TypeError):
            val = 0
        if a.get("category") == "Real Estate":
            real += val
        else:
            personal += val
    if personal > 0:
        data["personalPropertyValue"] = f"{personal:,.2f}"
    if real > 0:
        data["realPropertyValue"] = f"{real:,.2f}"


def compute_bond_status(data):
    """Determine bond status from distributee dispositions.
    Returns: ('dispense', reason) or ('require', reason)
    """
    dists = [d for d in data.get("distributees", []) if d.get("name")]
    has_minors = any(d.get("isMinor") for d in dists)
    has_citations = any(d.get("disposition") == "citation" for d in dists)
    all_waive = all(d.get("disposition") == "waiver" for d in dists) and len(dists) > 0

    if has_minors:
        return ("require", "minor/person under disability exists")
    if has_citations:
        return ("require", "citation required for one or more parties")
    if all_waive:
        return ("dispense", "all distributees waiving")
    if data.get("dispenseBond"):
        return ("dispense", "per will / petitioner request")
    return ("dispense", "")


def compute_interested_persons(data, pet, pet_addr, letters_to):
    """Build the canonical list of interested persons used by both the
    Probate Petition and the Notice of Probate.

    Merges ``data.distributees`` with ``data.willBeneficiaries``,
    auto-inserts the petitioner / successor executor / trustee / guardian,
    applies EPTL 4-1.1 first-surviving-class logic to decide who actually
    qualifies for the "Distributee" label, and prepends Executor when the
    name matches the recipient of Letters Testamentary.

    The ``interest`` string on each entry is the canonical nature-of-
    interest text — both the petition and the notice render this same
    string, so the two documents cannot drift.

    Args:
        data: form data dict.
        pet: petitioner full name (e.g. ``"Amy Sue Nathan"``).
        pet_addr: pre-formatted petitioner address line.
        letters_to: name of the executor receiving Letters Testamentary
                    (typically equals ``pet``).

    Returns:
        list of dicts, each with ``name``, ``relationship``, ``address``,
        ``citizenship``, ``interest``, ``beneficiaryType``, ``isMinor``,
        ``dob``, ``guardianInfo``.
    """
    # ── Fuzzy dedup of data.distributees ────────────────────────────────────
    # The same person is sometimes saved multiple times — once when entered
    # by hand (e.g. "Amy Sue Nathan") and again from Smart Intake (e.g.
    # "Amy Nathan", with a richer interest field). Collapse by
    # (first_token, last_meaningful_token) and keep the richest entry —
    # otherwise the person who matters ends up with a worse description on
    # both the petition and the notice.
    _SUFFIXES = {"jr", "sr", "ii", "iii", "iv", "v"}
    def _name_key_for_dedup(name):
        if not name:
            return None
        toks = [t for t in re.split(r'\s+', name.strip()) if t]
        while toks and toks[-1].lower().rstrip('.') in _SUFFIXES:
            toks.pop()
        if not toks:
            return None
        first = toks[0].lower()
        last = toks[-1].lower() if len(toks) > 1 else ""
        return (first, last)

    def _entry_score(d):
        # Higher = more useful entry. Prefer entries with a populated
        # interest, then address, then a non-"Unknown" relationship.
        s = 0
        if (d.get("interest") or "").strip():
            s += 100
        if (d.get("address") or "").strip():
            s += 10
        rel = (d.get("relationship") or "").strip().lower()
        if rel and rel != "unknown":
            s += 5
        if (d.get("dob") or "").strip():
            s += 1
        return s

    _dedup_seen = {}      # key -> entry index in all_dists
    all_dists = []
    for d in (data.get("distributees") or []):
        if not d.get("name"):
            continue
        key = _name_key_for_dedup(d.get("name", ""))
        if key is None:
            all_dists.append(d)
            continue
        if key not in _dedup_seen:
            _dedup_seen[key] = len(all_dists)
            all_dists.append(d)
        else:
            existing = all_dists[_dedup_seen[key]]
            if _entry_score(d) > _entry_score(existing):
                all_dists[_dedup_seen[key]] = d

    # Auto-insert petitioner in §6a if not already listed
    pet_lower = pet.strip().lower()
    if pet_lower and not any(d.get("name", "").strip().lower() == pet_lower for d in all_dists):
        pet_interest_label = data.get("petitionerInterest", "Executor(s) named in decedent's Will")
        all_dists.insert(0, {
            "name": pet,
            "address": pet_addr,
            "citizenship": data.get("petitionerCitizenship", "U.S.A."),
            "interest": pet_interest_label,
            "relationship": data.get("petitionerRelationship", ""),
            "beneficiaryType": "primary",
            "isMinor": False,
        })

    # Auto-insert successor executor into §7 if provided and not already listed
    succ_exec = (data.get("successorExecutor") or "").strip()
    if succ_exec and not any(d.get("name", "").strip().lower() == succ_exec.lower() for d in all_dists):
        all_dists.append({
            "name": succ_exec,
            "relationship": "",
            "address": "",
            "citizenship": "U.S.A.",
            "interest": "Successor Executor named in Will",
            "beneficiaryType": "successor",
            "isMinor": False,
        })

    # Auto-insert trustee into §7 if provided and not already listed
    trustee = (data.get("trusteeName") or "").strip()
    trust_name = (data.get("trustName") or "").strip()
    if trustee and not any(d.get("name", "").strip().lower() == trustee.lower() for d in all_dists):
        interest = f"Trustee of {trust_name}" if trust_name else "Trustee named in Will"
        all_dists.append({
            "name": trustee,
            "relationship": "Trustee",
            "address": "",
            "citizenship": "U.S.A.",
            "interest": interest,
            "beneficiaryType": "successor",
            "isMinor": False,
        })

    # Auto-insert guardian into §7 if provided and not already listed
    guardian = (data.get("guardianName") or "").strip()
    if guardian and not any(d.get("name", "").strip().lower() == guardian.lower() for d in all_dists):
        all_dists.append({
            "name": guardian,
            "relationship": "Guardian",
            "address": "",
            "citizenship": "U.S.A.",
            "interest": "Guardian of minor(s) named in Will",
            "beneficiaryType": "successor",
            "isMinor": False,
        })

    # ── Merge will beneficiaries (extracted by Smart Intake) ─────────────────
    # The structured willBeneficiaries list is the source of truth for legacy
    # descriptions. Match by name (case-insensitive); if the person already
    # appears in all_dists (e.g., a distributee who is also a will legatee),
    # replace their generic interest with the standardized willBenef text.
    # If they don't appear, append as a new entry routed to ¶7.
    def _norm_name(n):
        return re.sub(r'\s+', ' ', (n or '').strip().lower())

    for wb in (data.get("willBeneficiaries") or []):
        wb_name = (wb.get("name") or "").strip()
        if not wb_name:
            continue
        # Skip deceased beneficiaries entirely — they're not necessary parties
        # (predeceased: bequest lapses; post-deceased: their estate would be
        # noticed separately, captured as a TODO until we add fiduciary-of-
        # estate handling).
        if wb.get("deceased"):
            continue
        wb_interest = _scrub_interest((wb.get("interest") or "").strip())
        wb_rel = (wb.get("relationship") or "").strip()
        wb_addr = (wb.get("address") or "").strip()
        wb_minor = bool(wb.get("isMinor"))
        wb_is_dist = bool(wb.get("isDistributee"))
        wb_dob = (wb.get("dob") or "").strip()
        wb_guard = (wb.get("guardianInfo") or "").strip()
        wb_norm = _norm_name(wb_name)
        match = next((d for d in all_dists if _norm_name(d.get("name", "")) == wb_norm), None)
        if match is not None:
            if wb_interest:
                match["interest"] = wb_interest
            if wb_minor:
                match["isMinor"] = True
            if wb_is_dist:
                # Promote: AI says this person IS a distributee. Carry it
                # through so the Distributee label below picks them up
                # even when their relationship string is something the
                # keyword map doesn't match (in-laws, "Sister of Spouse",
                # etc.).
                match["isDistributee"] = True
            if wb_rel and not match.get("relationship"):
                match["relationship"] = wb_rel
            if wb_addr and not match.get("address"):
                match["address"] = wb_addr
            if wb_dob:
                match["dob"] = wb_dob
            if wb_guard:
                match["guardianInfo"] = wb_guard
        else:
            ben_type = "primary" if wb_is_dist else "successor"
            all_dists.append({
                "name": wb_name,
                "relationship": wb_rel,
                "address": wb_addr,
                "citizenship": "U.S.A.",
                "interest": wb_interest,
                "beneficiaryType": ben_type,
                "isMinor": wb_minor,
                "isDistributee": wb_is_dist,
                "dob": wb_dob,
                "guardianInfo": wb_guard,
            })

    # ── Re-route entries by actual EPTL distributee status ──────────────────
    # §6 of the petition lists distributees (and the petitioner). When a
    # surviving spouse and/or child takes under EPTL 4-1.1(a)(1)/(2)/(3),
    # parents/siblings/cousins are NOT distributees regardless of whether
    # the form data lists them. Anyone in data.distributees who's neither
    # an actual EPTL distributee nor a will beneficiary is dropped from
    # the petition + Notice of Probate entirely (matches the "if not a
    # beneficiary, not on notice" rule).
    exec_lower = letters_to.strip().lower()
    pet_lower = pet.strip().lower()
    _first_surv = first_surviving_class(data)
    _dist_classes = distributee_classes(data)

    def _is_actual_distributee(dist):
        if dist.get("isDistributee") is True:
            return True
        rel = (dist.get("relationship") or "").strip().lower()
        rel_class = next((ci for kw, ci in DIST_REL_MAP.items() if kw in rel), None)
        return rel_class is not None and rel_class in _dist_classes

    pruned = []
    for dist in all_dists:
        name_lower = dist.get("name", "").strip().lower()
        interest = (dist.get("interest") or "").strip()

        # Petitioner, successor executor, trustee, guardian — keep as-is.
        # These were auto-inserted by this function and have stable roles.
        if name_lower == pet_lower:
            pruned.append(dist)
            continue
        if dist.get("beneficiaryType") == "successor" and interest:
            # Already routed to §7 by the will-beneficiary merge; preserve.
            pruned.append(dist)
            continue

        if _is_actual_distributee(dist):
            dist["beneficiaryType"] = "primary"
            pruned.append(dist)
        elif interest:
            # Not a distributee but has a real will-side interest → §7.
            dist["beneficiaryType"] = "successor"
            pruned.append(dist)
        # Else: stale data (no distributee status, no will interest) → drop.

    all_dists = pruned

    # ── Enhance interest descriptions with roles ─────────────────────────────
    # Prepend Executor and Distributee roles in the canonical interest
    # string. "Distributee" is set when either the explicit flag is true
    # or the relationship maps to an actual EPTL distributee class.
    for dist in all_dists:
        name_lower = dist.get("name", "").strip().lower()
        interest = (dist.get("interest") or "").strip()
        is_primary = (dist.get("beneficiaryType") or "primary") == "primary"
        prefix_parts = []

        if is_primary and "distributee" not in interest.lower():
            if _is_actual_distributee(dist):
                prefix_parts.append("Distributee")

        if name_lower and name_lower == exec_lower and "executor" not in interest.lower():
            prefix_parts.append("Executor")

        if prefix_parts:
            prefix = ", ".join(prefix_parts)
            dist["interest"] = f"{prefix}; {interest}" if interest else prefix

    # Final pass: scrub every entry's interest text through INTEREST_PHRASE_MAP
    # so distributees from data.distributees (which weren't scrubbed during the
    # willBenef merge) also get cleaned.
    for dist in all_dists:
        if dist.get("interest"):
            dist["interest"] = _scrub_interest(dist["interest"])

    return all_dists


def _build_probate_fields(data):
    """Build field name→value dict for Probate Petition + Oath.pdf."""
    proceeding = data.get("proceedingType", "Probate")
    # Auto-compute property values from asset tracker
    _auto_compute_property(data)

    county   = data.get("county", "").upper()
    dec      = decedent_full(data)
    pet      = petitioner_full(data)
    lt       = data.get("lettersType", "")
    # Letters always issued to the petitioner's full legal name. One
    # source of truth — data.lettersTo is intentionally ignored, since
    # using it as an override produced inconsistencies between the
    # petition caption (full name) and the "Letters Testamentary to:"
    # line / waiver "be issued to" (short name).
    letters_to = petitioner_full(data)
    witnesses = " and ".join(filter(None, [data.get("witness1", ""), data.get("witness2", "")]))
    pet_addr  = ", ".join(filter(None, [
        data.get("petitionerStreet", ""), data.get("petitionerCity", ""),
        data.get("petitionerState", ""), data.get("petitionerZip", ""),
    ]))

    # Surviving relatives → Dropdown 5a–5g (EPTL 4-1.1 order, 7 classes)
    #
    # Per NY Surrogate's Court practice for Form P-1 §5:
    #   5a (spouse):   "Yes" / "No"     — binary, always answered.
    #   5b (children): "No" / count     — never "X". Required for the
    #                                     EPTL 4-1.1(a)(1) spouse-plus-
    #                                     issue share calculation.
    #   5c–5g:         "X" / "No" / count — "X" when a closer class is
    #                                     the first surviving (cuts off
    #                                     intestate share); otherwise
    #                                     count if survivors, else "No".
    ft = data.get("ft") or {}
    first_surviving = first_surviving_class(data)

    # Dedup data.distributees by fuzzy name key before counting per-class
    # survivors. Saved cases sometimes carry the same person two or three
    # times (Smart Intake re-runs, hand-edits) — naive iteration would
    # report Amy twice as a surviving spouse.
    _count_seen = set()
    class_counts = [0] * 7
    for dist in (data.get("distributees") or []):
        nm = (dist.get("name") or "").strip()
        if not nm:
            continue
        toks = [t for t in re.split(r'\s+', nm) if t]
        suffixes = {"jr", "sr", "ii", "iii", "iv", "v"}
        while toks and toks[-1].lower().rstrip('.') in suffixes:
            toks.pop()
        key = (toks[0].lower(), toks[-1].lower()) if len(toks) > 1 else (toks[0].lower() if toks else "", "")
        if key in _count_seen:
            continue
        _count_seen.add(key)
        rel = (dist.get("relationship") or "").strip().lower()
        for keyword, cls_idx in DIST_REL_MAP.items():
            if keyword in rel:
                class_counts[cls_idx] += 1
                break

    dropdown_vals = []
    for idx in range(7):
        if idx == 0:
            # Spouse: binary Yes/No
            spouse_alive = (ft.get("spouse") is True) or class_counts[0] > 0
            dropdown_vals.append("Yes" if spouse_alive else "No")
        elif idx == 1:
            # Children: count or "No" — never "X"
            children_alive = (ft.get("children") is True) or class_counts[1] > 0
            if children_alive:
                count = class_counts[1]
                dropdown_vals.append(str(count) if count > 0 else "1")
            else:
                dropdown_vals.append("No")
        else:
            # 5c–5g: X if a closer class is first surviving
            if first_surviving is None:
                dropdown_vals.append("No")
            elif idx < first_surviving:
                dropdown_vals.append("No")
            elif idx == first_surviving:
                count = class_counts[idx]
                dropdown_vals.append(str(count) if count > 0 else "1")
            else:
                dropdown_vals.append("X")

    fields = {
        # ── Petition (pages 1-4) ────────────────────────────────────────────────
        "COUNTY OF": county,
        "To the Surrogates Court County of": county,
        # The "PROBATE PROCEEDING 1" widget sits on the second line of the
        # caption block right after "PROBATE PROCEEDING,". The line is meant
        # to be blank — the decedent name belongs in the "decedent" /
        # "WILL OF:" field, not here.
        "PROBATE PROCEEDING 1": "",
        "File No": data.get("fileNo", ""),
        "decedent": dec,
        "a Name": dec,
        "aka": data.get("decedentAKA", ""),
        "aka2": "",  # second AKA line in caption (renamed from field "1")
        "Name_petitioner": pet,
        "Domicile or Principal Office": data.get("petitionerStreet", ""),
        "City Village or Town": data.get("petitionerCity", ""),
        "State": data.get("petitionerState", ""),
        "Zip Code": data.get("petitionerZip", ""),
        "Citizen of": data.get("petitionerCitizenship", "U.S.A."),
        "b Date of death": data.get("decedentDOD", ""),
        "c Place of death": data.get("decedentPlaceOfDeath", ""),
        "d Domicile Street": data.get("decedentStreet", ""),
        "City Town Village": data.get("decedentCity", ""),
        "County": data.get("decedentCounty", ""),
        "State_2": data.get("decedentState", ""),
        "e Citizen of": data.get("decedentCitizenship", "U.S.A."),
        "Date of Will": data.get("willDate", ""),
        "Names of All Witnesses to Will": witnesses,
        "Date of Codicil": data.get("codicilDate", ""),
        "follows Enter NONE or specify 1": data.get("noOtherWill", "NONE"),
        "the nature of the confidential relationship 1": data.get("confidentialRelationships", "NONE"),
        "Improved real property in New York State": data.get("improvedRealProperty", ""),
        "Unimproved real property in New York State": data.get("unimprovedRealProperty", ""),
        "Estimated gross rents for a period of 18 months": data.get("grossRents18mo", ""),
        "the estate except as follows Enter NONE or specify": data.get("otherAssets", "NONE"),
        "but less than": data.get("personalPropertyValue", ""),
        # Surviving relatives dropdowns
        "Dropdown 5a": dropdown_vals[0],
        "Dropdown 5b": dropdown_vals[1],
        "Dropdown 5c": dropdown_vals[2],
        "Dropdown 5d": dropdown_vals[3],
        "Dropdown 5e": dropdown_vals[4],
        "Dropdown 5f": dropdown_vals[5],
        "Dropdown 5g": dropdown_vals[6],
        # Prayer / letters (page 4) — checkboxes + "to" fields
        "requested": "X" if "Testamentary" in lt else "",           # Letters Testamentary checkbox
        "Petitioner_1": letters_to if "Testamentary" in lt else "",
        "Petitioner_2": "",
        "undefined_7": "X" if "Trusteeship" in lt else "",          # Letters of Trusteeship checkbox
        "Letters of Trusteeship to 1": letters_to if "Trusteeship" in lt else "",
        "undefined_8": "X" if "c.t.a" in lt else "",                # Letters of Admin c.t.a. checkbox
        "Letters of Administration cta to": letters_to if "c.t.a" in lt else "",
        "Dated": "",
        "Print Name": pet,

        # ── Oath and Designation (page 5) ───────────────────────────────────────
        "STATE OF NEW YORK": "New York",
        "COUNTY OF_2": county,
        "OATH OF": "X",  # checkbox — always checked (the oath applies to the petitioner)
        "Surrogates Court of": county,
        "My domicile is": pet_addr,
        "Street Address": "",  # signature line — leave blank for petitioner to sign
        # Print Name_3 = petitioner (the oath-taker). Print Name_4 stays
        # the attorney (notary/preparer block at the bottom).
        "Print Name_3": pet,
        "Print Name_4": data.get("attorneyName") or "Jessica Wilson, Esq.",
        "Firm Name": data.get("firmName") or "Law Office of Jessica Wilson",
        "Tel No": data.get("attorneyPhone") or "(212) 739-1736",
        "Email": data.get("attorneyEmail") or "jwilson@jessicawilsonlaw.com",
        "Address of Attorney": data.get("firmAddress") or "221 Columbia Street, Brooklyn NY 11231",

        # ── Attesting Witness (page 10) ─────────────────────────────────────────
        "COUNTY OF_7": county,
        "WILL OF 1": " ".join(filter(None, [
            data.get("decedentFirstName", ""),
            data.get("decedentMiddleName", ""),
        ])),
        "WILL OF 2": data.get("decedentLastName", ""),
        "aka 1": data.get("decedentAKA", ""),
        "aka 2": "",  # second AKA line on attesting witness
        "File_2": data.get("fileNo", ""),
        "STATE OF NEW YORK_5": "",  # Leave blank — witness fills in their own state
        "COUNTY OF_8": "",  # Leave blank — witness fills in their own county
        "I have been shown check one": "X",  # "the original instrument" checkbox
        # Oath ¶2 checkbox group is set in the letters-type branch below.
        # Field names are visually misaligned with their labels (each
        # widget sits BEFORE its label text on the page), so the field
        # called 'OATH OF' is actually the Executor box, 'EXECUTOR' is
        # the Administrator c.t.a. box, etc. The unconditional
        # "OATH OF": "X" line that used to live here was producing two
        # X marks on every probate petition (the Executor box AND the
        # Administrator c.t.a. box).
        "the original instrument dated": data.get("willDate", ""),
        "purporting to be the last Will and TestamentCodicil of the abovenamed decedent": "",  # "court-certified photographic reproduction" — leave unchecked
        # 'and I saw the other witness es' and 'Print Name_9' are filled
        # per-affiant by generate_probate_docs (one affidavit per witness,
        # each one references the OTHER witness in paragraph 3) — leave
        # blank in the base fields dict.
        "and I saw the other witness es": "",
        "Print Name_9": "",
        "I am making this affidavit at the request of 1": pet,
    }

    # Letters type checkboxes — two distinct groups, both set per letters type:
    #   Page 1 "PETITION FOR PROBATE AND" — fields are named after their
    #     visible labels (Letters Testamentary, Letters of Trusteeship, etc.)
    #   Page 5 "OATH OF [Executor / Administrator c.t.a. / Trustee]" — field
    #     names are visually misaligned. Each widget sits BEFORE its label,
    #     so 'OATH OF' is the Executor box, 'EXECUTOR' is the Administrator
    #     c.t.a. box, 'ADMINISTRATOR cta' is the Trustee box. Set ONLY ONE.
    fields["Letters Testamentary"] = ""
    fields["Letters of Trusteeship"] = ""
    fields["Letters of Administration cta"] = ""
    fields["Temporary Administration"] = ""
    fields["OATH OF"] = ""           # Page-5 Executor option (visually)
    fields["EXECUTOR"] = ""           # Page-5 Administrator c.t.a. option (visually)
    fields["ADMINISTRATOR cta"] = ""  # Page-5 Trustee option (visually)
    fields["Executor"] = ""           # mixed-case — preserved for legacy use elsewhere
    fields["Administrator cta"] = ""  # mixed-case — preserved for legacy use elsewhere

    if "Testamentary" in lt:
        fields["Letters Testamentary"] = "X"
        fields["OATH OF"] = "X"        # Oath of Executor
        fields["Executor"] = "X"
    elif "Trusteeship" in lt:
        fields["Letters of Trusteeship"] = "X"
        fields["ADMINISTRATOR cta"] = "X"  # Oath of Trustee (yes, misnamed)
    elif "c.t.a" in lt:
        fields["Letters of Administration cta"] = "X"
        fields["EXECUTOR"] = "X"       # Oath of Administrator c.t.a. (yes, misnamed)
        fields["Administrator cta"] = "X"
    elif "Temporary" in lt:
        fields["Temporary Administration"] = "X"

    # Petitioner interest — default to Executor for probate
    pet_interest = data.get("petitionerInterest", "")
    if "Executor" in pet_interest or not pet_interest:
        fields["Executor s named in decedents Will"] = "X"
    else:
        fields["Other Specify Check"] = "X"
        fields["Other Specify"] = pet_interest
    if data.get("petitionerIsAttorney") == "Yes":
        fields["is"] = "X"
    else:
        fields["is not an attorney"] = "X"

    # ¶1(c) — Attorney-draftsperson / then-affiliated attorney
    if data.get("petitionerIsDraftsperson") == "Yes":
        fields["is_2"] = "X"
    else:
        fields["is not the attorneydraftsperson a thenaffiliated attorney"] = "X"

    # ── Distributees + will beneficiaries — single canonical source ──────────
    # Both this petition and the Notice of Probate use the same helper, so the
    # two documents render identical "Nature of Interest" text.
    all_dists = compute_interested_persons(data, pet, pet_addr, letters_to)

    # Split into 4 groups
    primary_adults    = [d for d in all_dists if (d.get("beneficiaryType") or "primary") == "primary" and not d.get("isMinor")]
    primary_minors    = [d for d in all_dists if (d.get("beneficiaryType") or "primary") == "primary" and d.get("isMinor")]
    successor_adults  = [d for d in all_dists if d.get("beneficiaryType") == "successor" and not d.get("isMinor")]
    successor_minors  = [d for d in all_dists if d.get("beneficiaryType") == "successor" and d.get("isMinor")]

    def _interest(dist):
        interest = (dist.get("interest") or "").strip()
        if interest:
            return interest
        rel = (dist.get("relationship") or "").strip()
        if proceeding == "Probate":
            return f"Legatee/Devisee under Will" if not rel else f"Legatee/Devisee under Will ({rel})"
        else:
            return f"EPTL 4-1.1 distributee" if not rel else f"EPTL 4-1.1 distributee ({rel})"

    def _name_with_rel(dist):
        name = dist.get("name", "")
        rel = (dist.get("relationship") or "").strip()
        # Strip trailing parenthetical context like "(Amy's sister)" so we
        # don't render nested parens — petition wants the bare blood/
        # marital relation only.
        rel = re.sub(r'\s*\([^)]*\)\s*$', '', rel).strip()
        return f"{name} ({rel})" if rel else name

    def _minor_desc(dist):
        """Build the 7b description: name, DOB, relationship, domicile, guardian."""
        parts = [dist.get("name", "")]
        if dist.get("dob"):
            parts.append(f"DOB: {dist['dob']}")
        if dist.get("relationship"):
            parts.append(dist["relationship"])
        if dist.get("address"):
            parts.append(dist["address"])
        if dist.get("guardianInfo"):
            parts.append(f"Guardian: {dist['guardianInfo']}")
        return "; ".join(parts)

    def _split_interest(text, max_chars=48):
        """Split long interest text into chunks that each fit on ONE
        visible line of an interest cell at 7pt (cell width ~172.7px,
        char width ~3.5px → ~49 chars per visible line). Single source
        of the cell-line capacity — nothing else should pass an
        override. Splits at word boundaries via rfind(' ')."""
        if len(text) <= max_chars:
            return [text]
        lines = []
        remaining = text
        while remaining:
            if len(remaining) <= max_chars:
                lines.append(remaining)
                break
            split_at = remaining.rfind(' ', 0, max_chars)
            if split_at == -1:
                split_at = max_chars
            lines.append(remaining[:split_at + 1].strip())
            remaining = remaining[split_at + 1:].strip()
        return lines

    def _build_rows(persons, max_rows, name_fn=None):
        """One row per chunk of interest text. Name and address appear
        only on the FIRST row of each person; continuation chunks use
        blank name and address — by petition convention, lines without
        a new name belong to the person above. Caps at max_rows."""
        if name_fn is None:
            name_fn = _name_with_rel
        rows = []
        for dist in persons:
            name = name_fn(dist)
            addr = dist.get("address", "")
            interest = _interest(dist)
            chunks = _split_interest(interest)
            rows.append((name, addr, chunks[0]))
            for chunk in chunks[1:]:
                rows.append(("", "", chunk))
            if len(rows) >= max_rows:
                break
        return rows[:max_rows]

    # Page 2, section 6a — Distributees + executor
    p2_6a_name = ["1_2", "2_2", "3", "4", "5", "6", "7"]
    p2_6a_addr = ["1_3", "2_3", "3_2", "4_2", "5_2", "6_2", "7_2"]
    p2_6a_int  = [f"Interest or Nature of Fiduciary Status {i}" for i in range(1, 9)]
    rows_6a = _build_rows(primary_adults, max_rows=7)
    for row, (nm, addr, interest) in enumerate(rows_6a):
        fields[p2_6a_name[row]] = nm
        fields[p2_6a_addr[row]] = addr
        fields[p2_6a_int[row]]  = interest

    # Page 2, section 6b — Primary beneficiaries under disability (6 rows)
    p2_7b_name = ["1_4", "2_4", "3_3", "4_3", "5_3", "6_3"]
    p2_7b_addr = ["1_5", "2_5", "3_4", "4_4", "5_4", "6_4"]
    p2_7b_int  = [f"Interest or Nature of Fiduciary Status {i}_2" for i in range(1, 7)]
    rows_6b = _build_rows(primary_minors, max_rows=6, name_fn=_minor_desc)
    for row, (nm, addr, interest) in enumerate(rows_6b):
        fields[p2_7b_name[row]] = nm
        fields[p2_7b_addr[row]] = addr
        fields[p2_7b_int[row]]  = interest

    # Page 3, section 7a — Other beneficiaries, trustees, successor executors
    p3_7a_name = ["1_9", "2_9", "3_5", "4_5", "5_5", "6_5", "7_3"]
    p3_7a_addr = ["1_10", "2_10", "3_6", "4_6", "5_6", "6_6", "7_4"]
    p3_7a_int  = ["Interest or Nature of Fiduciary Status 1_3",
                  "Interest or Nature of Fiduciary Status 2_3",
                  "Interest or Nature of Fiduciary Status 3_3",
                  "Interest or Nature of Fiduciary Status 4_3",
                  "Interest or Nature of Fiduciary Status 5_3",
                  "Interest or Nature of Fiduciary Status 6_3",
                  "Interest or Nature of Fiduciary Status 7_2",  # not 7_3!
                  "Interest or Nature of Fiduciary Status 8_2"]  # not 8_3!
    rows_7a = _build_rows(successor_adults, max_rows=7)
    for row, (nm, addr, interest) in enumerate(rows_7a):
        fields[p3_7a_name[row]] = nm
        fields[p3_7a_addr[row]] = addr
        fields[p3_7a_int[row]]  = interest

    # Page 3, section 7b — Persons under disability from section 7a (7 rows)
    p3_7b_name = ["1_11", "2_11", "3_7", "4_7", "5_7", "6_7", "7_5"]
    p3_7b_addr = ["1_12", "2_12", "3_8", "4_8", "5_8", "6_8", "7_6"]
    p3_7b_int  = [f"Interest or Nature of Fiduciary Status {i}_4" for i in range(1, 8)]
    rows_7b = _build_rows(successor_minors, max_rows=7, name_fn=_minor_desc)
    for row, (nm, addr, interest) in enumerate(rows_7b):
        fields[p3_7b_name[row]] = nm
        fields[p3_7b_addr[row]] = addr
        fields[p3_7b_int[row]]  = interest

    return fields


def generate_probate_docs(data):
    """
    Returns list of (filename, bytes) for the full probate packet:
      - P-1 Petition (pages 1-4)
      - Combined Verification, Oath and Designation (page 5)
      - Affidavit of Attesting Witness (page 10) — omitted if self-proving will
    Fills the source PDF only once for efficiency.
    """
    template = os.path.join(PROBATE_TEMPLATES_DIR, "Probate Petition + Oath.pdf")
    fields = _build_probate_fields(data)
    # Font overrides: caption fields match template (10pt), interest fields consistent (8pt)
    font_overrides = {}
    # §5 surviving-relatives dropdowns — Yes/No/count/X
    for f in ["Dropdown 5a", "Dropdown 5b", "Dropdown 5c",
              "Dropdown 5d", "Dropdown 5e", "Dropdown 5f", "Dropdown 5g"]:
        font_overrides[f] = 8
    # §1–§5 text fields — caption block, petitioner, decedent, will, no-other-will
    for f in [
        # Caption block
        "COUNTY OF",
        "PROBATE PROCEEDING 1",
        "decedent",
        "aka",
        "aka2",
        "File No",
        "To the Surrogates Court County of",
        # §1(a) Petitioner
        "Name_petitioner",
        "Domicile or Principal Office",
        "City Village or Town",
        "State",
        "Zip Code",
        "Mailing Address",
        "Citizen of",
        "Other Specify",
        # §2 Decedent
        "a Name",
        "b Date of death",
        "c Place of death",
        "d Domicile Street",
        "City Town Village",
        "County",
        "State_2",
        "e Citizen of",
        # §3 Will + codicils
        "Date of Will",
        "Names of All Witnesses to Will",
        "Date of Codicil",
        "Names of All Witnesses to Codicil",
        # §4 No-other-will
        "follows Enter NONE or specify 1",
        # Affidavit of Attesting Witness — witness-side fields
        "and I saw the other witness es",
        "Print Name_9",
        "Street Address_2",
        "TownStateZip",
        "I am making this affidavit at the request of 1",
    ]:
        font_overrides[f] = 8
    # All "X" checkbox cells now auto-size by default in fill_pdf
    # (fontsize=0). No per-field override needed.
    # All interest/description fields at consistent 7pt for readability + fit
    for i in range(1, 9):
        font_overrides[f"Interest or Nature of Fiduciary Status {i}"] = 7
        font_overrides[f"Interest or Nature of Fiduciary Status {i}_2"] = 7
        font_overrides[f"Interest or Nature of Fiduciary Status {i}_3"] = 7
        font_overrides[f"Interest or Nature of Fiduciary Status {i}_4"] = 7
    font_overrides["Interest or Nature of Fiduciary Status 7_2"] = 7
    font_overrides["Interest or Nature of Fiduciary Status 8_2"] = 7
    # §6a name + address fields
    for f in ["1_2","2_2","3","4","5","6","7",        # names
              "1_3","2_3","3_2","4_2","5_2","6_2","7_2"]: # addresses
        font_overrides[f] = 8
    # §6b name + address fields
    for f in ["1_4","2_4","3_3","4_3","5_3","6_3",    # names
              "1_5","2_5","3_4","4_4","5_4","6_4"]:   # addresses
        font_overrides[f] = 8
    # §7a name + address fields
    for f in ["1_9","2_9","3_5","4_5","5_5","6_5","7_3",
              "1_10","2_10","3_6","4_6","5_6","6_6","7_4"]:
        font_overrides[f] = 8
    # §7b name + address fields
    for f in ["1_11","2_11","3_7","4_7","5_7","6_7","7_5",
              "1_12","2_12","3_8","4_8","5_8","6_8","7_6"]:
        font_overrides[f] = 8
    filled = fill_pdf(template, fields, font_overrides=font_overrides)
    last = data.get("decedentLastName", "estate").replace(" ", "_")
    docs = [
        (f"02_Petition_P1_{last}.pdf",        _extract_pages(filled, [0, 1, 2, 3])),
        (f"03_Oath_Designation_{last}.pdf",   _extract_pages(filled, [4])),
    ]

    # ── Affidavit(s) of Attesting Witness ──────────────────────────────────
    # Each witness signs their own affidavit swearing they saw the OTHER
    # witness sign the will. So with two witnesses, generate two
    # affidavits — paragraph 3 references the other witness, Print Name_9
    # holds the affiant's name, and the address fields hold the affiant's
    # address. Skip entirely when the will has a self-proving affidavit.
    if not data.get("selfProvingAffidavit"):
        w1      = (data.get("witness1") or "").strip()
        w1_addr = (data.get("witness1Address") or "").strip()
        w2      = (data.get("witness2") or "").strip()
        w2_addr = (data.get("witness2Address") or "").strip()
        # Pairs: (affiant_name, affiant_addr, other_witness_name)
        affiants = []
        if w1 and w2:
            affiants.append((w1, w1_addr, w2))
            affiants.append((w2, w2_addr, w1))
        elif w1 or w2:
            only      = w1 or w2
            only_addr = w1_addr or w2_addr
            other     = w2 if w1 else w1  # blank
            affiants.append((only, only_addr, other))

        for idx, (affiant, affiant_addr, other) in enumerate(affiants):
            af_fields = dict(fields)
            af_fields["and I saw the other witness es"] = other
            af_fields["Print Name_9"] = affiant
            af_fields["Street Address_2"] = affiant_addr
            af_fields["TownStateZip"] = ""  # leave blank — could be parsed out of addr later
            af_filled = fill_pdf(template, af_fields, font_overrides=font_overrides)
            af_page = _extract_pages(af_filled, [9])
            suffix = chr(ord("a") + idx) if len(affiants) > 1 else ""
            safe_aff = re.sub(r'[^A-Za-z0-9]+', '_', affiant).strip("_") or f"W{idx+1}"
            docs.append((
                f"04{suffix}_Affidavit_Attesting_Witness_{safe_aff}_{last}.pdf",
                af_page,
            ))

    return docs


def fill_probate_pdf(data):
    template = os.path.join(PROBATE_TEMPLATES_DIR, "Probate Petition + Oath.pdf")
    return extract_pdf_pages(template, _build_probate_fields(data), [0, 1, 2, 3])


# ─── ANCILLARY ADMIN PDF (AA-1) ───────────────────────────────────────────────

def fill_ancillary_pdf(data):
    """Fill the AA-1 Ancillary Administration Petition PDF form.

    Field mappings verified against admin_ancil.pdf template:
    - Text Field 19 = Mailing Address (NOT citizenship)
    - Text Field 20 = Citizen of (petitioner 1)
    - Radio Button 2 = Interest of petitioner (/0=Admin, /1=Distributee, /2=Creditor, /3=Other)
    - Text Field 28 = Distributee relationship text
    - Text Field 29 = Other/specify text for interest
    - Text Field 76 = WHEREFORE "Letters to" name (parent-child field)
    - Radio Button 3 = WHEREFORE prayer type (/0=Ancillary Letters, /1=d.b.n.)
    - Text Field 75 = "No other persons interested" paragraph (NOT WHEREFORE)
    """
    dec = decedent_full(data)
    pet = petitioner_full(data)
    # Letters always issued to the petitioner's full legal name. One
    # source of truth — data.lettersTo is intentionally ignored, since
    # using it as an override produced inconsistencies between the
    # petition caption (full name) and the "Letters Testamentary to:"
    # line / waiver "be issued to" (short name).
    letters_to = petitioner_full(data)
    county = data.get("county", "")
    foreign_state = data.get("foreignState", "")

    def v(key, default=""):
        val = str(data.get(key, "") or "").strip()
        return val if val else default

    # Compute total NY property value
    try:
        total = sum(float(data.get(k) or 0) for k in [
            "personalPropertyValue", "improvedRealProperty",
            "unimprovedRealProperty", "grossRents18mo"
        ])
        total_str = f"{total:,.2f}" if total > 0 else "0.00"
    except Exception:
        total_str = ""

    petitioner_address = ", ".join(filter(None, [
        data.get("petitionerStreet", ""),
        data.get("petitionerCity", ""),
        data.get("petitionerState", ""),
        data.get("petitionerZip", "")
    ]))

    # Petitioner interest logic
    pet_interest = v("petitionerInterest", "Distributee")
    is_distributee = pet_interest.lower() == "distributee"

    # Radio button values
    radio_interest_val = "/1" if is_distributee else "/3"
    if pet_interest.lower() == "administrator":
        radio_interest_val = "/0"
    elif pet_interest.lower() == "creditor":
        radio_interest_val = "/2"

    fields = {
        # ── PAGE 1 ────────────────────────────────────────────────
        "Text Field 8":  county,
        "Text Field 9":  dec,
        "Text Field 10": v("decedentAKA"),
        "Text Field 11": foreign_state,
        "Text Field 12": v("fileNo"),
        "Text Field 13": county,

        "Text Field 14": pet,
        "Text Field 15": v("petitionerStreet"),
        "Text Field 16": v("petitionerCity"),
        "Text Field 17": v("petitionerState"),
        "Text Field 18": v("petitionerZip"),
        "Text Field 19": petitioner_address,
        "Text Field 20": v("petitionerCitizenship", "U.S.A."),

        # Interest of petitioner (radio + text)
        "Radio Button 2": radio_interest_val,
        "Text Field 28": v("petitionerRelationship") if is_distributee else "",
        "Text Field 29": "" if is_distributee else pet_interest,

        # Para 2 — Decedent
        "Text Field 30": v("decedentDOD"),
        "Text Field 31": v("decedentPlaceOfDeath"),
        "Text Field 32": v("decedentStreet"),
        "Text Field 33": v("decedentCity"),
        "Text Field 34": v("decedentCounty"),
        "Text Field 35": foreign_state,
        "Text Field 36": v("decedentZip"),
        "Text Field 37": v("decedentCitizenship", "U.S.A."),

        # ── PAGE 2 ────────────────────────────────────────────────
        "Text Field 38": v("foreignLettersDate"),
        "Text Field 39": v("foreignLettersIssuedTo", letters_to),
        "Text Field 40": v("foreignCourtName"),
        "Text Field 41": foreign_state,
        "Text Field 42": v("foreignBondAmount", "0"),

        "Text Field 43": v("personalPropertyValue", "0.00"),
        "Text Field 44": v("improvedRealProperty", "0.00"),
        "Text Field 45": v("unimprovedRealProperty", "0.00"),
        "Text Field 46": v("grossRents18mo", "0.00"),
        "Text Field 47": total_str,

        "Text Field 48": v("otherAssets", "NONE"),
        "Text Field 49": "",

        "Text Field 50": "N/A",

        # ── PAGE 3 ────────────────────────────────────────────────
        # WHEREFORE clause
        "Text Field 76": letters_to,
        "Radio Button 3": "/0",
        "Text Field 1065": "",
        "Text Field 77":   "",
        "Text Field 79":   "NONE",
        "Text Field 80":   "",

        # ── PAGE 4 — Combined Verification, Oath and Designation ──────────────
        "Text Field 85": v("petitionerState", "New York"),
        "Text Field 87": county,
        "Text Field 89": county,
        "Text Field 91": petitioner_address,
        "Text Field 97": pet,
    }

    # Para 6(a) distributees — 3 rows (name / address / interest)
    dist_rows = [
        ("Text Field 57", "Text Field 58", "Text Field 59"),
        ("Text Field 60", "Text Field 61", "Text Field 62"),
        ("Text Field 63", "Text Field 64", "Text Field 65"),
    ]
    for i, dist in enumerate(data.get("distributees", [])[:3]):
        if dist.get("name"):
            nf, af, rf = dist_rows[i]
            fields[nf] = dist["name"]
            fields[af] = dist.get("address", "")
            fields[rf] = dist.get("relationship", "")

    template = os.path.join(PDFS_DIR, "admin_ancil.pdf")
    return fill_pdf(template, fields)


# ─── HELPERS ──────────────────────────────────────────────────────────────────


def decedent_full(data):
    return " ".join(filter(None, [
        data.get("decedentFirstName", ""),
        data.get("decedentMiddleName", ""),
        data.get("decedentLastName", "")
    ]))

def petitioner_full(data):
    return " ".join(filter(None, [
        data.get("petitionerFirstName", ""),
        data.get("petitionerMiddleName", ""),
        data.get("petitionerLastName", "")
    ]))




# ─── ADMINISTRATION PETITION (A-1) ────────────────────────────────────────────

def fill_administration_pdf(data):
    """Fill the A-1 Administration Petition + Oath PDF form."""
    _auto_compute_property(data)
    county    = data.get("county", "")
    dec       = decedent_full(data)
    pet       = petitioner_full(data)
    lt        = data.get("lettersType", "Letters of Administration")
    lt_lower  = lt.lower()
    # Letters always issued to the petitioner's full legal name.
    letters_to = petitioner_full(data)

    def v(key, default=""):
        return str(data.get(key, "") or "").strip() or default

    # Letters type flags
    is_limited    = "limited" in lt_lower and "limitation" not in lt_lower
    is_limitation = "limitation" in lt_lower
    is_temporary  = "temporary" in lt_lower
    is_standard   = not any([is_limited, is_limitation, is_temporary])

    # Citizenship flags
    pet_cit = v("petitionerCitizenship", "U.S.A.")
    dec_cit = v("decedentCitizenship",   "U.S.A.")
    pet_us  = "U.S.A" in pet_cit or "usa" in pet_cit.lower()
    dec_us  = "U.S.A" in dec_cit or "usa" in dec_cit.lower()

    is_attorney = data.get("petitionerIsAttorney") == "Yes"

    # ── Petitioner Interest ────────────────────────────────────────────
    # Administration = no will. The petitioner is normally a distributee
    # (spouse/relative). Override any will-related interest (e.g.,
    # "Executor named in Will" left over from probate-era data) — it
    # never applies to an admin proceeding.
    pet_interest_raw = v("petitionerInterest", "Distributee")
    if "executor" in pet_interest_raw.lower() or "will" in pet_interest_raw.lower():
        pet_interest_raw = "Distributee"
    is_pet_distributee = pet_interest_raw.lower() == "distributee"

    # ── §6 distributees — use the canonical EPTL-routing helper ────────
    pet_addr = ", ".join(filter(None, [
        v("petitionerStreet"), v("petitionerCity"),
        v("petitionerState"), v("petitionerZip"),
    ]))
    all_dists = compute_interested_persons(data, pet, pet_addr, letters_to)
    primary_adults = [d for d in all_dists if (d.get("beneficiaryType") or "primary") == "primary" and not d.get("isMinor")]
    primary_minors = [d for d in all_dists if (d.get("beneficiaryType") or "primary") == "primary" and d.get("isMinor")]

    # ── §5 surviving relatives — ft-driven (matches P-1's rule), 8 classes ─
    # A-1 has 8 classes (Spouse / Children / Issue / Parents / Siblings /
    # Grandparents / Aunts-Uncles / First Cousins). P-1's
    # first_surviving_class returns a 7-class index (no separate Issue).
    # Map: P-1 0→A-1 0, P-1 1→A-1 1, P-1 2+→A-1 3+ (bump for Issue slot).
    ft = data.get("ft") or {}
    p1_first = first_surviving_class(data)
    if p1_first is None:
        a1_first = None
    elif p1_first <= 1:
        a1_first = p1_first
    else:
        a1_first = p1_first + 1

    # Count survivors per A-1 class from the deduped distributee list
    A1_REL_MAP = {
        "spouse": 0, "husband": 0, "wife": 0,
        "son": 1, "daughter": 1, "child": 1, "children": 1,
        "grandchild": 2, "grandson": 2, "granddaughter": 2, "issue": 2,
        "mother": 3, "father": 3, "parent": 3,
        "sister": 4, "brother": 4, "sibling": 4, "niece": 4, "nephew": 4,
        "grandmother": 5, "grandfather": 5, "grandparent": 5,
        "aunt": 6, "uncle": 6, "cousin": 6,
    }
    class_counts = [0] * 8
    for d in primary_adults + primary_minors:
        rel = (d.get("relationship") or "").strip().lower()
        for kw, ci in A1_REL_MAP.items():
            if kw in rel:
                class_counts[ci] += 1
                break

    dropdown_vals = []
    for idx in range(8):
        if idx == 0:
            # Spouse: Yes/No
            alive = (ft.get("spouse") is True) or class_counts[0] > 0
            dropdown_vals.append("Yes" if alive else "No")
        elif idx == 1:
            # Children: No or count — never X
            alive = (ft.get("children") is True) or class_counts[1] > 0
            if alive:
                cnt = class_counts[1]
                dropdown_vals.append(str(cnt) if cnt > 0 else "1")
            else:
                dropdown_vals.append("No")
        else:
            # 6c–6h: X if a closer class is first surviving, else No/count.
            # (6c "Issue of predeceased children" follows the same rule —
            # unlike 6b which must always be No/count for the spouse+issue
            # share calculation.)
            if a1_first is None:
                dropdown_vals.append("No")
            elif idx < a1_first:
                dropdown_vals.append("No")
            elif idx == a1_first:
                cnt = class_counts[idx]
                dropdown_vals.append(str(cnt) if cnt > 0 else "1")
            else:
                dropdown_vals.append("X")

    # Debts
    debt_lines = []
    for key, label in [("mortgageAmount",    "Outstanding Mortgage: ${}"),
                       ("funeralPaid",        "Funeral Expenses Paid: ${}"),
                       ("funeralOutstanding", "Funeral Expenses Outstanding: ${}"),
                       ("miscDebts",          "Misc Debts: {}")]:
        val = (data.get(key, "") or "").strip()
        if val:
            debt_lines.append(label.format(val))
    if not debt_lines:
        debt_lines = ["NONE"]

    pet_addr = ", ".join(filter(None, [
        v("petitionerStreet"), v("petitionerCity"),
        v("petitionerState"), v("petitionerZip"),
    ]))

    fields = {
        # ── PAGE 1: Caption ──────────────────────────────────────────
        "COUNTY OF":                        county.upper(),
        "Estate of 1":                      dec,
        "aka":                              v("decedentAKA"),
        "File No":                          v("fileNo"),
        "TO THE SURROGATES COURT COUNTY OF": county.upper(),

        # Caption checkboxes (letters type)
        "petition for letters of admin":    is_standard,
        "limited admin":                    is_limited,
        "limited admin with lim":           is_limitation,
        "temp admin":                       is_temporary,

        # ── PAGE 1: Petitioner ───────────────────────────────────────
        "Name":                             pet,
        "Domicile":                         v("petitionerStreet"),
        "County":                           v("petitionerCity"),
        "State":                            v("petitionerState"),
        "Zip":                              v("petitionerZip"),
        "yes us citizen":                   pet_us,
        "NO us citizen":                    not pet_us,
        # Distributee path (default for admin) fills the relationship and
        # checks "Distributee". "Other" path is used only for explicit
        # non-distributee interests (Creditor / Designee).
        "Distributee of decedent state relationship":
            v("petitionerRelationship") if is_pet_distributee else "",
        "Otherspecify":                     "" if is_pet_distributee else pet_interest_raw,
        "Mark if Distributee":              is_pet_distributee,
        "Mark if other and then specifiy":  not is_pet_distributee,
        "yes attorney":                     is_attorney,
        "NO not an attorney":               not is_attorney,
        "not a convicted felon":            True,

        # ── PAGE 1: Decedent ─────────────────────────────────────────
        "Name_2":                           dec,
        "Domicile_2":                       v("decedentStreet"),
        "City/Town/Village":                v("decedentCity"),
        "State_2":                          v("decedentState"),
        "Zip Code":                         v("decedentZip"),
        "Township of":                      v("decedentCounty", v("decedentCity")),
        "Date of Death":                    v("decedentDOD"),
        "Place of Death":                   v("decedentPlaceOfDeath"),
        "yes us citizen 1":                 dec_us,
        "NO not US Citizen 2":              not dec_us,

        # ── PAGE 2: Property values ──────────────────────────────────
        # Personal property defaults to "0" (NY convention when no
        # personalty exists or hasn't been valued yet). Real property
        # auto-computes from the asset tracker via _auto_compute_property.
        "gross value personal":             v("personalPropertyValue", "0"),
        "gross value real property":        v("realPropertyValue", "0"),
        "improved":                         bool(nonzero(data.get("improvedRealProperty"))),
        "unimproved":                       bool(nonzero(data.get("unimprovedRealProperty"))),
        "A brief description of each parcel is as follows":
                                            v("realPropertyDescription"),
        "c The estimated gross rent for a period of eighteen 18 months is the sum of":
                                            v("grossRents18mo"),
        # ── §3(d) "right of action / wrongful death asset" — default NONE
        "and the person against whom it exists including names and carrier 1":
                                            v("rightOfAction", "NONE"),
        "and the person against whom it exists including names and carrier 2": "",
        "and the person against whom it exists including names and carrier 3": "",

        # Surviving relatives dropdowns
        "Dropdown 6a": dropdown_vals[0],
        "Dropdown 6b": dropdown_vals[1],
        "Dropdown 6c": dropdown_vals[2],
        "Dropdown 6d": dropdown_vals[3],
        "Dropdown 6e": dropdown_vals[4],
        "Dropdown 6f": dropdown_vals[5],
        "Dropdown 6g": dropdown_vals[6],
        "Dropdown 6h": dropdown_vals[7],

        # ── PAGE 4: Prayer for relief ────────────────────────────────
        "a-process issue letters":          True,
        "c a decree award letters of":      True,
        "9c1":                              is_standard,
        "9c2":                              is_limited,
        "9c3":                              is_limitation,
        "9c4":                              is_temporary,
        "Administration to":                letters_to if is_standard   else "",
        "Limited Administration to":        letters_to if is_limited    else "",
        "Administration with Limitation to": letters_to if is_limitation else "",
        "Temporary Administration to":      letters_to if is_temporary  else "",
        "Dated":                            "",
        "Print Name":                       pet,

        # ── PAGE 1: Petitioner phone (also used on page 5) ───────────
        "Telephone Number":                 v("petitionerPhone", "(212) 739-1736"),

        # ── PAGE 5: Combined Verification, Oath & Designation ────────
        "ss":                               v("petitionerState", "New York"),
        "County of":                        county.upper(),
        "My domicile is":                   pet_addr,
        "before me personally came":        pet,
        "Print Name_3":                     v("attorneyName", "Jessica Wilson, Esq."),
        "Firm Name":                        v("attorneyFirm", "Law Office of Jessica Wilson"),
        "TelNo":                            v("attorneyPhone", "(212) 739-1736"),
        "Address of Attorney":              v("attorneyAddress", "221 Columbia Street, Brooklyn NY 11231"),

        # Wrongful death (always No for standard admin)
        "yes wrongful death":               False,
    }

    # ── PAGE 3: Distributees — driven by compute_interested_persons ──
    # Primary (full-age, sound mind) distributees fill rows 1–8.
    # Minor / under-disability distributees fill rows 1_2–8_2.
    for i, dist in enumerate(primary_adults[:8]):
        n = str(i + 1)
        fields[f"Name {n}"]                         = dist.get("name", "")
        fields[f"Relationship {n}"]                 = dist.get("relationship", "")
        fields[f"Domicile and Mailing Address {n}"] = dist.get("address", "")
        fields[f"Citizenship {n}"]                  = dist.get("citizenship", "U.S.A.")
    for i, dist in enumerate(primary_minors[:8]):
        n = str(i + 1)
        fields[f"Name {n}_2"]                         = dist.get("name", "")
        fields[f"Relationship {n}_2"]                 = dist.get("relationship", "")
        fields[f"Domicile and Mailing Address {n}_2"] = dist.get("address", "")
        fields[f"Citizenship {n}_2"]                  = dist.get("citizenship", "U.S.A.")

    # ── PAGE 3: Debts ────────────────────────────────────────────────
    debt_key = "8 There are no outstanding debts or funeral expenses except Write NONE or state same {}"
    for i, line in enumerate(debt_lines[:9]):
        fields[debt_key.format(i + 1)] = line

    # ── Font overrides ──────────────────────────────────────────────
    font_overrides = {
        # Page-5 "County of" widget is 92x9 — too short for 8pt baseline
        # to land inside. fontsize=0 lets Acrobat auto-size to fit.
        "County of": 0,
        # The 8-class §5 dropdowns are similar.
        "Dropdown 6a": 8, "Dropdown 6b": 8, "Dropdown 6c": 8, "Dropdown 6d": 8,
        "Dropdown 6e": 8, "Dropdown 6f": 8, "Dropdown 6g": 8, "Dropdown 6h": 8,
    }

    template = os.path.join(ADMIN_TEMPLATES_DIR, "Admin Petition + Oath.pdf")
    return fill_pdf(template, fields, font_overrides=font_overrides)


def fill_nondom_pdf(data):
    """Fill the Non-Domiciliary Administration Petition + Oath PDF form.

    Uses the same field mapping as fill_administration_pdf but with the
    Non Dom template which has additional non-domiciliary specific fields.
    """
    county    = data.get("county", "")
    dec       = decedent_full(data)
    pet       = petitioner_full(data)
    lt        = data.get("lettersType", "Letters of Administration")
    lt_lower  = lt.lower()
    # Letters always issued to the petitioner's full legal name. One
    # source of truth — data.lettersTo is intentionally ignored, since
    # using it as an override produced inconsistencies between the
    # petition caption (full name) and the "Letters Testamentary to:"
    # line / waiver "be issued to" (short name).
    letters_to = petitioner_full(data)

    def v(key, default=""):
        return str(data.get(key, "") or "").strip() or default

    is_limited    = "limited" in lt_lower and "limitation" not in lt_lower
    is_limitation = "limitation" in lt_lower
    is_temporary  = "temporary" in lt_lower
    is_standard   = not any([is_limited, is_limitation, is_temporary])

    pet_cit = v("petitionerCitizenship", "U.S.A.")
    dec_cit = v("decedentCitizenship",   "U.S.A.")
    pet_us  = "U.S.A" in pet_cit or "usa" in pet_cit.lower()
    dec_us  = "U.S.A" in dec_cit or "usa" in dec_cit.lower()

    is_attorney = data.get("petitionerIsAttorney") == "Yes"

    surv_keys = [
        "survivingSpouse", "survivingChildren", "survivingIssue",
        "survivingParents", "survivingSiblings", "survivingGrandparents",
        "survivingAuntsUncles", "survivingFirstCousinsOnceRemoved",
    ]
    first_surviving = None
    for idx, key in enumerate(surv_keys):
        raw = data.get(key)
        if raw and str(raw).strip().lower() not in ("false", "0", "no", ""):
            first_surviving = idx
            break
    dropdown_vals = []
    for idx, key in enumerate(surv_keys):
        raw = data.get(key)
        if first_surviving is None:
            dropdown_vals.append("No")
        elif idx < first_surviving:
            dropdown_vals.append("No")
        elif idx == first_surviving:
            s = str(raw).strip()
            dropdown_vals.append(s if s.lower() not in ("true", "yes") else "Yes")
        else:
            dropdown_vals.append("X")

    debt_lines = []
    for key, label in [("mortgageAmount",    "Outstanding Mortgage: ${}"),
                       ("funeralPaid",        "Funeral Expenses Paid: ${}"),
                       ("funeralOutstanding", "Funeral Expenses Outstanding: ${}"),
                       ("miscDebts",          "Misc Debts: {}")]:
        val = (data.get(key, "") or "").strip()
        if val:
            debt_lines.append(label.format(val))
    if not debt_lines:
        debt_lines = ["NONE"]

    pet_addr = ", ".join(filter(None, [
        v("petitionerStreet"), v("petitionerCity"),
        v("petitionerState"), v("petitionerZip"),
    ]))

    # Foreign letters info for non-domiciliary
    foreign_state = v("foreignState", v("decedentState"))

    fields = {
        "COUNTY OF":                        county.upper(),
        "Estate of 1":                      dec,
        "aka":                              v("decedentAKA"),
        "File No":                          v("fileNo"),
        "TO THE SURROGATES COURT COUNTY OF": county.upper(),

        "Name":                             pet,
        "Domicile":                         v("petitionerStreet"),
        "County":                           v("petitionerCity"),
        "State":                            v("petitionerState"),
        "Zip":                              v("petitionerZip"),
        "Mailing address is":               pet_addr,
        "yes us citizen":                   pet_us,
        "NO us citizen":                    not pet_us,
        "Distributee of decedent state relationship":
            v("petitionerRelationship") if v("petitionerInterest", "").lower() in ("", "distributee") else "",
        "Otherspecify":
            "" if v("petitionerInterest", "").lower() in ("", "distributee") else v("petitionerInterest"),
        "Mark if Distributee":
            v("petitionerInterest", "").lower() in ("", "distributee"),
        "Mark if other and then specifiy":
            bool(v("petitionerInterest")) and v("petitionerInterest", "").lower() != "distributee",
        "yes attorney":                     is_attorney,
        "NO not an attorney":               not is_attorney,
        "not a convicted felon":            True,

        "Name_2":                           dec,
        "Domicile_2":                       v("decedentStreet"),
        "City/Town/Village":                v("decedentCity"),
        "State_2":                          v("decedentState"),
        "Zip Code":                         v("decedentZip"),
        "Township of":                      v("decedentCounty", v("decedentCity")),
        "Date of Death":                    v("decedentDOD"),
        "Place of Death":                   v("decedentPlaceOfDeath"),
        "yes us citizen 1":                 dec_us,
        "NO not US Citizen 2":              not dec_us,

        "gross value personal":             v("personalPropertyValue", "0"),
        "gross value real property":        v("realPropertyValue", "0"),
        "improved":                         bool(nonzero(data.get("improvedRealProperty"))),
        "unimproved":                       bool(nonzero(data.get("unimprovedRealProperty"))),
        "A brief description of each parcel is as follows":
                                            v("realPropertyDescription"),
        "c The estimated gross rent for a period of eighteen 18 months is the sum of":
                                            v("grossRents18mo"),

        "Dropdown 6a": dropdown_vals[0],
        "Dropdown 6b": dropdown_vals[1],
        "Dropdown 6c": dropdown_vals[2],
        "Dropdown 6d": dropdown_vals[3],
        "Dropdown 6e": dropdown_vals[4],
        "Dropdown 6f": dropdown_vals[5],
        "Dropdown 6g": dropdown_vals[6],
        "Dropdown 6h": dropdown_vals[7],

        "a-process issue letters":          True,
        "c a decree award letters of":      True,
        "9c1":                              is_standard,
        "9c2":                              is_limited,
        "9c3":                              is_limitation,
        "9c4":                              is_temporary,
        "Administration to":                letters_to if is_standard   else "",
        "Limited Administration to":        letters_to if is_limited    else "",
        "Administration with Limitation to": letters_to if is_limitation else "",
        "Temporary Administration to":      letters_to if is_temporary  else "",
        "Dated":                            "",
        "Print Name":                       pet,

        "Telephone Number":                 v("petitionerPhone", "(212) 739-1736"),

        "ss":                               v("petitionerState", "New York"),
        "My domicile is":                   pet_addr,
        "before me personally came":        pet,
        "Print Name_3":                     v("attorneyName", "Jessica Wilson, Esq."),
        "Firm Name":                        v("attorneyFirm", "Law Office of Jessica Wilson"),
        "TelNo":                            v("attorneyPhone", "(212) 739-1736"),
        "Address of Attorney":              v("attorneyAddress", "221 Columbia Street, Brooklyn NY 11231"),

        "yes wrongful death":               False,
    }

    # Distributees — full age / sound mind (rows 1-8)
    for i, dist in enumerate(data.get("distributees", [])[:8]):
        if dist.get("name"):
            n = str(i + 1)
            fields[f"Name {n}"]                        = dist["name"]
            fields[f"Relationship {n}"]                = dist.get("relationship", "")
            fields[f"Domicile and Mailing Address {n}"] = dist.get("address", "")
            fields[f"Citizenship {n}"]                 = dist.get("citizenship", "U.S.A.")

    # Debts
    debt_key = "8 There are no outstanding debts or funeral expenses except Write NONE or state same {}"
    for i, line in enumerate(debt_lines[:9]):
        fields[debt_key.format(i + 1)] = line

    template = os.path.join(ADMIN_TEMPLATES_DIR, "Non Dom Petition + Oath.pdf")
    return fill_pdf(template, fields)


# ─── FAMILY TREE WORKSHEET (FT-1) ─────────────────────────────────────────────

def fill_ft1_pdf(data):
    """Fill the actual FT-1 Family Tree Affidavit court form PDF."""
    dec_name    = decedent_full(data)
    aka         = data.get("decedentAKA", "")
    file_no     = data.get("fileNo", "")
    pet_name    = petitioner_full(data)
    pet_addr    = ", ".join(filter(None, [
        data.get("petitionerStreet", ""),
        data.get("petitionerCity", ""),
        data.get("petitionerState", "NY"),
        data.get("petitionerZip", ""),
    ]))
    pet_rel     = data.get("petitionerRelationship", "")
    marital     = (data.get("maritalStatus") or "").strip()
    spouse_name = (data.get("spouseName") or "").strip()
    divorce_yr  = (data.get("divorceYear") or "").strip()

    # Distribute distributees into sections by relationship keyword
    all_dists = data.get("distributees", [])

    def _match(d, *keywords):
        return any(k in (d.get("relationship") or "").lower() for k in keywords)

    children  = [d for d in all_dists if _match(d, "child", "son", "daughter")]
    siblings  = [d for d in all_dists if _match(d, "brother", "sister", "sibling")]
    nieces    = [d for d in all_dists if _match(d, "niece", "nephew")]
    mat_aunts = [d for d in all_dists if _match(d, "maternal aunt", "maternal uncle")]
    pat_aunts = [d for d in all_dists if _match(d, "paternal aunt", "paternal uncle")]
    cousins   = [d for d in all_dists if _match(d, "cousin")]

    fields = {}

    # ── Header ──────────────────────────────────────────────────────────────────
    fields["128"]         = dec_name
    fields["230"]         = aka
    fields["412"]         = file_no
    letters_type = (data.get("lettersType") or "Letters of Administration").upper()
    fields["Combo Box00"] = letters_type

    # ── Deponent — the family-tree affiant (form requires someone with no
    # financial interest; UI collects deponent* fields, falling back to the
    # petitioner when blank). Field map (verified against template geometry):
    #   5   = "I, ____" (deponent name)
    #   5a5 = "I AM OVER 18 AND RESIDE AT: ____"
    #   5b6 = "MY RELATIONSHIP TO THE DECEDENT IS ____"
    #   5c7 = "I KNEW THE DECEDENT FOR ____ YEARS"
    #   0   = "BASED ON ____" (left blank for the affiant)
    fields["5"]   = (data.get("deponentName") or "").strip() or pet_name
    fields["5a5"] = (data.get("deponentAddress") or "").strip() or pet_addr
    fields["5b6"] = (data.get("deponentRelationship") or "").strip() or pet_rel
    fields["5c7"] = (data.get("yearsKnown") or "").strip()

    # ── Section 1a: Marriages ───────────────────────────────────────────────────
    if marital == "never_married":
        fields["Check Box01h"] = True
    elif marital == "married" and spouse_name:
        fields["6a9"] = spouse_name
    elif marital == "divorced" and spouse_name:
        fields["6b10"] = spouse_name
        fields["Check Box01a"] = True
        if divorce_yr:
            fields["6a9"] = f"divorced {divorce_yr}"
    elif marital == "widowed" and spouse_name:
        fields["6b10"] = spouse_name
        fields["Check Box01b"] = True

    # ── Section 1b: Children (6 slots; name + date-of-death columns) ───────────
    child_name_f = ["816",  "917",  "1018",  "1119",  "1220",  "1321"]
    child_dod_f  = ["8a22", "9a23", "10a24", "11a25", "12a26", "13a27"]
    for i, c in enumerate(children[:6]):
        if c.get("name"):
            fields[child_name_f[i]] = c["name"]
            dod = c.get("postDeceasedDOD") or c.get("dateOfDeath") or ""
            if dod:
                fields[child_dod_f[i]] = dod

    # ── Section 3a: Siblings (6 slots, page 2) ─────────────────────────────────
    sib_name_f = ["27", "28", "29", "30", "31", "32"]
    for i, s in enumerate(siblings[:6]):
        if s.get("name"):
            fields[sib_name_f[i]] = s["name"]

    # ── Section 3b: Nieces/Nephews (7 slots, page 2) ───────────────────────────
    nie_name_f = ["33","34","35","36","37","38","39"]
    for i, n in enumerate(nieces[:7]):
        if n.get("name"):
            fields[nie_name_f[i]] = n["name"]

    # ── Section 4b: Maternal Aunts/Uncles (7 slots, page 3) ────────────────────
    mat_name_f = ["49","50","51","52","53","54","55"]
    for i, a in enumerate(mat_aunts[:7]):
        if a.get("name"):
            fields[mat_name_f[i]] = a["name"]

    # ── Section 5b: Paternal Aunts/Uncles (7 slots, page 4) ────────────────────
    pat_name_f = ["71","72","73","74","75","76","77"]
    for i, a in enumerate(pat_aunts[:7]):
        if a.get("name"):
            fields[pat_name_f[i]] = a["name"]

    template = os.path.join(ADMIN_TEMPLATES_DIR, "Family_Tree_Affidavit_Fill-In.pdf")
    return fill_pdf(template, fields)


def generate_ft1(data):
    return fill_ft1_pdf(data)


# ─── ACCOUNTING EXCEL ─────────────────────────────────────────────────────────

def _calc_commission(total):
    t1 = min(total, 100000)
    t2 = min(max(total - 100000, 0), 200000)
    t3 = min(max(total - 300000, 0), 700000)
    t4 = max(total - 1000000, 0)
    return t1 * 0.05 + t2 * 0.04 + t3 * 0.03 + t4 * 0.025


def generate_accounting_excel(form_data, assets_data):
    """Generate a full Schedules A–H accounting workbook from asset list data."""
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.utils import get_column_letter

    decedent = decedent_full(form_data)

    wb = Workbook()
    ws = wb.active
    ws.title = "Accounting"

    # ── Styles ────────────────────────────────────────────────────────────────
    GOLD_FILL   = PatternFill("solid", fgColor="7A5C1E")
    LIGHT_FILL  = PatternFill("solid", fgColor="FDF8EE")
    TOTAL_FILL  = PatternFill("solid", fgColor="F4F1EB")
    HDR_FONT    = Font(name="Arial", bold=True, color="FFFFFF", size=11)
    BOLD        = Font(name="Arial", bold=True, size=11)
    NORMAL      = Font(name="Arial", size=11)
    MONEY       = Font(name="Courier New", size=11)
    LABEL_FONT  = Font(name="Arial", bold=True, size=11)
    thin = Side(style="thin", color="DDDDDD")
    BORDER = Border(bottom=Side(style="thin", color="DDDDDD"))

    def money_fmt(cell):
        cell.number_format = '#,##0.00'
        cell.font = MONEY

    def section_header(row, title):
        ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=3)
        c = ws.cell(row=row, column=1, value=title)
        c.font = HDR_FONT
        c.fill = GOLD_FILL
        c.alignment = Alignment(horizontal="left", vertical="center", indent=1)
        ws.row_dimensions[row].height = 22

    def col_headers(row, *headers):
        for i, h in enumerate(headers, 1):
            c = ws.cell(row=row, column=i, value=h)
            c.font = BOLD
            c.fill = LIGHT_FILL
            c.alignment = Alignment(horizontal="left" if i < len(headers) else "right")

    def total_row(row, label, value):
        c1 = ws.cell(row=row, column=1, value=label)
        c1.font = BOLD
        c1.fill = TOTAL_FILL
        c3 = ws.cell(row=row, column=3, value=value)
        c3.font = Font(name="Courier New", bold=True, size=11)
        c3.fill = TOTAL_FILL
        c3.number_format = '#,##0.00'
        c3.alignment = Alignment(horizontal="right")

    def blank_rows(start_row, count):
        for r in range(start_row, start_row + count):
            ws.cell(row=r, column=1).border = BORDER
            ws.cell(row=r, column=3).border = BORDER
            ws.cell(row=r, column=3).number_format = '#,##0.00'
            ws.cell(row=r, column=3).alignment = Alignment(horizontal="right")

    # ── Column widths ─────────────────────────────────────────────────────────
    ws.column_dimensions['A'].width = 40
    ws.column_dimensions['B'].width = 20
    ws.column_dimensions['C'].width = 16

    # ── Title ─────────────────────────────────────────────────────────────────
    r = 1
    ws.merge_cells(start_row=r, start_column=1, end_row=r, end_column=3)
    title_cell = ws.cell(row=r, column=1,
        value=f"Estate of {decedent} — Informal Accounting")
    title_cell.font = Font(name="Arial", bold=True, size=13)
    title_cell.alignment = Alignment(horizontal="center")
    ws.row_dimensions[r].height = 24
    r += 1

    date_cell = ws.cell(row=r, column=1, value=f"Generated: {today()}")
    date_cell.font = NORMAL
    r += 2

    # ── Schedule A — Estate Assets ────────────────────────────────────────────
    section_header(r, "Schedule A — Estate Assets"); r += 1
    col_headers(r, "Institution / Description", "Category", "Value ($)"); r += 1

    sched_a_total = 0.0
    for a in assets_data:
        val = 0.0
        try:
            val = float(str(a.get("value", "0")).replace(",", "").replace("$", "").strip() or 0)
        except Exception:
            pass
        sched_a_total += val
        c1 = ws.cell(row=r, column=1, value=a.get("institution") or a.get("category", ""))
        c1.font = NORMAL
        c2 = ws.cell(row=r, column=2, value=a.get("category", ""))
        c2.font = NORMAL
        c3 = ws.cell(row=r, column=3, value=val)
        money_fmt(c3)
        c3.alignment = Alignment(horizontal="right")
        r += 1
    total_row(r, "Schedule A Total", sched_a_total); r += 2

    # ── Schedule B — Income / Receipts ────────────────────────────────────────
    section_header(r, "Schedule B — Income / Receipts"); r += 1
    col_headers(r, "Description", "", "Amount ($)"); r += 1
    b_start = r
    blank_rows(r, 10); r += 10
    total_row(r, "Schedule B Subtotal", 0); r += 2

    # ── Schedule C — Disbursements ────────────────────────────────────────────
    section_header(r, "Schedule C — Disbursements"); r += 1
    col_headers(r, "Description", "", "Amount ($)"); r += 1
    blank_rows(r, 10); r += 10
    total_row(r, "Schedule C Subtotal", 0); r += 2

    # ── Schedule D — Prior Distributions ─────────────────────────────────────
    section_header(r, "Schedule D — Prior Distributions"); r += 1
    ws.cell(row=r, column=1, value="Prior Distributions").font = NORMAL
    d_cell = ws.cell(row=r, column=3)
    d_cell.number_format = '#,##0.00'
    d_cell.alignment = Alignment(horizontal="right")
    d_cell.border = BORDER
    r += 1
    total_row(r, "Schedule D Total", 0); r += 2

    # ── Schedule E — Commission ───────────────────────────────────────────────
    section_header(r, "Schedule E — Executor/Administrator Commission (NY SCPA)"); r += 1
    commission = _calc_commission(sched_a_total)
    tiers = [
        ("First $100,000 × 5%",       min(sched_a_total, 100000),         0.05),
        ("Next $200,000 × 4%",         min(max(sched_a_total - 100000, 0), 200000), 0.04),
        ("Next $700,000 × 3%",         min(max(sched_a_total - 300000, 0), 700000), 0.03),
        ("Balance over $1,000,000 × 2.5%", max(sched_a_total - 1000000, 0), 0.025),
    ]
    note = ws.cell(row=r, column=1,
        value=f"Commission base (Schedule A total): ${sched_a_total:,.2f}")
    note.font = Font(name="Arial", italic=True, size=10, color="888888")
    r += 1
    for label, base, rate in tiers:
        if base > 0:
            c1 = ws.cell(row=r, column=1, value=label)
            c1.font = NORMAL
            c3 = ws.cell(row=r, column=3, value=base * rate)
            money_fmt(c3)
            c3.alignment = Alignment(horizontal="right")
            r += 1
    total_row(r, "Total Commission", commission); r += 2

    # ── Schedule F — Estate Account Balance ───────────────────────────────────
    section_header(r, "Schedule F — Balance in Estate Account"); r += 1
    ws.cell(row=r, column=1, value="Current balance in estate account").font = NORMAL
    f_cell = ws.cell(row=r, column=3)
    f_cell.number_format = '#,##0.00'
    f_cell.alignment = Alignment(horizontal="right")
    f_cell.border = BORDER
    r += 1
    total_row(r, "Schedule F Balance", 0); r += 2

    # ── Schedule G — Reconciliation ───────────────────────────────────────────
    section_header(r, "Schedule G — Reconciliation"); r += 1
    rows_g = [
        ("Schedule A + B (Total Receipts)", sched_a_total),
        ("Less: Schedule C (Disbursements)", 0),
        ("Less: Schedule D (Prior Distributions)", 0),
        ("Net Balance", sched_a_total),
        ("Schedule F (Estate Account Balance)", 0),
        ("Difference (should be zero)", sched_a_total),
    ]
    for label, val in rows_g:
        c1 = ws.cell(row=r, column=1, value=label)
        c1.font = BOLD if "Net Balance" in label or "Difference" in label else NORMAL
        c3 = ws.cell(row=r, column=3, value=val)
        money_fmt(c3)
        c3.alignment = Alignment(horizontal="right")
        r += 1
    r += 1

    # ── Schedule H — Distribution Plan ────────────────────────────────────────
    section_header(r, "Schedule H — Distribution Plan"); r += 1
    col_headers(r, "Beneficiary / Purpose", "", "Amount ($)"); r += 1
    c1 = ws.cell(row=r, column=1, value="Executor/Administrator Commission (from Sched E)")
    c1.font = NORMAL
    c3 = ws.cell(row=r, column=3, value=commission)
    money_fmt(c3)
    c3.alignment = Alignment(horizontal="right")
    r += 1
    blank_rows(r, 10); r += 10
    total_row(r, "Schedule H Total", 0); r += 1

    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    return buf.read()


# ─── LETTER OF AUTHORIZATION ──────────────────────────────────────────────────

def generate_auth_letter(data, asset):
    """Pre-letters letter from nominated executor/administrator authorizing the
    law office to speak with the institution."""
    lt = data.get("lettersType", "")
    role = "executor" if "Testamentary" in lt else "administrator"
    decedent = decedent_full(data)
    petitioner = petitioner_full(data)
    institution = asset.get("institution", "").strip() or "Financial Institution"
    account_no = asset.get("accountNumber", "").strip() or "N/A"

    doc = Document()

    FONT = "Times New Roman"
    SIZE = Pt(12)

    def _run(para, text, bold=False):
        r = para.add_run(text)
        r.font.name = FONT
        r.font.size = SIZE
        r.bold = bold
        return r

    def line(text="", bold=False, space_after=6):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(space_after)
        if text:
            _run(p, text, bold=bold)
        return p

    line(today(), space_after=12)
    line("")
    line(institution, bold=True)
    line(f"Re: Estate of {decedent}")
    line(f"    Account No.: {account_no}", space_after=12)
    line("")
    line("To Whom It May Concern:", space_after=12)
    line("")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    _run(p, f"I, {petitioner}, am the nominated {role} of the Estate of "
         f"{decedent}, deceased. I hereby authorize the Law Office of Jessica Wilson "
         f"to discuss, obtain information about, and act on my behalf with respect to "
         f"the above-referenced account and any other accounts held by the above-named estate.")
    line("Please extend your full cooperation to our office upon request.", space_after=24)
    line("")
    line("Sincerely,", space_after=48)
    line("")
    line("")
    line(petitioner)

    _validate_docx(doc, "generate_auth_letter")
    return make_docx_bytes(doc)


# ─── LETTER OF INSTRUCTION ────────────────────────────────────────────────────

def generate_instruction_letter(data, asset, marshal_action="check"):
    """Post-letters letter requesting the institution marshal assets."""
    lt = data.get("lettersType", "")
    role = "executor" if "Testamentary" in lt else "administrator"
    letters_label = "Letters Testamentary" if "Testamentary" in lt else "Letters of Administration"
    decedent = decedent_full(data)
    petitioner = petitioner_full(data)
    county = data.get("county", "")
    dod = data.get("decedentDOD", "")
    institution = asset.get("institution", "").strip() or "Financial Institution"
    account_no = asset.get("accountNumber", "").strip() or "N/A"
    signer_key = data.get("signer", "Jessica Wilson")
    signer = SIGNERS.get(signer_key, signer_key)

    if marshal_action == "transfer":
        marshal_text = "transfer all funds to the estate account"
    else:
        marshal_text = f"remit payment by check payable to 'Estate of {decedent}'"

    doc = Document()

    FONT = "Times New Roman"
    SIZE = Pt(12)

    def _run(para, text, bold=False):
        r = para.add_run(text)
        r.font.name = FONT
        r.font.size = SIZE
        r.bold = bold
        return r

    def line(text="", bold=False, space_after=6):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(space_after)
        if text:
            _run(p, text, bold=bold)
        return p

    line(today(), space_after=12)
    line("")
    line(institution, bold=True)
    line(f"Re: Estate of {decedent}")
    line(f"    Account No.: {account_no}", space_after=12)
    line("")
    line("Dear Sir or Madam:", space_after=12)
    line("")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    _run(p, f"Our office represents {petitioner}, the duly appointed {role} of the "
         f"Estate of {decedent}, who died on {dod}. "
         f"{letters_label} were issued by the Surrogate's Court, {county} County.")
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(12)
    _run(p2, f"Please marshal all assets held in the above-referenced account and "
         f"{marshal_text} at your earliest convenience. "
         f"Please find enclosed a certified copy of the Letters.")
    line("Please do not hesitate to contact our office should you require any additional "
         "information or documentation.", space_after=24)
    line("")
    line("Very truly yours,", space_after=48)
    line("")
    line("")
    line(signer)

    _validate_docx(doc, "generate_instruction_letter")
    return make_docx_bytes(doc)


# ─── SCHEDULE D(a) — POST-DECEASED DISTRIBUTEE ──────────────────────────────

def fill_schedule_da_pdf(data, dist):
    """Fill the Schedule D(a) form for a distributee who post-deceased the decedent.

    Field mapping (by rect position on page):
    - County:         County (caption — replaces hard-coded "NEW YORK")
    - Text Field167:  File #
    - Text Field164:  Estate of (decedent name)
    - Text Field165:  a/k/a
    - Text Field168:  1. Name of post-deceased distributee
    - Text Field169:  Date of Death of post-deceased
    - Text Field170:  Relationship to decedent
    - Text Field171:  Last permanent address (domicile)
    - Check Box01:    Yes/No fiduciary appointed
    - Text Field174:  3(a) Fiduciary row 1 (name / address / citizenship / court)
    - Text Field176:  3(a) Fiduciary row 2
    - Text Field175:  3(b) Distributee row 1
    - Text Field177:  3(b) Distributee row 2
    - Text Field178:  3(b) Distributee row 3
    - Text Field179:  3(b) Distributee row 4
    - Text Field180:  3(b) Distributee row 5
    - Text Field181:  3(b) Distributee row 6
    """
    dec = decedent_full(data)
    aka = data.get("decedentAKA", "")
    county = data.get("county", "")
    file_no = data.get("fileNo", "")

    fields = {}

    # Header
    fields["County"] = county
    fields["Text Field167"] = file_no
    fields["Text Field164"] = dec
    fields["Text Field165"] = aka

    # Section 1 — post-deceased distributee info
    fields["Text Field168"] = dist.get("name", "")
    fields["Text Field169"] = dist.get("postDeceasedDOD", "")
    fields["Text Field170"] = dist.get("relationship", "")
    fields["Text Field171"] = dist.get("address", "")

    # Section 2 — fiduciary yes/no
    has_fiduciary = dist.get("hasFiduciary", False)
    if has_fiduciary:
        fields["Check Box01"] = True

    # Section 3(a) — fiduciary details (2 rows)
    fid = dist.get("fiduciary", {})
    if has_fiduciary and fid:
        row1_parts = [fid.get("name", ""), fid.get("address", ""),
                      fid.get("citizenship", ""), fid.get("court", "")]
        fields["Text Field174"] = "     ".join(p for p in row1_parts if p)
        fields["Text Field176"] = fid.get("row2", "")

    # Section 3(b) — post-deceased person's distributees (up to 6 rows)
    pd_dists = dist.get("postDeceasedDistributees", [])
    row_fields = ["Text Field175", "Text Field177", "Text Field178",
                  "Text Field179", "Text Field180", "Text Field181"]
    for idx, pd in enumerate(pd_dists[:6]):
        parts = [pd.get("name", ""), pd.get("address", ""),
                 pd.get("citizenship", ""), pd.get("relationship", "")]
        fields[row_fields[idx]] = "     ".join(p for p in parts if p)

    template = os.path.join(PROBATE_TEMPLATES_DIR,
                            "Schedule D(a)- Distributee Who Post-Deceased Decedent.pdf")
    return fill_pdf(template, fields)


# ─── ADMIN CTA (ADMINISTRATION C.T.A.) ──────────────────────────────────────

def fill_cta_pdf(data):
    """Fill the Administration C.T.A. petition PDF (SCPA 1418/1419).

    Used when a will exists but the named executor cannot serve
    (died, resigned, or was removed).

    Template: templates/Probate/probcta.pdf  (7 pages)
    Page 0: CTA-1 Petition (sections 1-2)
    Page 1: CTA-1 Petition cont (sections 3-7, WHEREFORE, signatures)
    Page 2: Combined Verification, Oath & Designation
    Page 3: Corporate Verification, Consent & Designation
    Page 4: CTA Citation
    Page 5: CTA-3 Waiver/Renunciation
    Page 6: P-12 Affidavit of No Debt
    """
    county   = data.get("county", "")
    dec      = decedent_full(data)
    aka      = data.get("decedentAKA", "")
    pet      = petitioner_full(data)
    file_no  = data.get("fileNo", "")
    # Letters always issued to the petitioner's full legal name. One
    # source of truth — data.lettersTo is intentionally ignored, since
    # using it as an override produced inconsistencies between the
    # petition caption (full name) and the "Letters Testamentary to:"
    # line / waiver "be issued to" (short name).
    letters_to = petitioner_full(data)

    def v(key, default=""):
        return str(data.get(key, "") or "").strip() or default

    pet_street = v("petitionerStreet")
    pet_city   = v("petitionerCity")
    pet_state  = v("petitionerState", "NY")
    pet_zip    = v("petitionerZip")
    pet_cit    = v("petitionerCitizenship", "US Citizen")
    dec_street = v("decedentStreet")
    dec_city   = v("decedentCity")
    dec_county = v("decedentCounty", county)
    dec_state  = v("decedentState", "NY")
    dec_zip    = v("decedentZip")

    # CTA-specific fields
    orig_county   = v("ctaOriginalCounty")
    orig_date     = v("ctaOriginalDate")
    orig_executor = v("ctaOriginalExecutor")
    exec_reason   = v("ctaExecutorReason")  # died / resigned / removed
    pet_interest  = v("ctaPetitionerInterest", "Residuary Beneficiary")
    is_attorney   = v("ctaAdminIsAttorney", "no")

    # Estate values
    personal = v("personalPropertyValue", "0")
    real_imp = v("improvedRealProperty", "0")
    real_unimp = v("unimprovedRealProperty", "0")
    gross_rents = v("grossRents18mo", "0")

    # Distributees
    dists = data.get("distributees", [])

    fields = {}

    # ═══ PAGE 0: CTA-1 Petition ═══════════════════════════════════════════
    fields["Decedent_Name"]   = dec
    fields["Decedent_AKA"]    = aka
    fields["File_No"]         = file_no
    # "County" covers both the caption COUNTY OF and the
    # "TO THE SURROGATE'S COURT, COUNTY OF" line (same field name).
    fields["County"]          = county

    # Section 1(a) — Petitioner info
    # TextField13 = trailing blank on the "is/are as follows:" lead-in (left
    # empty); 14 = Name; 15/16 = Domicile (Street and Number)/(City, Village
    # or Town); 17-20 = County/State/Zip/Telephone; 21 = Mailing address.
    fields["TextField13[0]"]  = ""
    fields["TextField14[0]"]  = pet                   # Petitioner name
    fields["TextField15[0]"]  = pet_street            # Street and Number
    fields["TextField16[0]"]  = pet_city              # City/Village/Town
    fields["TextField17[0]"]  = v("petitionerCounty", county)  # County
    fields["TextField18[0]"]  = pet_state             # State
    fields["TextField19[0]"]  = pet_zip               # Zip
    fields["TextField20[0]"]  = ""                    # Telephone
    fields["TextField21[0]"]  = ""                    # Mailing address if different

    # Citizenship checkboxes (normalize "U.S.A." → "usa" before matching)
    _cit = pet_cit.lower().replace(".", "").replace(" ", "")
    if "us" in _cit or "citizen" in _cit or "america" in _cit:
        fields["CheckBox1[0]"] = True   # USA
    else:
        fields["CheckBox2[0]"] = True   # Other
        fields["TextField22[0]"] = pet_cit

    # Second petitioner (leave blank)
    fields["TextField23[0]"]  = ""

    # Interest checkboxes
    # Template geometry: CheckBox5 = Sole Beneficiary, CheckBox7 = Residuary
    # Beneficiary (same row, right), CheckBox6 = Other [Specify] (next row).
    if "sole" in pet_interest.lower():
        fields["CheckBox5[0]"] = True   # Sole Beneficiary
    elif "residuary" in pet_interest.lower():
        fields["CheckBox7[0]"] = True   # Residuary Beneficiary
    else:
        fields["CheckBox6[0]"] = True   # Other
        fields["TextField32[0]"] = pet_interest

    # 1(b) — Is admin CTA an attorney?
    if is_attorney.lower() == "yes":
        fields["CheckBox8[0]"] = True   # is an attorney
    else:
        fields["CheckBox9[0]"] = True   # is not an attorney

    # Section 2 — Original probate info
    fields["TextField33[0]"]  = orig_county           # County where probated
    fields["TextField34[0]"]  = orig_date             # Date probated
    fields["TextField35[0]"]  = orig_executor         # Original executor name
    fields["TextField36[0]"]  = ""                    # "who on [date]..."

    # Reason checkboxes
    if exec_reason == "died":
        fields["CheckBox10[0]"] = True
    elif exec_reason == "resigned":
        fields["CheckBox11[0]"] = True
    elif exec_reason == "removed":
        fields["CheckBox12[0]"] = True

    # ═══ PAGE 1: Petition continued ═══════════════════════════════════════
    # Section 3 — Persons with prior/equal right (SCPA 1418)
    if len(dists) > 0:
        d = dists[0]
        fields["TextField37[0]"]  = d.get("name", "")
        fields["TextField38[0]"]  = ""                 # Description of legacy
        fields["TextField39[0]"]  = d.get("relationship", "")
        fields["TextField40[0]"]  = ""                 # Mailing address
        fields["TextField41[0]"]  = ""                 # Fiduciary status
        fields["TextField42[0]"]  = d.get("address", "")
        fields["TextField43[0]"]  = ""                 # Additional line

    # Section 4 — Other beneficiaries
    if len(dists) > 1:
        d = dists[1]
        fields["TextField44[0]"]  = d.get("name", "")
        fields["TextField45[0]"]  = ""
        fields["TextField46[0]"]  = d.get("relationship", "")
        fields["TextField47[0]"]  = ""
        fields["TextField67[0]"]  = ""
        fields["TextField49[0]"]  = d.get("address", "")
        fields["TextField50[0]"]  = ""

    # Section 6 — Debts
    fields["TextField51[0]"]  = ""    # Debts/funeral expenses (leave for manual)

    # Section 7 — Estate values
    fields["TextField52[0]"]  = personal    # Personal property
    fields["TextField53[0]"]  = real_imp    # Improved real property
    fields["TextField54[0]"]  = real_unimp  # Unimproved real property
    fields["TextField55[0]"]  = gross_rents # Estimated gross rents 18 months
    fields["TextField56[0]"]  = ""          # Other assets / cause of action

    # WHEREFORE
    fields["Petitioner"]      = letters_to   # Letters of Admin CTA to
    fields["TextField58[0]"]  = ""           # Other relief
    fields["TextField59[0]"]  = today()      # Dated

    # Petitioner signature block — signature lines stay blank for wet
    # signatures; only the Print Name line is filled.
    fields["TextField60[0]"]  = ""           # Signature line 1 (wet signature)
    fields["TextField61[0]"]  = ""           # Signature line 2
    fields["TextField62[0]"]  = pet          # Print name 1
    fields["TextField63[0]"]  = ""           # Print name 2

    # ═══ PAGE 2: Verification, Oath & Designation ═════════════════════════
    fields["STATE OF_F02"]             = "NEW YORK"
    fields["COUNTY OF_F13"]            = county.upper()
    fields["of_F24"]                   = county  # Designation county
    fields["(Street Address)_F35"]     = pet_street
    fields["(City/Town/Village)_F46"]  = pet_city
    fields["(State)_F57"]              = pet_state
    fields["(Print Name)_F68"]         = pet
    fields["came_F79"]                 = pet     # "came [name]"
    fields["Date0"]                    = today()
    fields["Year1"]                    = ""
    # Attorney block at the bottom of the oath page
    atty_name  = data.get("attorneyName")  or "Jessica Wilson, Esq."
    atty_firm  = data.get("firmName")      or "Law Office of Jessica Wilson PC"
    atty_phone = data.get("attorneyPhone") or "(212) 739-1736"
    atty_addr  = data.get("attorneyAddress") or ", ".join(filter(None, [
        data.get("firmAddress", "221 Columbia Street"),
        data.get("firmAddress2", "Brooklyn, New York 11231"),
    ]))
    fields["Print Name_F911"]          = atty_name
    fields["Firm Name_F1012"]          = atty_firm
    fields["Tel No_F1113"]             = atty_phone
    fields["Address of Attorney_F1214"] = atty_addr

    # ═══ PAGE 3: Corporate Verification ═══════════════════════════════════
    fields["TextField86[0]"]  = ""    # State
    fields["TextField87[0]"]  = ""    # County
    # Corporate fields left blank (filled when corporate petitioner)

    # ═══ PAGE 4: Citation ═════════════════════════════════════════════════
    fields["TextField108[0]"] = file_no               # File No
    fields["TextField109[0]"] = county                 # County
    # TO lines (cite parties)
    cite_names = [d.get("name", "") for d in dists if d.get("disposition") == "citation"]
    to_fields = ["TextField110[0]", "TextField111[0]", "TextField112[0]",
                 "TextField113[0]", "TextField114[0]"]
    for i, name in enumerate(cite_names[:5]):
        fields[to_fields[i]] = name

    fields["TextField115[0]"] = pet                    # Petitioner name
    fields["TextField116[0]"] = pet_street             # Petitioner domicile
    fields["TextField117[0]"] = f"{pet_city}, {pet_state}"
    fields["TextField118[0]"] = county                 # County
    # TextField120 = citation return date ("on ___") — court fills this in.
    fields["TextField120[0]"] = ""
    fields["TextField123[0]"] = f"{dec_street}, {dec_city}, {dec_state}"  # Domicile
    fields["TextField124[0]"] = dec                    # "estate of ___"
    fields["TextField125[0]"] = ""                     # Surrogate name
    fields["TextField126[0]"] = letters_to             # Letters to
    fields["TextField135[0]"] = atty_name              # Attorney for petitioner
    fields["TextField136[0]"] = atty_phone             # Telephone
    fields["TextField137[0]"] = atty_addr              # Address of attorney

    # ═══ PAGE 5: CTA-3 Waiver/Renunciation ═══════════════════════════════
    fields["TextField138[0]"] = county                 # County
    fields["TextField139[0]"] = dec                    # Will of
    fields["TextField140[0]"] = aka                    # a/k/a
    fields["TextField141[0]"] = file_no                # File No
    fields["TextField142[0]"] = ""                     # Undersigned name (filled by signer)

    # Interest checkboxes (page 5)
    # CheckBox13 = beneficiary with equal/prior right
    # CheckBox14 = beneficiary of estate
    # CheckBox15 = creditor
    # CheckBox16 = other

    # TextField144 = "appears ... in the Surrogate's Court of ___ County"
    fields["TextField144[0]"] = county
    # TextField145 = ¶3 "Consents that Letters ... be granted by the Court to ___"
    fields["TextField145[0]"] = letters_to
    # Jurat COUNTY OF (TextField152) left blank — the notary fills it in
    # wherever the waiver is actually signed (per firm practice).
    # Attorney block at the bottom
    fields["TextField158[0]"] = atty_name              # Name of Attorney
    fields["TextField159[0]"] = atty_phone             # Tel. No.
    fields["TextField160[0]"] = atty_addr              # Address of Attorney

    # ═══ PAGE 6: P-12 Affidavit of No Debt ═══════════════════════════════
    fields["TextField161[0]"] = county                 # County
    fields["TextField162[0]"] = dec                    # Will of
    fields["TextField163[0]"] = aka                    # a/k/a
    fields["TextField164[0]"] = file_no                # File No
    fields["TextField165[0]"] = county                 # County (SS:)
    fields["TextField166[0]"] = pet                    # Deponent name
    fields["TextField170[0]"] = pet_street             # Resides at
    fields["TextField168[0]"] = dec_county             # County of residence
    fields["TextField169[0]"] = pet_state              # State
    fields["TextField171[0]"] = personal               # Estate value
    fields["TextField172[0]"] = v("miscDebts", "NONE") # [If "none", write "NONE"]
    fields["TextField189[0]"] = pet                    # Print Name (signature stays blank)
    fields["TextField191[0]"] = atty_name              # Name of Attorney
    fields["TextField192[0]"] = atty_phone             # Tel. No.
    fields["TextField193[0]"] = atty_addr              # Address of Attorney

    template = os.path.join(PROBATE_TEMPLATES_DIR, "probcta.pdf")
    return fill_pdf(template, fields)


# ─── WAIVER OF CONSENT AND RENUNCIATION (A-8 Individual) ────────────────────

def fill_waiver_individual_pdf(data, dist):
    """Fill the A-8 Waiver, Renunciation and Consent to Appoint Administrator
    (official NYSBA Form A-8, 2 pages) for an individual distributee.

    Template: templates/Admin/Waiver of Consent and Renunciation.pdf — built
    from the official form with fields added at every blank. Field names:

    Page 1: county, estate_of, aka, file_no, court_county,
            cb_letters / cb_letters_limits / cb_limited (checkboxes),
            be_issued_to, cb_bond_dispensed / cb_bond_amount, bond_amount,
            dated, signature, print_name, street, city, state, zip,
            country, relationship
    Page 2: notary_state, notary_county, notary_date, notary_appeared,
            atty_name, atty_firm, atty_phone, atty_address, atty_email
            (attorney block runs across the bottom under the notary).
    """
    import re as _re
    dec = decedent_full(data)
    aka = data.get("decedentAKA", "")
    county = data.get("county", "")
    file_no = data.get("fileNo", "")
    # Letters always issued to the petitioner's full legal name. One
    # source of truth — data.lettersTo is intentionally ignored.
    letters_to = petitioner_full(data)

    dist_name = dist.get("name", "")
    dist_addr = (dist.get("address", "") or "").strip()
    dist_rel = dist.get("relationship", "")

    # Split "street, city, ST zip" into the form's separate boxes.
    street, city, state, zipc = dist_addr, "", "", ""
    m = _re.match(r"^(.*?),\s*(.*?),\s*([A-Za-z]{2})\.?\s*(\d{5}(?:-\d{4})?)?\s*$", dist_addr)
    if m:
        street, city, state, zipc = m.group(1), m.group(2), m.group(3).upper(), m.group(4) or ""
    elif ", " in dist_addr:
        street, city = dist_addr.split(", ", 1)

    lt = (data.get("lettersType") or "Letters of Administration").lower()
    cb_limits  = "limitation" in lt
    cb_limited = lt.startswith("limited")
    cb_letters = not (cb_limits or cb_limited)

    atty_addr = data.get("attorneyAddress") or ", ".join(filter(None, [
        data.get("firmAddress", "221 Columbia Street"),
        data.get("firmAddress2", "Brooklyn, New York 11231"),
    ]))

    fields = {
        "county":        county.upper(),
        "estate_of":     dec,
        "aka":           aka,
        "file_no":       file_no,
        "court_county":  county,
        "cb_letters":        cb_letters,
        "cb_letters_limits": cb_limits,
        "cb_limited":        cb_limited,
        "be_issued_to":  letters_to,
        "cb_bond_dispensed": bool(data.get("dispenseBond")),
        "print_name":    dist_name,
        "street":        street,
        "city":          city,
        "state":         state,
        "zip":           zipc,
        "country":       "USA" if (dist.get("citizenship", "") or "").upper().startswith("U") else "",
        "relationship":  dist_rel,
        # Notary STATE OF / COUNTY OF left blank — the notary fills these
        # in wherever the waiver is actually signed.
        "atty_name":     data.get("attorneyName")  or "Jessica Wilson, Esq.",
        "atty_firm":     data.get("firmName")      or "Law Office of Jessica Wilson PC",
        "atty_phone":    data.get("attorneyPhone") or "(212) 739-1736",
        "atty_address":  atty_addr,
        "atty_email":    data.get("attorneyEmail") or "jwilson@jessicawilsonlaw.com",
    }

    template = os.path.join(ADMIN_TEMPLATES_DIR, "Waiver of Consent and Renunciation.pdf")
    return fill_pdf(template, fields)


# ─── WAIVER & CONSENT CORPORATE (A-9) ───────────────────────────────────────

def fill_waiver_corporate_pdf(data, dist):
    """Fill the A-9 Waiver & Consent form for a corporate distributee.

    Field mapping (Waiver & Consent Corp.pdf):
    - county of 112:                     County
    - Estate of 112:                     Estate name (decedent)
    - aka of 112:                        a/k/a
    - File No_8:                         File number
    - Name of Corporation:               Corporation name
    - a citation ... be issued to:       Letters to (administrator name)
    - COUNTY OF_6:                       County (notary section)
    - Name of Attorney_2:               Attorney name
    - Firm Name_2:                       Firm name
    - Address_2:                         Attorney address
    - Telephone Number_3:               Attorney phone
    - Email_2:                           Attorney email
    """
    dec = decedent_full(data)
    aka = data.get("decedentAKA", "")
    county = data.get("county", "")
    file_no = data.get("fileNo", "")
    # Letters always issued to the petitioner's full legal name. One
    # source of truth — data.lettersTo is intentionally ignored, since
    # using it as an override produced inconsistencies between the
    # petition caption (full name) and the "Letters Testamentary to:"
    # line / waiver "be issued to" (short name).
    letters_to = petitioner_full(data)

    corp_name = dist.get("name", "")

    fields = {
        "county of 112":     county,
        "Estate of 112":     dec,
        "aka of 112":        aka,
        "File No_8":         file_no,
        "Name of Corporation": corp_name,
        "a citation in this matter and consents that Letters of Administration be issued to": letters_to,
        # Notary COUNTY OF left blank (filled by the notary at signing)
        # Attorney block — two rows across the bottom under the notary:
        #   Name of Attorney | Firm Name
        #   Address          | Telephone | Email
        "Name of Attorney_2": data.get("attorneyName")  or "Jessica Wilson, Esq.",
        "Firm Name_2":        data.get("firmName")      or "Law Office of Jessica Wilson PC",
        "Address_2":          data.get("attorneyAddress") or ", ".join(filter(None, [
                                  data.get("firmAddress", "221 Columbia Street"),
                                  data.get("firmAddress2", "Brooklyn, New York 11231")])),
        "Telephone Number_3": data.get("attorneyPhone") or "(212) 739-1736",
        "Email_2":            data.get("attorneyEmail") or "jwilson@jessicawilsonlaw.com",
    }

    template = os.path.join(ADMIN_TEMPLATES_DIR, "Waiver & Consent Corp.pdf")
    return fill_pdf(template, fields)


# ─── CITATION (Admin) ───────────────────────────────────────────────────────

def fill_citation_pdf(data):
    """Fill the Citation PDF form (post-filing document).

    Field mapping (Citation.pdf):
    - SURROGATES COURT:                  County
    - File No_2:                         File number
    - A petition having been duly filed by: Petitioner name
    - who is domicilied at:              Petitioner address
    - county 444:                        County
    - decree should not be made in the estate of: Decedent name
    - decree should not be made in the estate of 222: a/k/a
    - lately domiciled at:               Decedent address
    - lately domiciled at the county of: Decedent county
    - estate of decentt to 222345:       Letters to name
    - Attorney for Petitioner:           Attorney name
    - TelNo_2:                           Attorney phone
    - Address of Attorney_2:             Attorney address
    """
    dec = decedent_full(data)
    pet = petitioner_full(data)
    county = data.get("county", "")
    file_no = data.get("fileNo", "")
    # Letters always issued to the petitioner's full legal name. One
    # source of truth — data.lettersTo is intentionally ignored, since
    # using it as an override produced inconsistencies between the
    # petition caption (full name) and the "Letters Testamentary to:"
    # line / waiver "be issued to" (short name).
    letters_to = petitioner_full(data)

    pet_addr = ", ".join(filter(None, [
        data.get("petitionerStreet", ""),
        data.get("petitionerCity", ""),
        data.get("petitionerState", ""),
        data.get("petitionerZip", ""),
    ]))
    dec_addr = ", ".join(filter(None, [
        data.get("decedentStreet", ""),
        data.get("decedentCity", ""),
        data.get("decedentState", ""),
        data.get("decedentZip", ""),
    ]))

    fields = {
        "SURROGATES COURT":                 county,
        "File No_2":                        file_no,
        "A petition having been duly filed by": pet,
        "who is domicilied at":             pet_addr,
        "county 444":                       county,
        "decree should not be made in the estate of": dec,
        "decree should not be made in the estate of 222": data.get("decedentAKA", ""),
        "lately domiciled at":              dec_addr,
        "lately domiciled at the county of": data.get("decedentCounty", county),
        "estate of decentt to 222345":      letters_to,
        "Attorney for Petitioner":          data.get("attorneyName", "Jessica Wilson, Esq."),
        "TelNo_2":                          data.get("attorneyPhone", "(212) 739-1736"),
        "Address of Attorney_2":            data.get("attorneyAddress", "221 Columbia Street, Brooklyn NY 11231"),
    }

    template = os.path.join(ADMIN_TEMPLATES_DIR, "Citation.pdf")
    return fill_pdf(template, fields)


# ─── AFFIDAVIT OF SERVICE ───────────────────────────────────────────────────

def fill_affidavit_of_service_pdf(data):
    """Fill the Affidavit of Service PDF header fields (post-filing document).

    Only fills county, estate, and file number. Person-served details are
    completed manually after service is actually made.

    Field mapping (Affid of Service.pdf):
    - county of 113:    County
    - Estate of 113:    Estate name (decedent)
    - File No_9:        File number
    """
    dec = decedent_full(data)
    county = data.get("county", "")
    file_no = data.get("fileNo", "")

    fields = {
        "county of 113":  county,
        "Estate of 113":  dec,
        "File No_9":      file_no,
    }

    template = os.path.join(ADMIN_TEMPLATES_DIR, "Affid of Service.pdf")
    return fill_pdf(template, fields)


# ─── NOTICE OF APPLICATION (SCPA 1005) ──────────────────────────────────────

def fill_notice_of_application_pdf(data):
    """Fill the Notice of Application (SCPA 1005) PDF form.

    Field mapping (Notice of App SCPA 1005.pdf):
    - County of 56:      County
    - Estate of 56:      Estate name (decedent)
    - aka of 56:         a/k/a
    - File No_3:         File number
    - petitioner:        Petitioner name
    - 3 petitioner prays...: Letters to name
    - Name of Distributee 1/2/3:              Distributee names (section 4a)
    - Domicile and Post Office Address 1/2/3: Distributee addresses (section 4a)
    - Name of Distributee 1_2/2_2/3_2:       Distributee names (section 4b)
    - Domicile and Post Office Address 1_2/2_2/3_2: Distributee addresses (section 4b)
    - Attorney for Petitioner_2:             Attorney name
    - Print Name_4:                          Attorney print name
    - Address Office:                        Attorney address
    """
    dec = decedent_full(data)
    pet = petitioner_full(data)
    county = data.get("county", "")
    file_no = data.get("fileNo", "")
    aka = data.get("decedentAKA", "")
    # Letters always issued to the petitioner's full legal name. One
    # source of truth — data.lettersTo is intentionally ignored, since
    # using it as an override produced inconsistencies between the
    # petition caption (full name) and the "Letters Testamentary to:"
    # line / waiver "be issued to" (short name).
    letters_to = petitioner_full(data)

    fields = {
        "County of 56":    county,
        "Estate of 56":    dec,
        "aka of 56":       aka,
        "File No_3":       file_no,
        "petitioner":      pet,
        "1 an application for Letters of Administration upon the estate of the abovenamed decedent has been made": "",
        "3 petitioner prays that a decree be made directing the issuance of Letters of Administration to": letters_to,
        "Attorney for Petitioner_2": data.get("attorneyName", "Jessica Wilson, Esq."),
        "Print Name_4":    data.get("attorneyName", "Jessica Wilson, Esq."),
        "Address Office":  data.get("attorneyAddress", "221 Columbia Street, Brooklyn NY 11231"),
    }

    # Section 4a — distributees with full age and sound mind (up to 3)
    dists = data.get("distributees", [])
    name_fields_a = ["Name of Distributee 1", "Name of Distributee 2", "Name of Distributee 3"]
    addr_fields_a = ["Domicile and Post Office Address 1", "Domicile and Post Office Address 2",
                     "Domicile and Post Office Address 3"]
    name_fields_b = ["Name of Distributee 1_2", "Name of Distributee 2_2", "Name of Distributee 3_2"]
    addr_fields_b = ["Domicile and Post Office Address 1_2", "Domicile and Post Office Address 2_2",
                     "Domicile and Post Office Address 3_2"]

    for i, dist in enumerate(dists[:3]):
        if dist.get("name"):
            fields[name_fields_a[i]] = dist["name"]
            fields[addr_fields_a[i]] = dist.get("address", "")

    template = os.path.join(ADMIN_TEMPLATES_DIR, "Notice of App SCPA 1005.pdf")
    return fill_pdf(template, fields)


# ─── AFFIDAVIT OF MAILING ───────────────────────────────────────────────────

def fill_affidavit_of_mailing_pdf(data):
    """Fill the Affidavit of Mailing PDF form header and distributee addresses.

    Field mapping (Affid of Mailing.pdf):
    - County of 57:      County
    - Estate of 57:      Estate name (decedent)
    - aka of 57:         a/k/a
    - File No_4:         File number
    - COUNTY OF_2:       County (venue)
    - whose post office address is / _2 / _3 / _4 / _5 / _6 / _7 / _8:
                         Distributee addresses (up to 8)
    """
    dec = decedent_full(data)
    county = data.get("county", "")
    file_no = data.get("fileNo", "")
    aka = data.get("decedentAKA", "")

    fields = {
        "County of 57":   county,
        "Estate of 57":   dec,
        "aka of 57":      aka,
        "File No_4":      file_no,
        "COUNTY OF_2":    county,
    }

    # Fill distributee addresses (up to 8 slots)
    addr_fields = [
        "whose post office address is",
        "whose post office address is_2",
        "whose post office address is_3",
        "whose post office address is_4",
        "whose post office address is_5",
        "whose post office address is_6",
        "whose post office address is_7",
        "whose post office address is_8",
    ]
    dists = data.get("distributees", [])
    for i, dist in enumerate(dists[:8]):
        if dist.get("name"):
            addr = dist.get("address", "")
            fields[addr_fields[i]] = f"{dist['name']}, {addr}" if addr else dist["name"]

    template = os.path.join(ADMIN_TEMPLATES_DIR, "Affid of Mailing.pdf")
    return fill_pdf(template, fields)


# ─── AFFIDAVIT OF REGULARITY ────────────────────────────────────────────────

def fill_affidavit_of_regularity_pdf(data):
    """Fill the Affidavit of Regularity PDF form (post-filing document).

    Field mapping (Affid of Regularity.pdf):
    - county of 10:      County
    - Estate of 10:      Estate name (decedent)
    - aka of10:          a/k/a
    - File No_6:         File number
    - COUNTY OF_4:       County (venue)
    - being duly sworn deposes and says: Attorney name (deponent)
    - 1 That heshe is the attorney for: Petitioner name
    - Name 1_5 / Name 2_5:              Waiver distributee names (section c)
    - Address 1_3 / Address 2_3:         Waiver distributee addresses (section c)
    - Name 1_3 / Name 2_3:              Citation distributee names (section a)
    - Address 1 / Address 2:            Citation distributee addresses (section a)
    """
    dec = decedent_full(data)
    pet = petitioner_full(data)
    county = data.get("county", "")
    file_no = data.get("fileNo", "")
    aka = data.get("decedentAKA", "")
    attorney = data.get("attorneyName", "Jessica Wilson, Esq.")

    fields = {
        "county of 10":    county,
        "Estate of 10":    dec,
        "aka of10":        aka,
        "File No_6":       file_no,
        "COUNTY OF_4":     county,
        "being duly sworn deposes and says": attorney,
        "1 That heshe is the attorney for": pet,
    }

    # Separate distributees by disposition
    dists = data.get("distributees", [])
    waiver_dists = [d for d in dists if d.get("disposition") == "waiver" and d.get("name")]
    citation_dists = [d for d in dists if d.get("disposition") == "citation" and d.get("name")]

    # Section (c) — waivers (up to 2)
    waiver_name_fields = ["Name 1_5", "Name 2_5"]
    waiver_addr_fields = ["Address 1_3", "Address 2_3"]
    for i, d in enumerate(waiver_dists[:2]):
        fields[waiver_name_fields[i]] = d["name"]
        fields[waiver_addr_fields[i]] = d.get("address", "")

    # Section (a) — citations (up to 2)
    cite_name_fields = ["Name 1_3", "Name 2_3"]
    cite_addr_fields = ["Address 1", "Address 2"]
    for i, d in enumerate(citation_dists[:2]):
        fields[cite_name_fields[i]] = d["name"]
        fields[cite_addr_fields[i]] = d.get("address", "")

    template = os.path.join(ADMIN_TEMPLATES_DIR, "Affid of Regularity.pdf")
    return fill_pdf(template, fields)


# ─── PROPOSED DECREE ─────────────────────────────────────────────────────────

def fill_proposed_decree_pdf(data):
    """Fill the Proposed Decree PDF form (post-filing document).

    Field mapping (Proposed Decree.pdf):
    - in and for the County of:          County
    - Estate of 9:                       Estate name (decedent)
    - aka of 9:                          a/k/a
    - FileNo:                            File number
    - A petition having been filed by:   Petitioner name
    - of the goods chattels...:          Letters to name
    - that:                              Petitioner name (competency statement)
    - ORDERED AND DECREED...:            Letters to name
    - ORDERED AND DECREED... 22:         Letters to name (bond dispensed)
    - bond having been filed and approved...: Bond amount
    - bond having been filed:            Checkbox — bond filed
    - bond having been dispensed:        Checkbox — bond dispensed
    """
    dec = decedent_full(data)
    pet = petitioner_full(data)
    county = data.get("county", "")
    file_no = data.get("fileNo", "")
    aka = data.get("decedentAKA", "")
    # Letters always issued to the petitioner's full legal name. One
    # source of truth — data.lettersTo is intentionally ignored, since
    # using it as an override produced inconsistencies between the
    # petition caption (full name) and the "Letters Testamentary to:"
    # line / waiver "be issued to" (short name).
    letters_to = petitioner_full(data)
    bond_amount = data.get("bondAmount", "")

    fields = {
        "in and for the County of":       county,
        "Estate of 9":                    dec,
        "aka of 9":                       aka,
        "FileNo":                         file_no,
        "A petition having been filed by": pet,
        "of the goods chattels and credits of the abovenamed decedent be granted to": letters_to,
        "that":                           pet,
        "is in all respects competent to act as administrat": "",
        "ORDERED AND DECREED that Letters of Administration issue to": letters_to,
        "ORDERED AND DECREED that Letters of Administration issue to 22": letters_to,
    }

    # Bond: filed vs dispensed
    if bond_amount and bond_amount.strip() not in ("0", ""):
        fields["bond having been filed"] = True
        fields["bond having been filed and approved in the amount of"] = bond_amount
    else:
        fields["bond having been dispensed"] = True

    template = os.path.join(ADMIN_TEMPLATES_DIR, "Proposed Decree.pdf")
    return fill_pdf(template, fields)


# ─── SCHEDULE A — NONMARITAL PERSONS ────────────────────────────────────────

def fill_schedule_a_pdf(data, dist):
    """Fill Schedule A (Nonmarital Persons) for a per-distributee schedule.

    Field mapping (Schedule A Nonmarital Persons.pdf):
    - County of 2:             County
    - Estate of 2:             Estate name (decedent)
    - aka of 2:                a/k/a
    - File:                    File number
    - Name of alleged distributee: Distributee name
    - Date of birth:           Distributee DOB
    - Relationship to decedent: Relationship
    - Name of father:          Father name
    - Name of mother:          Mother name
    """
    dec = decedent_full(data)
    county = data.get("county", "")
    file_no = data.get("fileNo", "")
    aka = data.get("decedentAKA", "")

    fields = {
        "County of 2":              county,
        "Estate of 2":              dec,
        "aka of 2":                 aka,
        "File":                     file_no,
        "Name of alleged distributee": dist.get("name", ""),
        "Date of birth":            dist.get("dob", ""),
        "Relationship to decedent": dist.get("relationship", ""),
        "Name of father":           dist.get("fatherName", ""),
        "Name of mother":           dist.get("motherName", ""),
    }

    template = os.path.join(ADMIN_TEMPLATES_DIR, "Schedule A Nonmarital Persons.pdf")
    return fill_pdf(template, fields)


# ─── SCHEDULE B — ADOPTION ──────────────────────────────────────────────────

def fill_schedule_b_pdf(data, dist):
    """Fill Schedule B (Adoption) for a per-distributee schedule.

    Field mapping (Sched B Adoption.pdf):
    - County of 3:             County
    - Estate of 3:             Estate name (decedent)
    - aka of 3:                a/k/a
    - File_2:                  File number
    - Name of child:           Adopted child name
    - Relationship to decedent prior to adoption: Prior relationship
    - Date of adoption:        Adoption date
    - If yesname of adoptive father or mother: Adoptive parent name
    - Name of the adoptive parent: Adoptive parent name
    """
    dec = decedent_full(data)
    county = data.get("county", "")
    file_no = data.get("fileNo", "")
    aka = data.get("decedentAKA", "")

    fields = {
        "County of 3":     county,
        "Estate of 3":     dec,
        "aka of 3":        aka,
        "File_2":          file_no,
        "Name of child":   dist.get("name", ""),
        "Relationship to decedent prior to adoption": dist.get("priorRelationship", ""),
        "Date of adoption": dist.get("adoptionDate", ""),
        "If yesname of adoptive father or mother": dist.get("adoptiveParent", ""),
        "Name of the adoptive parent": dist.get("adoptiveParent", ""),
    }

    template = os.path.join(ADMIN_TEMPLATES_DIR, "Sched B Adoption.pdf")
    return fill_pdf(template, fields)


# ─── SCHEDULE C — INFANTS ───────────────────────────────────────────────────

def fill_schedule_c_pdf(data, dist):
    """Fill Schedule C (Infants) for a per-distributee schedule.

    Field mapping (Sched C Infants.pdf):
    - County of 4:             County
    - Estate of 4:             Estate name (decedent)
    - aka of 4:                a/k/a
    - File_3:                  File number
    - Name_3:                  Infant name
    - Date of birth 1:         DOB line 1
    - Date of birth 2:         DOB line 2
    - Relationship to the decedent: Relationship
    - With whom does the infant reside: Residence info
    - Name of mother_2:        Mother name
    - Is she alive:            Mother alive
    - Name of Father:          Father name
    - Is he alive:             Father alive
    - If yes name and address of guardian: Guardian info
    """
    dec = decedent_full(data)
    county = data.get("county", "")
    file_no = data.get("fileNo", "")
    aka = data.get("decedentAKA", "")

    fields = {
        "County of 4":     county,
        "Estate of 4":     dec,
        "aka of 4":        aka,
        "File_3":          file_no,
        "Name_3":          dist.get("name", ""),
        "Date of birth 1": dist.get("dob", ""),
        "Relationship to the decedent": dist.get("relationship", ""),
        "With whom does the infant reside": dist.get("residesWithWhom", ""),
        "Name of mother_2": dist.get("motherName", ""),
        "Is she alive":    dist.get("motherAlive", ""),
        "Name of Father":  dist.get("fatherName", ""),
        "Is he alive":     dist.get("fatherAlive", ""),
        "If yes name and address of guardian": dist.get("guardianInfo", ""),
    }

    template = os.path.join(ADMIN_TEMPLATES_DIR, "Sched C Infants.pdf")
    return fill_pdf(template, fields)


# ─── SCHEDULE D — DISABILITY ────────────────────────────────────────────────

def fill_schedule_d_pdf(data, dist):
    """Fill Schedule D (Disability) for a per-distributee schedule.

    Field mapping (Sched D Disability.pdf):
    - County of 5:             County
    - Estate of 5:             Estate name (decedent)
    - aka of 5:                a/k/a
    - File_4:                  File number
    - 1 Name:                  Person's name
    - Relationship:            Relationship to decedent
    - Residence:               Residence address
    - With whom does this person reside: Caretaker info
    - If this person is in prison name of prison: Prison name
    - If yesgive nametitle and address 1: Court-appointed attorney line 1
    - If yesgive nametitle and address 2: Court-appointed attorney line 2
    - If nodescribe nature of disability 1: Disability description line 1
    - If nodescribe nature of disability 2: Disability description line 2
    - If nogive name and address of relative or friend... 1: Interested person line 1
    - If nogive name and address of relative or friend... 2: Interested person line 2
    """
    dec = decedent_full(data)
    county = data.get("county", "")
    file_no = data.get("fileNo", "")
    aka = data.get("decedentAKA", "")

    fields = {
        "County of 5":     county,
        "Estate of 5":     dec,
        "aka of 5":        aka,
        "File_4":          file_no,
        "1 Name":          dist.get("name", ""),
        "Relationship":    dist.get("relationship", ""),
        "Residence":       dist.get("address", ""),
        "With whom does this person reside": dist.get("residesWithWhom", ""),
        "If this person is in prison name of prison": dist.get("prisonName", ""),
        "If yesgive nametitle and address 1": dist.get("courtAttorneyInfo", ""),
        "If yesgive nametitle and address 2": "",
        "If nodescribe nature of disability 1": dist.get("disabilityDescription", ""),
        "If nodescribe nature of disability 2": "",
        "If nogive name and address of relative or friend interested in his or her welfare 1":
            dist.get("interestedPerson", ""),
        "If nogive name and address of relative or friend interested in his or her welfare 2": "",
    }

    template = os.path.join(ADMIN_TEMPLATES_DIR, "Sched D Disability.pdf")
    return fill_pdf(template, fields)


# ─── WORD TEMPLATE GENERATORS ───────────────────────────────────────────────

def generate_waiver_probate(data, dist):
    """Generate the P-4 Waiver of Process, Consent to Probate Word document
    for a Probate proceeding distributee.

    Template placeholders (Waiver_Probate.docx):
    - _________________  (county, in 'County of ___')
    - No actual bracket-style placeholders; uses blanks for manual fill.
    We replace the county blank and leave signature blanks for manual completion.
    """
    import re as _re
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document(os.path.join(WORD_TEMPLATES_DIR, "Waiver_Probate.docx"))
    county    = data.get("county", "")
    dec       = decedent_full(data)
    aka       = data.get("decedentAKA", "")
    file_no   = data.get("fileNo", "")
    will_date = data.get("willDate", "") or "____________"
    pet       = petitioner_full(data)
    dist_name = dist.get("name", "")
    dist_addr = dist.get("address", "")
    dist_rel  = dist.get("relationship", "")
    year      = datetime.now().strftime("%Y")


    def _all_paras():
        for p in doc.paragraphs:
            yield p
        for t in doc.tables:
            for row in t.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        yield p

    def _sub(p, rx, repl):
        """Regex-substitute inside paragraph p, preserving run formatting
        when the match sits inside one run; else collapse into runs[0]."""
        if not rx.search(p.text):
            return False
        hit = False
        for run in p.runs:
            if rx.search(run.text):
                run.text = rx.sub(repl, run.text, count=1)
                hit = True
        # Only fall back to paragraph-level collapse when no single run
        # contained the match (split across runs). Re-testing p.text
        # would double-apply when the replacement itself still matches.
        if not hit:
            full = rx.sub(repl, p.text, count=1)
            if p.runs:
                p.runs[0].text = full
                for r in p.runs[1:]:
                    r.text = ""
            else:
                p.add_run(full)
        return True

    # ── Caption table ──────────────────────────────────────────────
    cap = doc.tables[0]
    for p in cap.rows[1].cells[0].paragraphs:          # COUNTY OF ______
        _sub(p, _re.compile(r"COUNTY OF\s*_+"), f"COUNTY OF {county.upper()}")
    for p in cap.rows[5].cells[1].paragraphs:          # WILL OF  ______
        _sub(p, _re.compile(r"_+"), dec)
    if aka:
        for p in cap.rows[6].cells[1].paragraphs:      # a/k/a  ______
            _sub(p, _re.compile(r"_+"), aka)
    else:
        # No a/k/a — blank out both the label and the underline so the
        # caption doesn't show a dangling "a/k/a ________".
        for ci in (0, 1):
            for p in cap.rows[6].cells[ci].paragraphs:
                for r in p.runs:
                    r.text = ""
    for p in cap.rows[6].cells[3].paragraphs:          # File No.
        _sub(p, _re.compile(r"File No\.\s*"), f"File No. {file_no}")

    # ── Body text ──────────────────────────────────────────────────
    for p in doc.paragraphs:
        _sub(p, _re.compile(r"County of\s*_+"), f"County of {county}")
        _sub(p, _re.compile(r"Testament\s+dated\s*_+"), f"Testament dated {will_date}")
        # Notary jurat year: "____, 2025," → current year
        _sub(p, _re.compile(r"(On\s+_+,\s*)\d{4}"), rf"\g<1>{year}")

    # ── Letters / signature table ──────────────────────────────────
    sig = doc.tables[1]
    for p in sig.rows[0].cells[0].paragraphs:          # [X] Letters Testamentary issue to ____
        _sub(p, _re.compile(r"issue to\s*_+"), f"issue to {pet}")
    for p in sig.rows[4].cells[0].paragraphs:          # date line "_____, 2024"
        _sub(p, _re.compile(r"(_+,\s*)\d{4}"), rf"\g<1>{year}")
    # (signature line in rows[4].cells[1] stays blank for the wet signature)
    for p in sig.rows[5].cells[1].paragraphs:          # "Signature of ____"
        _sub(p, _re.compile(r"Signature of\s*_+"), f"Signature of {dist_name}")
    for p in sig.rows[4].cells[3].paragraphs:          # relationship (was hardcoded "Son")
        _sub(p, _re.compile(r"\S.*"), dist_rel or "______________")
    if dist_addr:                                      # address line above "(address)"
        # Give the address column more room (template cell is only ~1.2")
        # by borrowing from the wide signature column to its left.
        from docx.shared import Inches as _In
        from docx.oxml.ns import qn as _qn
        grid = sig._tbl.tblGrid.findall(_qn("w:gridCol"))
        if len(grid) >= 3:
            grid[1].set(_qn("w:w"), str(int(2.15 * 1440)))
            grid[2].set(_qn("w:w"), str(int(1.81 * 1440)))
        for row in list(sig.rows)[4:7]:
            row.cells[1].width = _In(2.15)
            row.cells[2].width = _In(1.81)
        for p in sig.rows[5].cells[2].paragraphs:
            if not p.text.strip():
                # Street on one line, city/state/zip on the next
                parts = dist_addr.split(", ", 1)
                r = p.add_run("\n".join(parts))
                r.font.name = "Arial"
                r.font.size = Pt(9)
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                break

    # ── Strip the yellow fill-in highlights everywhere ─────────────
    for p in _all_paras():
        for r in p.runs:
            if r.font.highlight_color is not None:
                r.font.highlight_color = None

    # Attorney block: the template's own stacked table (Print Name / Firm
    # Name / Tel No. / Address of Attorney) at the bottom is kept exactly
    # as-is — that's the layout the firm files (see Grego waiver).

    _validate_docx(doc, "generate_waiver_probate")
    return make_docx_bytes(doc)


def generate_notice_of_probate(data):
    """Generate the Notice of Probate + Affidavit of Mailing Word document.

    Fills caption, will date, decedent domicile, petitioner address, today's
    date, and the recipient table (distributees + will beneficiaries with
    addresses and roles).
    """
    import re as _re
    doc = Document(os.path.join(WORD_TEMPLATES_DIR, "Notice_of_Probate.docx"))
    dec       = decedent_full(data)
    pet       = petitioner_full(data)
    county    = data.get("county", "")
    aka       = data.get("decedentAKA", "")
    will_date = data.get("willDate", "") or "___________"

    def _addr_str(street, city, state, zipc, blank):
        # Only fill if at least street is present — partial fragments look broken
        if not street:
            return blank
        return ", ".join(filter(None, [street, city, state, zipc]))

    dec_addr = _addr_str(
        data.get("decedentStreet", ""), data.get("decedentCity", ""),
        data.get("decedentState", ""), data.get("decedentZip", ""),
        "_________________",
    )
    pet_addr = _addr_str(
        data.get("petitionerStreet", ""), data.get("petitionerCity", ""),
        data.get("petitionerState", ""), data.get("petitionerZip", ""),
        "_______________________",
    )

    today_str = datetime.now().strftime("%B %d, %Y")

    # Strip the orphan "aka [DECEDENT AKA]" / "a/k/a [DECEDENT AKA]" phrase
    # before token replacement when aka is empty — otherwise the literal
    # "aka " run remains in the caption with nothing after it. Run-level
    # scrub preserves the bold "NOTICE OF PROBATE" run that lives in the
    # same paragraph.
    if not aka:
        _aka_phrase_rx = _re.compile(r"(?i)(?:a/k/a|aka)\s*\[DECEDENT AKA\]\s*")
        def _scrub_para(p):
            if not _aka_phrase_rx.search(p.text):
                return
            # Pass 1: scrub within any single run that carries the full match
            for run in p.runs:
                if _aka_phrase_rx.search(run.text):
                    run.text = _aka_phrase_rx.sub("", run.text)
            # Pass 2: split across runs — clear runs that contributed to
            # the phrase. Each fragment is either "aka"/"a/k/a"
            # (case-insensitive, with surrounding whitespace) or
            # "[DECEDENT AKA]" with surrounding whitespace.
            if _aka_phrase_rx.search(p.text):
                _frag_rx = _re.compile(
                    r"(?i)^\s*(?:a/k/a|aka)\s*$|^\s*\[DECEDENT AKA\]\s*$"
                )
                for run in p.runs:
                    if _frag_rx.match(run.text):
                        run.text = ""
        for p in doc.paragraphs:
            _scrub_para(p)
        for t in doc.tables:
            for row in t.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        _scrub_para(p)

    # Caption + token-style replacements
    replace_in_doc(doc, {
        "[COUNTY]":       county.upper(),
        "[DECEDENT]":     dec,
        "[DECEDENT AKA]": aka,
        "[county]":       county,
        "[Petitioner]":   pet,
    })

    # Regex replacements over body paragraphs (underscored blanks)
    patterns = [
        (_re.compile(r"The Will dated _+"),
         f"The Will dated {will_date}"),
        (_re.compile(r"domiciled at _+(?=, County of)"),
         f"domiciled at {dec_addr}"),
        (_re.compile(r"whose address is _+(?=\.)"),
         f"whose address is {pet_addr}"),
        (_re.compile(r"Dated:\s*_+"),
         f"Dated: {today_str}"),
        # Affidavit of mailing — re-stamp the will date there too
        (_re.compile(r"copy of the Will dated _+"),
         f"copy of the Will dated {will_date}"),
    ]

    def _rewrite_para(p, rx, repl):
        """Substitute `rx` → `repl` inside paragraph p, preserving run
        formatting when the match fits within a single run. Same approach
        as replace_in_doc(): only collapse runs into runs[0] when the
        match is split across runs (there's no unambiguous way to
        attribute formatting to the rewritten span in that case)."""
        if not rx.search(p.text):
            return
        # Pass 1: replace inside each run that already contains the match.
        # Bold / italic / underline / font on these runs is preserved.
        for run in p.runs:
            if rx.search(run.text):
                run.text = rx.sub(repl, run.text)
        # Pass 2: if the pattern still matches at paragraph level, the
        # match crossed a run boundary — fall back to consolidation.
        if rx.search(p.text):
            full = rx.sub(repl, p.text)
            if p.runs:
                p.runs[0].text = full
                for r in p.runs[1:]:
                    r.text = ""
            else:
                p.add_run(full)

    def _apply_patterns(p):
        if not p.text:
            return
        for rx, repl in patterns:
            _rewrite_para(p, rx, repl)

    for p in doc.paragraphs:
        _apply_patterns(p)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    _apply_patterns(p)

    # ── Build recipients list from the canonical interested-persons helper ───
    # Fuzzy match key: (first_token, last_meaningful_token) lowercased, with
    # generational suffixes stripped. So "Amy Sue Nathan" and "Amy Nathan"
    # both reduce to ("amy", "nathan") and are treated as the same person.
    _SUFFIXES = {"jr", "sr", "ii", "iii", "iv", "v"}
    def _name_key(name):
        if not name:
            return None
        toks = [t for t in _re.split(r'\s+', name.strip()) if t]
        while toks and toks[-1].lower().rstrip('.') in _SUFFIXES:
            toks.pop()
        if not toks:
            return None
        first = toks[0].lower()
        last = toks[-1].lower() if len(toks) > 1 else ""
        return (first, last)

    # Petitioner key — they don't get noticed about their own filing
    pet_key = _name_key(pet)

    # Compute the canonical interested-persons list — same helper the petition
    # uses, so the two documents render identical "Nature of Interest" text.
    # Letters always issued to the petitioner's full legal name. One
    # source of truth — data.lettersTo is intentionally ignored, since
    # using it as an override produced inconsistencies between the
    # petition caption (full name) and the "Letters Testamentary to:"
    # line / waiver "be issued to" (short name).
    letters_to = petitioner_full(data)
    persons = compute_interested_persons(data, pet, pet_addr, letters_to)

    recipients = []
    seen_keys = set()
    if pet_key:
        seen_keys.add(pet_key)

    for p in persons:
        nm = (p.get("name") or "").strip()
        if not nm:
            continue
        k = _name_key(nm)
        if not k or k in seen_keys:
            continue
        seen_keys.add(k)
        addr = (p.get("address") or "").strip() or "___________"
        nature = (p.get("interest") or "").strip()
        recipients.append((nm, addr, nature))

    # ── Replace the recipient table with a clean 3-column structure ──────────
    # The template's table 0 has inconsistent gridSpan across rows (row 0:
    # [span2,span1,span1]; rows 1-2: [span1,span2,span1]) and column 3 holds
    # explanatory bracket text. Rather than fight the merge layout, we strip
    # the table's rows and grid and rebuild as a flat 3-column table:
    # NAME | MAILING ADDRESS | NATURE OF INTEREST OR STATUS.
    if doc.tables and recipients:
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        def _make_cell(text, bold=False):
            tc = OxmlElement('w:tc')
            tcPr = OxmlElement('w:tcPr')
            # Word ignores tblGrid column widths unless each cell also
            # carries a <w:tcW> width — without it Word auto-sizes columns
            # to content, collapsing the table layout.
            tcW = OxmlElement('w:tcW')
            tcW.set(qn('w:w'), '3000')
            tcW.set(qn('w:type'), 'dxa')
            tcPr.append(tcW)
            tc.append(tcPr)
            p = OxmlElement('w:p')
            r = OxmlElement('w:r')
            if bold:
                rPr = OxmlElement('w:rPr')
                rPr.append(OxmlElement('w:b'))
                r.append(rPr)
            t = OxmlElement('w:t')
            t.text = text
            t.set(qn('xml:space'), 'preserve')
            r.append(t)
            p.append(r)
            tc.append(p)
            return tc

        def _make_row(values, bold=False):
            tr = OxmlElement('w:tr')
            for v in values:
                tr.append(_make_cell(v, bold=bold))
            return tr

        tbl = doc.tables[0]._tbl
        # Strip existing rows
        for tr in list(tbl.findall(qn('w:tr'))):
            tbl.remove(tr)
        # Replace tblGrid with a 3-column equal-width grid
        old_grid = tbl.find(qn('w:tblGrid'))
        if old_grid is not None:
            tbl.remove(old_grid)
        new_grid = OxmlElement('w:tblGrid')
        for _ in range(3):
            gc = OxmlElement('w:gridCol')
            gc.set(qn('w:w'), '3000')
            new_grid.append(gc)
        tbl_pr = tbl.find(qn('w:tblPr'))
        if tbl_pr is not None:
            tbl_pr.addnext(new_grid)
        else:
            tbl.insert(0, new_grid)
        # Recipient rows. The template paragraph immediately above the
        # table already prints a tab-separated "NAME / MAILING ADDRESS /
        # NATURE OF INTEREST OR STATUS" header — adding another one
        # inside the table doubled the header on every output.
        for nm, addr, nature in recipients:
            tbl.append(_make_row([nm, addr, nature]))

    _validate_docx(doc, "generate_notice_of_probate")
    return make_docx_bytes(doc)


def generate_affidavit_of_comparison(data):
    """Generate the Affidavit of Comparison (P-13) Word document.

    Template placeholders (Affiv of Comparison.docx):
    - [COUNTY]:        Caption county (uppercase)
    - [DECEDENT]:      Decedent full name
    - [DECEDENT AKA]:  a/k/a
    - [FILE NO]:       File number
    - [STATE]:         Jurat state (defaults to NEW YORK)
    - [county]:        Jurat county (lowercase placeholder)
    - [AFFIANT]:       Person executing the affidavit (defaults to attorney name)
    """
    doc = Document(os.path.join(WORD_TEMPLATES_DIR, "Affiv of Comparison.docx"))
    dec = decedent_full(data)
    county = data.get("county", "")
    affiant = data.get("affiant") or data.get("attorneyName") or "Jessica Wilson, Esq."

    replace_in_doc(doc, {
        "[COUNTY]":       county.upper(),
        "[DECEDENT]":     dec,
        "[DECEDENT AKA]": data.get("decedentAKA", ""),
        "[FILE NO]":      data.get("fileNo", ""),
        "[STATE]":        data.get("affiantState", "NEW YORK"),
        "[county]":       data.get("affiantCounty", county),
        "[AFFIANT]":      affiant,
    })

    _validate_docx(doc, "generate_affidavit_of_comparison")
    return make_docx_bytes(doc)


def generate_bond_affidavit(data):
    """Generate the Bond Affidavit Word document.

    Template uses hardcoded sample data (PHILLIP WILSON-CAMHI / KINGS).
    We replace sample names/values with actual case data.

    Template placeholders (Bond_Affidavit.docx):
    - COUNTY OF KINGS:                     County
    - PHILLIP WILSON-CAMHI:                Petitioner name (appears 3x)
    - 9 Mills Road, Stony Brook, New York 11790: Petitioner address
    - 31 years of age:                     Decedent age at death
    - February 2018:                       Month/Year for notary
    """
    doc = Document(os.path.join(WORD_TEMPLATES_DIR, "Bond_Affidavit.docx"))
    pet = petitioner_full(data)
    county = data.get("county", "")

    # Calculate decedent age at death
    age_str = ""
    try:
        dob = data.get("decedentDOB", "")
        dod = data.get("decedentDOD", "")
        if dob and dod:
            from datetime import datetime as _dt
            dt_dob = _dt.strptime(dob, "%m/%d/%Y")
            dt_dod = _dt.strptime(dod, "%m/%d/%Y")
            age = dt_dod.year - dt_dob.year - (
                (dt_dod.month, dt_dod.day) < (dt_dob.month, dt_dob.day))
            age_str = str(age)
    except Exception:
        pass

    # Build relationship string
    pet_rel = data.get("petitionerRelationship", "Distributee")

    replace_in_doc(doc, {
        "COUNTY OF KINGS":         f"COUNTY OF {county.upper()}",
        "PHILLIP WILSON-CAMHI":    pet,
        "9 Mills Road, Stony Brook, New York 11790": ", ".join(filter(None, [
            data.get("petitionerStreet", ""),
            data.get("petitionerCity", ""),
            data.get("petitionerState", ""),
            data.get("petitionerZip", ""),
        ])),
        "31 years of age":         f"{age_str} years of age" if age_str else "__ years of age",
        "Distributee of said deceased": f"{pet_rel} of said deceased",
        "February 2018":           datetime.now().strftime("%B %Y"),
    })

    _validate_docx(doc, "generate_bond_affidavit")
    return make_docx_bytes(doc)


def generate_petition_scpa_2203(data):
    """Generate the Petition SCPA 2203 (Voluntary Accounting) Word document.

    This template uses hardcoded sample data. We replace the county header.
    Most fields require manual completion as the template is a filled sample.

    Template: Petition_SCPA_2203.docx
    """
    doc = Document(os.path.join(WORD_TEMPLATES_DIR, "Petition_SCPA_2203.docx"))
    county = data.get("county", "")

    replace_in_doc(doc, {
        "COUNTY OF BRONX": f"COUNTY OF {county.upper()}",
    })

    _validate_docx(doc, "generate_petition_scpa_2203")
    return make_docx_bytes(doc)


# ─── REFUNDING AGREEMENT (Receipt, Release, Indemnification & Refunding) ──────


def generate_refunding_agreement(data):
    """Generate the Receipt, Release, Indemnification & Refunding Agreement.

    Template: RRI_Refunding_Agreement.docx (converted from legacy .doc)

    Auto-fills case header info (county, decedent, executor, date of death).
    Bracketed optional clauses (e.g. [WHEREAS...], [his/her]) are left as-is
    for the attorney to select/edit manually in Word.

    Placeholders replaced:
    - COUNTY OF SUFFOLK          → actual county
    - DECEDENT (in header/body)  → decedent full name
    - EXECUTOR (in header/body)  → petitioner/executor name
    - "died on DATE"             → date of death (long format)
    - "County of COUNTY"         → county name
    - EXEC (commission para)     → executor name
    - BENE1 / BENE 1            → first distributee name (if available)
    """
    doc = Document(os.path.join(WORD_TEMPLATES_DIR, "RRI_Refunding_Agreement.docx"))

    dec = decedent_full(data)
    pet = petitioner_full(data)
    county = data.get("county", "")
    dod = data.get("decedentDOD", "")
    dod_long = format_date_long(dod) if dod else "________"

    # Get first distributee name if available
    dists = data.get("distributees", [])
    bene1_name = ""
    if dists:
        bene1_name = " ".join(filter(None, [
            dists[0].get("firstName", ""),
            dists[0].get("middleName", ""),
            dists[0].get("lastName", ""),
        ]))

    replacements = {
        "COUNTY OF SUFFOLK":  f"COUNTY OF {county.upper()}" if county else "COUNTY OF __________",
        "died on DATE":       f"died on {dod_long}",
        "County of COUNTY":   f"County of {county}" if county else "County of __________",
    }

    # Replace DECEDENT — but only the standalone placeholder, not inside
    # "Decedent" (which appears as a defined term in the body)
    if dec:
        replacements["DECEDENT,"] = f"{dec.upper()},"
        replacements["DECEDENT, (the"] = f"{dec.upper()}, (the"
        replacements["of EXECUTOR, as Executor"] = f"of {pet.upper()}, as Executor"
        replacements["EXECUTOR was appointed"] = f"{pet.upper()} was appointed"

    # Executor signature block and notary
    if pet:
        replacements["EXEC individually"] = f"{pet.upper()} individually"

    # Beneficiary name fill (first bene only — others need manual entry)
    if bene1_name:
        replacements["BENE1  hereby"] = f"{bene1_name.upper()}  hereby"
        replacements["BENE 1"] = bene1_name.upper()

    replace_in_doc(doc, replacements)

    _validate_docx(doc, "generate_refunding_agreement")
    return make_docx_bytes(doc)


# ─── FORMAL ACCOUNTING (Judicial Settlement) ─────────────────────────────────


def generate_formal_accounting(form_data, entries):
    """Generate a formal accounting document (Word) matching Surrogate's Court format.

    Produces cover page, summary statement, and Schedules A through K.
    """
    from docx.enum.section import WD_ORIENT

    doc = Document()

    # ── Page setup ────────────────────────────────────────────────────────────
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)

    dec = decedent_full(form_data)
    aka = form_data.get("decedentAKA", "")
    pet = petitioner_full(form_data)
    county = form_data.get("county", "")
    dod = form_data.get("decedentDOD", "")
    dod_long = format_date_long(dod) if dod else "________"
    file_no = form_data.get("fileNo", "")
    proc = form_data.get("proceedingType", "Administration")
    role = "Executor" if proc == "Probate" else "Administrator"

    # Group entries by schedule
    by_sched = {}
    for e in entries:
        s = e.get("schedule", "")
        by_sched.setdefault(s, []).append(e)

    def sched_total(s):
        return sum(float(e.get("amount", 0) or 0) for e in by_sched.get(s, []))

    def add_para(text, bold=False, size=12, alignment=None, space_after=6):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = bold
        run.font.size = Pt(size)
        run.font.name = "Times New Roman"
        if alignment is not None:
            p.alignment = alignment
        p.paragraph_format.space_after = Pt(space_after)
        return p

    def money(val):
        try:
            v = float(val or 0)
            return f"${v:,.2f}"
        except (ValueError, TypeError):
            return "$0.00"

    # ── COVER PAGE ────────────────────────────────────────────────────────────
    add_para("SURROGATE'S COURT OF THE STATE OF NEW YORK",
             bold=True, size=12, alignment=WD_ALIGN_PARAGRAPH.LEFT)
    add_para(f"COUNTY OF {county.upper()}" if county else "COUNTY OF __________",
             bold=True, size=12, alignment=WD_ALIGN_PARAGRAPH.LEFT)

    # Caption box
    dec_display = dec.upper()
    if aka:
        dec_display += f", a/k/a\n{aka.upper()}"

    caption_left = (
        f"In the Matter of the Judicial Settlement of the Final Account of\n\n"
        f"{pet.upper()}, as {role}\n\n"
        f"of the Estate of\n\n"
        f"{dec_display},\n\n"
        f"{'':>40}Deceased."
    )
    add_para(caption_left, size=12)

    file_line = f"File No:    {file_no}" if file_no else "File No:    __________"
    add_para(file_line, size=12, alignment=WD_ALIGN_PARAGRAPH.RIGHT)

    # Accounting type
    add_para(f"ACCOUNTING BY:\n  {role}", size=11, space_after=12)

    # Court address and period
    add_para(f"TO THE SURROGATE'S COURT OF THE COUNTY OF {county.upper() if county else '__________'}:",
             bold=True, size=11, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
    add_para(
        f"The undersigned does hereby render the account of the proceedings as follows:\n"
        f"Period of account from {dod_long} to {today()}\n"
        f"This is a first and final account containing the following schedules.",
        size=11, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=12
    )

    # ── TABLE OF CONTENTS ─────────────────────────────────────────────────────
    toc_items = [
        ("A", "Principal Received"),
        ("AA", "Subsequent Receipts of Principal"),
        ("A-1", "Realized Increases"),
        ("A-2", "Income Collected"),
        ("B", "Realized Decreases"),
        ("C", "Funeral and Administration Expenses and Taxes"),
        ("C-1", "Unpaid Administration Expenses"),
        ("D", "Creditors' Claims"),
        ("E", "Distributions of Principal"),
        ("F", "New Investments, Exchanges and Stock Distributions"),
        ("G", "Principal Remaining on Hand"),
        ("H", "Interested Parties"),
        ("I", "Computation of Commissions"),
        ("J", "Other Pertinent Facts and Cash Reconciliation"),
        ("K", "Estate Taxes Paid and Allocation of Estate Taxes"),
    ]

    add_para("PRINCIPAL", bold=True, size=11, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
    for sched, title in toc_items[:11]:
        add_para(f"Schedule {sched}        {title}", size=11, space_after=2)
    add_para("OTHER", bold=True, size=11, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
    for sched, title in toc_items[11:]:
        add_para(f"Schedule {sched}        {title}", size=11, space_after=2)

    doc.add_page_break()

    # ── SUMMARY STATEMENT ─────────────────────────────────────────────────────
    add_para("SUMMARY STATEMENT", bold=True, size=12,
             alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
    add_para("COMBINED ACCOUNT", bold=True, size=12,
             alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)

    tot_a = sched_total("A")
    tot_aa = sched_total("AA")
    tot_a1 = sched_total("A-1")
    tot_a2 = sched_total("A-2")
    tot_b = sched_total("B")
    tot_c = sched_total("C")
    tot_c1 = sched_total("C-1")
    tot_d = sched_total("D")
    tot_e = sched_total("E")
    tot_g = sched_total("G")

    # Unrealized: G inventory vs market
    unreal_inc = 0
    unreal_dec = 0
    for e in by_sched.get("G", []):
        inv = float(e.get("inventory_value", 0) or 0)
        mkt = float(e.get("market_value", 0) or float(e.get("amount", 0) or 0))
        diff = mkt - inv
        if diff > 0:
            unreal_inc += diff
        else:
            unreal_dec += abs(diff)

    charges = tot_a + tot_aa + tot_a1 + tot_a2 + unreal_inc
    credits = tot_b + tot_c + tot_d + tot_e + unreal_dec
    balance = charges - credits

    # Charges table
    add_para("CHARGES:", bold=True, size=11, space_after=4)
    charge_items = [
        ('Schedule "A"', "(Principal Received)", tot_a),
        ('Schedule "AA"', "(Subsequent Receipts)", tot_aa),
        ('Schedule "A-1"', "(Realized Increases)", tot_a1),
        ('Schedule "A-2"', "(Income Collected)", tot_a2),
        ('Schedule "G"', "(Unrealized Increases)", unreal_inc),
    ]
    for label, desc, val in charge_items:
        add_para(f"{label:20s} {desc:40s} {money(val):>15s}", size=11, space_after=1)
    add_para(f"{'Total Charges':20s} {'':40s} {money(charges):>15s}", bold=True, size=11, space_after=8)

    add_para("CREDITS:", bold=True, size=11, space_after=4)
    credit_items = [
        ('Schedule "B"', "(Realized Decreases)", tot_b),
        ('Schedule "C"', "(Funeral and Administration Expenses)", tot_c),
        ('Schedule "D"', "(Creditors' Claims Actually Paid)", tot_d),
        ('Schedule "E"', "(Distributions)", tot_e),
        ('Schedule "G"', "(Unrealized Decreases)", unreal_dec),
    ]
    for label, desc, val in credit_items:
        add_para(f"{label:20s} {desc:40s} {money(val):>15s}", size=11, space_after=1)
    add_para(f"{'Total Credits':20s} {'':40s} {money(credits):>15s}", bold=True, size=11, space_after=4)
    bal_label = 'Balance on Hand Shown by Schedule "G"'
    add_para(f"{bal_label:40s} {money(balance):>15s}",
             bold=True, size=11, space_after=8)

    doc.add_page_break()

    # ── SUMMARY NARRATIVE ─────────────────────────────────────────────────────
    add_para("SUMMARY STATEMENT", bold=True, size=12,
             alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)
    add_para(
        f"The foregoing balance of {money(balance)} consists of "
        f"cash and other property on hand as of {today()}. "
        f"It is subject to deductions of estimated principal commissions "
        f"amounting to {money(_calc_commission(charges))}, "
        f"shown in Schedule I and to the proper charge to principal of expenses of this "
        f"accounting.",
        size=11, space_after=12
    )
    add_para("The attached schedules are part of this account.",
             size=11, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)
    add_para(f"{'_' * 35}", size=11, space_after=2)
    add_para(pet, size=11, space_after=2)
    add_para(role, size=11, space_after=0)

    doc.add_page_break()

    # ── SCHEDULE GENERATION HELPER ────────────────────────────────────────────
    def add_schedule(sched_id, title, subtitle, cols, amt_col="amount"):
        """Add a schedule section with header and entry table."""
        estate_header = f"Estate of {dec}"
        if aka:
            estate_header += f", aka {aka}"
        add_para(estate_header, bold=True, size=11,
                 alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
        add_para(f"Schedule {sched_id}", bold=True, size=11,
                 alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
        add_para(subtitle, bold=True, size=11,
                 alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)

        sched_entries = by_sched.get(sched_id, [])

        if not sched_entries:
            add_para("None", size=11, space_after=6)
            total = 0
        else:
            # Build table
            table = doc.add_table(rows=1, cols=len(cols))
            table.style = "Table Grid"

            # Header row
            for i, (hdr, _) in enumerate(cols):
                cell = table.rows[0].cells[i]
                cell.text = hdr
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.bold = True
                        r.font.size = Pt(10)
                        r.font.name = "Times New Roman"

            total = 0
            for e in sched_entries:
                row = table.add_row()
                for i, (_, field) in enumerate(cols):
                    val = e.get(field, "") or ""
                    if field == amt_col or field in ("inventory_value", "market_value", "amount"):
                        try:
                            val = money(float(val or 0))
                        except (ValueError, TypeError):
                            val = ""
                    cell = row.cells[i]
                    cell.text = str(val)
                    for p in cell.paragraphs:
                        for r in p.runs:
                            r.font.size = Pt(10)
                            r.font.name = "Times New Roman"
                amt = float(e.get(amt_col, 0) or 0)
                total += amt

        add_para(f"\nTotal Schedule {sched_id}:  {money(total)}",
                 bold=True, size=11, space_after=6)

        doc.add_page_break()
        return total

    # ── GENERATE ALL SCHEDULES ────────────────────────────────────────────────
    add_schedule("A", "Schedule A", "Receipts",
                 [("Description", "description"), ("Institution", "institution"),
                  ("Inventory Value", "amount")])

    add_schedule("AA", "Schedule AA", "Statement of Subsequent Receipts of Principal",
                 [("Date Received", "date"), ("Description", "description"),
                  ("Inventory Value", "amount")])

    add_schedule("A-1", "Schedule A-1",
                 "Statement of Increases on Sales, Liquidation or Distribution",
                 [("Description", "description"),
                  ("Proceeds", "amount"), ("Inventory Value", "inventory_value")])

    add_schedule("A-2", "Schedule A-2", "Statement of All Income Collected",
                 [("Date", "date"), ("Description", "description"),
                  ("Institution", "institution"), ("Amount", "amount")])

    add_schedule("B", "Schedule B",
                 "Statement of Decreases Due to Sales, Liquidation, Collection, Distribution, or Uncollectibility",
                 [("Date", "date"), ("Description", "description"),
                  ("Proceeds", "amount"), ("Inventory Value", "inventory_value")])

    add_schedule("C", "Schedule C",
                 "Statement of Funeral and Administration Expenses and Taxes",
                 [("Date", "date"), ("Description", "description"), ("Amount", "amount")])

    add_schedule("C-1", "Schedule C-1", "Statement of Unpaid Administration Expenses",
                 [("Description", "description"), ("Amount", "amount")])

    add_schedule("D", "Schedule D", "Statement of All Creditors' Claims",
                 [("Description", "description"), ("Amount", "amount")])

    add_schedule("E", "Schedule E", "Distributions",
                 [("Description", "description"), ("Distribution Value", "amount")])

    add_schedule("F", "Schedule F",
                 "Statement of New Investments, Exchanges and Stock Distributions",
                 [("Date", "date"), ("Description", "description"),
                  ("Shares", "shares"), ("Inventory Value", "amount")])

    add_schedule("G", "Schedule G", "Balance On Hand",
                 [("Description", "description"), ("Shares", "shares"),
                  ("Market Value", "market_value"), ("Inventory Value", "inventory_value")])

    # ── SCHEDULE H — Interested Parties ───────────────────────────────────────
    estate_header = f"Estate of {dec}"
    if aka:
        estate_header += f", aka {aka}"
    add_para(estate_header, bold=True, size=11,
             alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    add_para("Schedule H", bold=True, size=11,
             alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    add_para("Statement of Interested Parties", bold=True, size=11,
             alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)

    h_entries = by_sched.get("H", [])
    if h_entries:
        table = doc.add_table(rows=1, cols=3)
        table.style = "Table Grid"
        for i, hdr in enumerate(["Name and Post Office Address", "Relationship", "Nature of Interest"]):
            cell = table.rows[0].cells[i]
            cell.text = hdr
            for p in cell.paragraphs:
                for r in p.runs:
                    r.bold = True
                    r.font.size = Pt(10)
                    r.font.name = "Times New Roman"
        for e in h_entries:
            row = table.add_row()
            row.cells[0].text = f"{e.get('description', '')}\n{e.get('institution', '')}"
            row.cells[1].text = e.get("category", "")
            row.cells[2].text = "Distributee"
            for cell in row.cells:
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.font.size = Pt(10)
                        r.font.name = "Times New Roman"
    else:
        add_para("None", size=11)

    add_para(
        "\nThe records of this Court have been searched for powers of attorney and "
        "assignments and encumbrances made and executed by any of the persons interested "
        "in or entitled to share in the estate. No such powers of attorney, assignments "
        "or encumbrances were found to have been filed or recorded in this Court, and the "
        "accounting party has no knowledge of the execution of any such power of attorney, "
        "assignment or encumbrance that is not so filed and recorded.",
        size=10, space_after=6
    )

    doc.add_page_break()

    # ── SCHEDULE I — Commission Computation ───────────────────────────────────
    add_para(estate_header, bold=True, size=11,
             alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    add_para("Schedule I", bold=True, size=11,
             alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    add_para("Statement of Computation of Commissions", bold=True, size=11,
             alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)

    # Receiving commission — use net equity (gross minus liens/mortgages)
    total_liens = sum(float(e.get("lien_amount", 0) or 0)
                      for e in by_sched.get("A", []) + by_sched.get("AA", []))
    net_principal = tot_a + tot_aa - total_liens

    add_para("For Receiving Principal", bold=True, size=11, space_after=4)
    recv_base = net_principal + tot_a1 + tot_a2 + unreal_inc
    add_para(f"Principal Received (Schedule A + AA)     {money(tot_a + tot_aa)}", size=11, space_after=1)
    if total_liens > 0:
        add_para(f"Less: Liens/Mortgages on Real Property  ({money(total_liens)})", size=11, space_after=1)
        add_para(f"Net Equity                              {money(net_principal)}", size=11, space_after=1)
    add_para(f"Increases on Principal (Schedule A-1)    {money(tot_a1)}", size=11, space_after=1)
    add_para(f"Income Collected (Schedule A-2)          {money(tot_a2)}", size=11, space_after=1)
    add_para(f"Unrealized Increases (Schedule G)        {money(unreal_inc)}", size=11, space_after=4)
    add_para(f"Commission Base                          {money(recv_base)}", bold=True, size=11, space_after=6)

    tiers = [
        (0.05, 100000), (0.04, 200000), (0.03, 700000), (0.025, float('inf')),
    ]
    remaining = recv_base
    recv_comm = 0
    for rate, bracket in tiers:
        if remaining <= 0:
            break
        base = min(remaining, bracket)
        comm = base * rate
        recv_comm += comm
        pct = int(rate * 100) if rate * 100 == int(rate * 100) else rate * 100
        add_para(f"  {pct}% on {money(base):>20s} = {money(comm):>15s}", size=11, space_after=1)
        remaining -= base

    recv_half = recv_comm / 2
    add_para(f"\n1/2 Thereof for Receiving               {money(recv_half)}",
             bold=True, size=11, space_after=8)

    # Paying commission
    paying_base = tot_c + tot_d + tot_e + tot_g
    add_para("For Paying Principal", bold=True, size=11, space_after=4)
    add_para(f"Funeral and Administration Expenses (Schedule C)  {money(tot_c)}", size=11, space_after=1)
    add_para(f"Payment of Debts (Schedule D)                     {money(tot_d)}", size=11, space_after=1)
    add_para(f"Distributions of Principal (Schedule E)           {money(tot_e)}", size=11, space_after=1)
    add_para(f"Principal on Hand (Schedule G)                    {money(tot_g)}", size=11, space_after=4)
    add_para(f"Total Principal                                   {money(paying_base)}",
             bold=True, size=11, space_after=6)

    remaining = paying_base
    pay_comm = 0
    for rate, bracket in tiers:
        if remaining <= 0:
            break
        base = min(remaining, bracket)
        comm = base * rate
        pay_comm += comm
        pct = int(rate * 100) if rate * 100 == int(rate * 100) else rate * 100
        add_para(f"  {pct}% on {money(base):>20s} = {money(comm):>15s}", size=11, space_after=1)
        remaining -= base

    pay_half = pay_comm / 2
    add_para(f"\n1/2 Thereof for Paying                  {money(pay_half)}",
             bold=True, size=11, space_after=12)

    total_comm = recv_half + pay_half
    add_para(f"Total Commissions Due Each {role}", bold=True, size=11, space_after=4)
    add_para(f"  Receiving     {money(recv_half)}", size=11, space_after=1)
    add_para(f"  Paying        {money(pay_half)}", size=11, space_after=4)
    add_para(f"  Total         {money(total_comm)}", bold=True, size=11, space_after=6)
    add_para(f"\nTotal commissions available for allocation:   {money(total_comm)}",
             bold=True, size=11, space_after=6)

    doc.add_page_break()

    # ── SCHEDULE J — Cash Reconciliation ──────────────────────────────────────
    add_para(estate_header, bold=True, size=11,
             alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    add_para("Schedule J", bold=True, size=11,
             alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    add_para("Statement of Other Pertinent Facts and Cash Reconciliation", bold=True,
             size=11, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)

    add_para("Other Pertinent Facts", bold=True, size=11, space_after=4)
    add_para("None", size=11, space_after=8)

    add_para("Reconciliation of Cash and Other Assets", bold=True, size=11, space_after=8)

    recon_items = [
        ("Schedule A", "Receipts", tot_a, "CREDITS"),
        ("Schedule AA", "Subsequent Receipts", tot_aa, "CREDITS"),
        ("Schedule A-2", "Income Collected", tot_a2, "CREDITS"),
        ("Schedule B", "Proceeds on Sales, Etc.", tot_b, "DEBITS"),
        ("Schedule C", "Admin/Funeral Expenses", tot_c, "DEBITS"),
        ("Schedule F", "Purchases, Etc.", sched_total("F"), "DEBITS"),
        ("Schedule G", "On Hand", tot_g, "DEBITS"),
    ]

    add_para(f"{'':30s} {'DEBITS':>15s} {'CREDITS':>15s}", bold=True, size=11, space_after=4)
    cash_debits = 0
    cash_credits = 0
    for label, desc, val, side in recon_items:
        debit = money(val) if side == "DEBITS" else ""
        credit = money(val) if side == "CREDITS" else ""
        if side == "DEBITS":
            cash_debits += val
        else:
            cash_credits += val
        add_para(f"{label:12s} {desc:18s} {debit:>15s} {credit:>15s}", size=11, space_after=1)

    add_para(f"\n{'Total':30s} {money(cash_debits):>15s} {money(cash_credits):>15s}",
             bold=True, size=11, space_after=6)

    doc.add_page_break()

    # ── SCHEDULE K — Estate Taxes ─────────────────────────────────────────────
    add_schedule("K", "Schedule K",
                 "Statement of Estate Taxes Paid and Allocation Thereof",
                 [("Description", "description"), ("Amount", "amount")])

    _validate_docx(doc, "generate_formal_accounting")
    return make_docx_bytes(doc)


# ─── SMALL ESTATE / VOLUNTARY ADMINISTRATION (SCPA Article 13) ───────────────
# Official 11/2019 court forms shipped as underscore-blank Word documents in
# templates/Small Estate/. Filled with the same regex-substitution approach as
# the P-4 waiver: match a labeled blank, replace inside the run that holds it,
# collapsing to the first run only when the blank is split across runs.

SMALL_ESTATE_DIR = os.path.join(TEMPLATES_DIR, "Small Estate")


def _se_sub(p, rx, repl):
    """Regex-substitute inside paragraph p (first match only)."""
    if not rx.search(p.text):
        return False
    for run in p.runs:
        if rx.search(run.text):
            run.text = rx.sub(repl, run.text, count=1)
            return True
    full = rx.sub(repl, p.text, count=1)
    if p.runs:
        p.runs[0].text = full
        for r in p.runs[1:]:
            r.text = ""
    else:
        p.add_run(full)
    return True


def _se_sub_collapse(p, rx, repl):
    """Like _se_sub but always substitutes on the full paragraph text and
    collapses to the first run — for blanks whose underscores span runs."""
    if not rx.search(p.text):
        return False
    full = rx.sub(repl, p.text, count=1)
    if p.runs:
        p.runs[0].text = full
        for r in p.runs[1:]:
            r.text = ""
    else:
        p.add_run(full)
    return True


_SE_BLANK = re.compile(r"_{3,}")


def _se_fill_blanks(p, values):
    """Replace the Nth underscore-run in paragraph p with values[N].
    A None/empty value keeps that blank as underscores. Collapses the
    paragraph into its first run (uniform formatting on these forms)."""
    if not _SE_BLANK.search(p.text):
        return
    vals = list(values)

    def _repl(m):
        if not vals:
            return m.group(0)
        v = vals.pop(0)
        return str(v) if v not in (None, "") else m.group(0)

    full = _SE_BLANK.sub(_repl, p.text)
    if full == p.text:
        return
    if p.runs:
        p.runs[0].text = full
        for r in p.runs[1:]:
            r.text = ""
    else:
        p.add_run(full)


def _se_check(p, which=0):
    """Mark the (which)th '[    ]' checkbox in paragraph p with an X."""
    boxes = list(re.finditer(r"\[\s+\]", p.text))
    if which >= len(boxes):
        return
    b = boxes[which]
    full = p.text[:b.start()] + "[ X ]" + p.text[b.end():]
    if p.runs:
        p.runs[0].text = full
        for r in p.runs[1:]:
            r.text = ""
    else:
        p.add_run(full)


def _se_money(v):
    try:
        return f"{float(str(v).replace(',', '').replace('$', '').strip()):,.2f}"
    except (ValueError, TypeError):
        return str(v or "")


def _se_attorney_block(data):
    return {
        "name": data.get("attorneyName") or "Jessica Wilson, Esq.",
        "firm": data.get("firmName") or "Law Office of Jessica Wilson",
        "phone": data.get("attorneyPhone") or "(212) 739-1736",
        "address": data.get("attorneyAddress") or data.get("firmAddress")
                   or "221 Columbia Street, Brooklyn NY 11231",
    }


def generate_se3a(data):
    """Fill the SE-3A Affidavit in Relation to Settlement of Estate Under
    Article 13, SCPA — the main (and only) petition document for a Small
    Estate / Voluntary Administration proceeding."""
    doc = Document(os.path.join(SMALL_ESTATE_DIR, "Small Estate - Aff SE-3A.docx"))
    P = doc.paragraphs

    county   = (data.get("county") or "").strip()
    file_no  = (data.get("fileNo") or "").strip()
    dec      = decedent_full(data)
    aka      = (data.get("decedentAKA") or "").strip()
    dec_name = f"{dec} a/k/a {aka}" if aka else dec
    pet      = petitioner_full(data)
    rel      = (data.get("petitionerRelationship") or "").strip()
    testate  = bool((data.get("willDate") or "").strip()
                    or (data.get("willBeneficiaries") or []))

    # Caption + venue
    _se_sub(P[1],  re.compile(r"COUNTY OF\s*_+"), f"COUNTY OF  {county.upper()}")
    _se_fill_blanks(P[7],  [dec_name, file_no])
    _se_sub_collapse(P[11], re.compile(r"STATE OF\s*_+"),  "STATE OF NEW YORK ")
    _se_sub_collapse(P[13], re.compile(r"COUNTY OF\s*_+"), f"COUNTY OF {county.upper()} ")
    _se_fill_blanks(P[15], [pet])

    # (1) Affiant address / email
    _se_fill_blanks(P[17], [f"{data.get('petitionerStreet', '')}"
                            f"                    {data.get('petitionerCity', '')}"])
    _se_fill_blanks(P[20], [f"{data.get('petitionerCounty', '')}"
                            f"                    {data.get('petitionerState', '')}"
                            f"          {data.get('petitionerZip', '')}"
                            f"               {data.get('petitionerPhone', '')}"])
    _se_fill_blanks(P[26], [data.get("petitionerEmail", "")])

    # (2) Interest
    if rel.lower() in ("", "other"):
        _se_check(P[32])
    else:
        _se_check(P[29])
        _se_fill_blanks(P[29], [rel])

    # (3) Decedent
    _se_fill_blanks(P[37], [dec_name])
    _se_fill_blanks(P[39], [f"{data.get('decedentStreet', '')}      "
                            f"{data.get('decedentCity', '')}      "
                            f"{county} County      {data.get('decedentState', '')}"])
    _se_fill_blanks(P[42], [format_date_long(data.get("decedentDOD", "")) + "      ",
                            data.get("decedentPlaceOfDeath", "")])
    _se_fill_blanks(P[45], [data.get("decedentCitizenship", "") or "United States"])

    # (4) Testate / intestate
    _se_check(P[49] if testate else P[48])

    # (6) Distributees (4 rows on the form)
    dists = [d for d in (data.get("distributees") or []) if d.get("name")]
    for row_idx, para_idx in enumerate((69, 71, 73, 75)):
        if row_idx < len(dists):
            d = dists[row_idx]
            _se_fill_blanks(P[para_idx], [d.get("name", ""),
                                          d.get("address", ""),
                                          d.get("relationship", "")])

    # (7) Will beneficiaries (4 rows) — only when testate
    if testate:
        bens = [b for b in (data.get("willBeneficiaries") or [])
                if b.get("name") and not b.get("deceased")]
        for row_idx, para_idx in enumerate((82, 84, 86, 88)):
            if row_idx < len(bens):
                b = bens[row_idx]
                _se_fill_blanks(P[para_idx], [b.get("name", ""),
                                              b.get("address", ""),
                                              (b.get("interest") or b.get("bequest", ""))])

    # (9) Personal property items (5 rows) + total
    assets = [a for a in (data.get("assets") or [])
              if a.get("institution") and a.get("category") != "Real Estate"]
    total = 0.0
    for row_idx, para_idx in enumerate((101, 103, 105, 107, 109)):
        if row_idx < len(assets):
            a = assets[row_idx]
            label = " — ".join(filter(None, [a.get("institution", ""),
                                             a.get("accountNumber", "")]))
            _se_fill_blanks(P[para_idx], [label, _se_money(a.get("value"))])
            try:
                total += float(str(a.get("value", "0")).replace(",", "").replace("$", ""))
            except (ValueError, TypeError):
                pass
    if total > 0:
        _se_fill_blanks(P[113], [f"{total:,.2f}"])
    elif data.get("personalPropertyValue"):
        _se_fill_blanks(P[113], [_se_money(data.get("personalPropertyValue"))])

    # Signature block — print name under signature line
    _se_fill_blanks(P[149], [pet])

    # Attorney block
    atty = _se_attorney_block(data)
    _se_fill_blanks(P[164], [atty["name"]])
    _se_fill_blanks(P[165], [atty["firm"] + "      ", atty["phone"]])
    _se_fill_blanks(P[166], [atty["address"]])

    _validate_docx(doc, "generate_se3a")
    return make_docx_bytes(doc)


def generate_se1c(data, dist):
    """Fill the SE-1C Renunciation of Voluntary Administration for one
    distributee who is renouncing the right to serve. Signature, notary and
    date blanks stay open for manual completion at signing."""
    doc = Document(os.path.join(SMALL_ESTATE_DIR, "Small Estate - Renunc SE-1C.docx"))
    P = doc.paragraphs

    county  = (data.get("county") or "").strip()
    file_no = (data.get("fileNo") or "").strip()
    dec     = decedent_full(data)
    name    = (dist.get("name") or "").strip()
    rel     = (dist.get("relationship") or "").strip()
    addr    = (dist.get("address") or "").strip()

    _se_sub(P[1], re.compile(r"COUNTY OF\s*_+"), f"COUNTY OF  {county.upper()}")
    # Caption decedent name sits on a mostly-blank line ending in a comma
    _se_sub(P[5], re.compile(r"^\s{20,},"), f"{dec},")
    if file_no:
        _se_sub(P[7], re.compile(r"(File No\.)\s*"), rf"\g<1> {file_no}  ")
    # Domiciliary address goes on the blank line under "...address is"
    if addr:
        target = P[12] if len(P) > 12 else None
        if target is not None and not target.text.strip():
            target.add_run(addr)
    # Distributee checkbox + relationship
    _se_check(P[21])
    if rel:
        _se_sub(P[21], re.compile(r"related as a\s*$"), f"related as a {rel}")
    # Renouncing party print name
    _se_fill_blanks(P[29], [name])
    # Acknowledgment venue
    _se_sub(P[31], re.compile(r"STATE OF\s{10,}"), "STATE OF NEW YORK" + " " * 20)
    _se_sub(P[33], re.compile(r"COUNTY OF\s{10,}"),
            f"COUNTY OF {county.upper()}" + " " * 20)
    # Attorney block
    atty = _se_attorney_block(data)
    _se_sub(P[44], re.compile(r"(Print Name of Attorney:)\s{5,}"),
            rf"\g<1> {atty['name']}")
    _se_sub(P[45], re.compile(r"(Firm Name:)\s{5,}"),
            rf"\g<1> {atty['firm']}      ")
    _se_sub(P[45], re.compile(r"(Tel\. No\.)\s*"), rf"\g<1> {atty['phone']}")
    _se_sub(P[46], re.compile(r"(Address of Attorney:)\s{5,}"),
            rf"\g<1> {atty['address']}")

    _validate_docx(doc, "generate_se1c")
    return make_docx_bytes(doc)


def generate_se1d(data):
    """Fill the caption of the SE-1D Report and Account in Settlement of
    Estate Pursuant to Article 13, SCPA. Item/disbursement rows are left
    blank — this form is completed after distribution with the actual
    receipts, so only the caption, venue, and names are pre-filled."""
    doc = Document(os.path.join(SMALL_ESTATE_DIR,
                                "Small Estate - Report & Acct SE-1D.docx"))
    P = doc.paragraphs

    county  = (data.get("county") or "").strip()
    file_no = (data.get("fileNo") or "").strip()
    dec     = decedent_full(data)
    pet     = petitioner_full(data)

    _se_sub(P[1], re.compile(r"COUNTY OF\s*_+"), f"COUNTY OF  {county.upper()}")
    for p in P[2:10]:
        if _se_sub(p, re.compile(r"^\s{20,},"), f"{dec},"):
            break
    if file_no:
        for p in P[2:12]:
            if _se_sub(p, re.compile(r"(File No\.)\s*$"), rf"\g<1> {file_no}"):
                break
    # Venue of the verification + affiant name on the ",, being duly sworn" line
    for p in P:
        _se_sub_collapse(p, re.compile(r"STATE OF\s*_+"), "STATE OF NEW YORK ")
        _se_sub_collapse(p, re.compile(r"COUNTY OF\s*_+"),
                         f"COUNTY OF {county.upper()} ")
        _se_sub_collapse(p, re.compile(r"^\s*,,"), f"{pet},")

    _validate_docx(doc, "generate_se1d")
    return make_docx_bytes(doc)
